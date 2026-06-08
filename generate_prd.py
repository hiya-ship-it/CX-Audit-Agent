"""
Generate the CX Audit Agent PRD Word Document
following Bajaj Finance Enterprise Agentic Requirement SOP V1.0
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ─────────────────────────────────────────────
# Colour palette
# ─────────────────────────────────────────────
BFL_BLUE   = RGBColor(0x00, 0x33, 0x87)
BFL_ORANGE = RGBColor(0xFF, 0x66, 0x00)
WHITE      = RGBColor(0xFF, 0xFF, 0xFF)
BLACK      = RGBColor(0x00, 0x00, 0x00)
AMBER      = RGBColor(0xBF, 0x36, 0x00)


# ─────────────────────────────────────────────
# Low-level XML helpers
# ─────────────────────────────────────────────

def set_cell_bg(cell, hex_color: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def set_cell_borders(cell, color="D9D9D9"):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side in ['top', 'left', 'bottom', 'right']:
        border = OxmlElement(f'w:{side}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), color)
        tcBorders.append(border)
    tcPr.append(tcBorders)


def para_spacing(para, before=0, after=0):
    pPr = para._p.get_or_add_pPr()
    spacing = OxmlElement('w:spacing')
    spacing.set(qn('w:before'), str(before))
    spacing.set(qn('w:after'), str(after))
    pPr.append(spacing)


def set_indent(para, left=0):
    pPr = para._p.get_or_add_pPr()
    ind = OxmlElement('w:ind')
    ind.set(qn('w:left'), str(left))
    pPr.append(ind)


def set_para_shading(para, fill_hex: str):
    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_hex)
    pPr.append(shd)


# ─────────────────────────────────────────────
# High-level building blocks
# ─────────────────────────────────────────────

def h1(doc, number: str, title: str):
    p = doc.add_paragraph()
    para_spacing(p, before=200, after=80)
    set_para_shading(p, '003387')
    set_indent(p, 120)
    run = p.add_run(f"{number}. {title.upper()}")
    run.bold = True
    run.font.size = Pt(13)
    run.font.color.rgb = WHITE
    return p


def h2(doc, number: str, title: str):
    p = doc.add_paragraph()
    para_spacing(p, before=160, after=60)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    left = OxmlElement('w:left')
    left.set(qn('w:val'), 'single')
    left.set(qn('w:sz'), '18')
    left.set(qn('w:space'), '4')
    left.set(qn('w:color'), 'FF6600')
    pBdr.append(left)
    pPr.append(pBdr)
    set_indent(p, 180)
    run = p.add_run(f"{number}  {title}")
    run.bold = True
    run.font.size = Pt(11.5)
    run.font.color.rgb = BFL_BLUE
    return p


def h3(doc, title: str):
    p = doc.add_paragraph()
    para_spacing(p, before=100, after=40)
    set_indent(p, 360)
    run = p.add_run(title)
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = BFL_ORANGE
    return p


def body(doc, text: str, bullet=False, indent_twips=0):
    if bullet:
        p = doc.add_paragraph(style='List Bullet')
    else:
        p = doc.add_paragraph()
    para_spacing(p, before=40, after=40)
    if indent_twips:
        set_indent(p, indent_twips)
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    return p


def bold_label(doc, label: str, text: str, bullet=False, indent_twips=0):
    if bullet:
        p = doc.add_paragraph(style='List Bullet')
    else:
        p = doc.add_paragraph()
    para_spacing(p, before=40, after=40)
    if indent_twips:
        set_indent(p, indent_twips)
    r1 = p.add_run(label)
    r1.bold = True
    r1.font.size = Pt(10.5)
    r1.font.color.rgb = BFL_BLUE
    r2 = p.add_run(text)
    r2.font.size = Pt(10.5)
    return p


def info_box(doc, text: str, fill="EBF0FA"):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    cell = tbl.cell(0, 0)
    set_cell_bg(cell, fill)
    set_cell_borders(cell, '003387')
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.italic = True
    run.font.size = Pt(10)
    run.font.color.rgb = BFL_BLUE if fill == "EBF0FA" else AMBER
    doc.add_paragraph()


def grid_table(doc, headers, rows_data, col_widths=None):
    n = len(headers)
    tbl = doc.add_table(rows=1 + len(rows_data), cols=n)
    tbl.style = 'Table Grid'
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT

    hrow = tbl.rows[0]
    for i, h in enumerate(headers):
        cell = hrow.cells[i]
        set_cell_bg(cell, '003387')
        set_cell_borders(cell, '003387')
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = WHITE

    for ri, row_data in enumerate(rows_data):
        row = tbl.rows[ri + 1]
        fill = 'F2F2F2' if ri % 2 == 0 else 'FFFFFF'
        for ci, val in enumerate(row_data):
            cell = row.cells[ci]
            set_cell_bg(cell, fill)
            set_cell_borders(cell, 'D9D9D9')
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            run.font.size = Pt(10)

    if col_widths:
        for i, w in enumerate(col_widths):
            for row in tbl.rows:
                row.cells[i].width = Inches(w)
    doc.add_paragraph()
    return tbl


def checklist_table(doc, rows_data):
    headers = ["Category", "Checkpoint", "Yes", "No", "Comments"]
    tbl = doc.add_table(rows=1 + len(rows_data), cols=5)
    tbl.style = 'Table Grid'

    hrow = tbl.rows[0]
    for i, h in enumerate(headers):
        cell = hrow.cells[i]
        set_cell_bg(cell, '003387')
        set_cell_borders(cell, '003387')
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = WHITE

    for ri, (cat, checkpoint) in enumerate(rows_data):
        row = tbl.rows[ri + 1]
        fill = 'F2F2F2' if ri % 2 == 0 else 'FFFFFF'
        for ci, val in enumerate([cat, checkpoint, "", "", ""]):
            cell = row.cells[ci]
            set_cell_bg(cell, fill)
            set_cell_borders(cell, 'D9D9D9')
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if ci >= 2 else WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(val)
            run.font.size = Pt(10)

    col_widths = [1.3, 3.2, 0.55, 0.55, 1.7]
    for i, w in enumerate(col_widths):
        for row in tbl.rows:
            row.cells[i].width = Inches(w)
    doc.add_paragraph()


# ═══════════════════════════════════════════════════════
# BUILD DOCUMENT
# ═══════════════════════════════════════════════════════

doc = Document()

# Page setup
sec = doc.sections[0]
sec.top_margin    = Cm(2.0)
sec.bottom_margin = Cm(2.0)
sec.left_margin   = Cm(2.5)
sec.right_margin  = Cm(2.5)

doc.styles['Normal'].font.name = 'Calibri'
doc.styles['Normal'].font.size = Pt(10.5)

# ──────────────────────────────────────────────────────
# COVER PAGE
# ──────────────────────────────────────────────────────
p = doc.add_paragraph()
para_spacing(p, before=800)

def centered_run(doc, text, size, color, bold=False, underline=False):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para_spacing(p, before=80, after=80)
    r = p.add_run(text)
    r.bold = bold
    r.underline = underline
    r.font.size = Pt(size)
    r.font.color.rgb = color
    return p

centered_run(doc, "BAJAJ FINANCE LIMITED", 22, BFL_BLUE, bold=True, underline=True)
centered_run(doc, "CX Audit Agent", 20, BFL_BLUE, bold=True)
centered_run(doc, "PRODUCT REQUIREMENTS DOCUMENT (PRD)", 16, BFL_ORANGE, bold=True, underline=True)
doc.add_paragraph()

meta = [
    ("Version :",        "1.0"),
    ("Date :",           "May 2026"),
    ("Classification :", "Internal Use"),
    ("Prepared by :",    "National Lead – Agentic AI"),
    ("Reviewed by :",    "Sr. Head – AI  |  Head – Agentic AI"),
    ("Approved by :",    "Chief Operating Officer – IT"),
]
for label, val in meta:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para_spacing(p, before=40, after=40)
    r1 = p.add_run(label + "  ")
    r1.bold = True; r1.font.size = Pt(11); r1.font.color.rgb = BFL_BLUE
    r2 = p.add_run(val)
    r2.font.size = Pt(11)

doc.add_page_break()

# ──────────────────────────────────────────────────────
# SIGN-OFF SHEET
# ──────────────────────────────────────────────────────
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
para_spacing(p, before=60, after=60)
set_para_shading(p, '003387')
r = p.add_run("SIGN-OFF SHEET")
r.bold = True; r.font.size = Pt(13); r.font.color.rgb = WHITE

so_tbl = doc.add_table(rows=4, cols=3)
so_tbl.style = 'Table Grid'
for i, h in enumerate(["Role", "Name / Designation", "Signature & Date"]):
    cell = so_tbl.rows[0].cells[i]
    set_cell_bg(cell, '003387')
    set_cell_borders(cell, '003387')
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(h); r.bold = True; r.font.size = Pt(10); r.font.color.rgb = WHITE

so_data = [
    ("Approved by", "Chief Operating Officer – IT", ""),
    ("Recommended & Reviewed by", "Sr. Head – AI\nHead – Agentic AI", ""),
    ("Prepared by", "National Lead – Agentic AI", ""),
]
for ri, (role, name, sig) in enumerate(so_data):
    row = so_tbl.rows[ri + 1]
    fill = 'F2F2F2' if ri % 2 == 0 else 'FFFFFF'
    for ci, val in enumerate([role, name, sig]):
        cell = row.cells[ci]
        set_cell_bg(cell, fill)
        set_cell_borders(cell, 'D9D9D9')
        p = cell.paragraphs[0]
        r = p.add_run(val)
        r.font.size = Pt(10)
        if ci == 0:
            r.bold = True; r.font.color.rgb = BFL_BLUE

for i, w in enumerate([2.0, 3.0, 2.3]):
    for row in so_tbl.rows:
        row.cells[i].width = Inches(w)
doc.add_paragraph()
doc.add_page_break()

# ──────────────────────────────────────────────────────
# TABLE OF CONTENTS
# ──────────────────────────────────────────────────────
p = doc.add_paragraph()
para_spacing(p, before=60, after=100)
r = p.add_run("TABLE OF CONTENTS")
r.bold = True; r.font.size = Pt(14); r.font.color.rgb = BFL_BLUE

toc = [
    ("1",  "Purpose"),
    ("2",  "Objective"),
    ("3",  "Scope"),
    ("4",  "Core Principles"),
    ("5",  "Standard Definitions"),
    ("6",  "Knowledge Clean-up and Standardization"),
    ("7",  "Use Case Definition"),
    ("8",  "Process Decomposition"),
    ("9",  "System Actions (Agent Execution Layer)"),
    ("10", "UI Change Policy (Vibe-Coded UI)"),
    ("11", "End-to-End Process Flow Description"),
    ("12", "Agent Definitions and Requirements"),
    ("13", "BRD Acceptance Criteria"),
    ("14", "Approvals and Ownership"),
    ("15", "Guiding Enterprise Rule"),
    ("16", "Go / No-Go Readiness Checklist"),
    ("Appendix A", "Output File Structure Reference"),
    ("Appendix B", "18 CX Dimension Reference"),
    ("Appendix C", "Version History"),
]
for num, title in toc:
    p = doc.add_paragraph()
    para_spacing(p, before=28, after=28)
    r1 = p.add_run(f"{num}.   ")
    r1.bold = True; r1.font.size = Pt(10.5); r1.font.color.rgb = BFL_BLUE
    r2 = p.add_run(title)
    r2.font.size = Pt(10.5)

doc.add_page_break()

# ══════════════════════════════════════════════════════
# SECTION 1 – PURPOSE
# ══════════════════════════════════════════════════════
h1(doc, "1", "Purpose")
body(doc,
    "This Product Requirements Document (PRD) establishes a comprehensive, enterprise-grade "
    "specification for the CX Audit Agent — an AI-powered platform that simulates real Indian "
    "customer journeys on fintech and BFSI websites and produces structured, evidence-based "
    "audit reports revealing exactly what specific customer types experience, feel, and decide "
    "as they navigate a live financial product website.")
body(doc,
    "This document is authored in full compliance with the Bajaj Finance Limited Enterprise "
    "Agentic Requirement SOP V1.0 (February 2026) and is submitted to the PMO team for review "
    "and approval prior to development commencement.")
body(doc, "The intent of this document is to ensure that:")
for pt in [
    "Business teams can clearly articulate what the CX Audit Agent does, in terms that reflect real-world operational needs and do not require technical or AI expertise.",
    "The Agentic AI system can execute all defined processes safely, deterministically, and in a fully auditable manner, with every decision grounded in visible on-screen evidence.",
    "Risk, compliance, and platform teams can govern, control, and scale the deployment effectively, with clearly defined boundaries, escalation paths, and human oversight mechanisms.",
    "All stakeholders — Business, Operations, AI Engineering, and Platforms — have a single, unambiguous reference document that eliminates assumptions and informal knowledge.",
]:
    body(doc, pt, bullet=True)
body(doc,
    "This PRD prioritises process clarity and business intent first. Agent execution logic is "
    "derived from that clarity, not the reverse.")

# ══════════════════════════════════════════════════════
# SECTION 2 – OBJECTIVE
# ══════════════════════════════════════════════════════
h1(doc, "2", "Objective")
body(doc, "The CX Audit Agent is designed to achieve the following business objectives:")
objectives = [
    ("CX Insight at Scale",
     "Provide structured, persona-specific insights into how real Indian customers of varying "
     "financial literacy, demographics, and intent experience the bajajfinserv.in website across "
     "all product categories — without any manual effort per persona."),
    ("Evidence-Based Audit",
     "Replace subjective heuristic website reviews with objective, screenshot-grounded, "
     "step-by-step evidence of exactly what each persona sees, feels, and decides during their journey."),
    ("Multi-Dimensional Scoring",
     "Evaluate the website experience across 18 CX dimensions plus independent Design, Content, "
     "and Accessibility modules — producing actionable scores and prioritised recommendations."),
    ("Persona-Driven Simulation",
     "Deploy AI agents that genuinely embody specific user personas and derive behaviour in real "
     "time from the persona description, without hardcoded rules, pre-scripted flows, or lookup tables."),
    ("Consistent Enterprise Standard",
     "Establish a repeatable, configurable audit methodology that can be re-run on any target URL, "
     "with any set of personas, at any time, producing comparable structured output."),
    ("Stakeholder-Ready Reporting",
     "Generate per-persona Markdown reports, a cross-persona master report, structured JSON data "
     "files, and an interactive visual dashboard requiring no backend server."),
    ("Full Audit Readiness",
     "All agent decisions, observations, screenshots, and evaluations are logged in full in "
     "structured JSONL files for audit, governance, and quality review at any time."),
]
for title, text in objectives:
    bold_label(doc, f"{title}: ", text, bullet=True)

# ══════════════════════════════════════════════════════
# SECTION 3 – SCOPE
# ══════════════════════════════════════════════════════
h1(doc, "3", "Scope")
body(doc, "This PRD and all components of the CX Audit Agent initiative apply to:")

h2(doc, "3.1", "In Scope")
in_scope = [
    "Anonymous (logged-out) journey simulation across all product categories accessible from the target website start URL.",
    "Authenticated (logged-in) journey simulation for personas requiring a post-login product application flow.",
    "Mobile viewport (360 x 740 Galaxy S8+ emulation with touch events and mobile user agent) and desktop viewport configurations, configurable per run.",
    "All pages reachable by interacting with visible elements from the start URL — no fabricated deep-path navigation.",
    "Journey simulation through the complete application funnel up to and including the last form step before the final Submit button. The agent stops before submitting.",
    "Login wall encounter simulation: the agent makes a genuine persona-driven decision to dismiss, enter mobile number and request OTP, or abandon.",
    "CX evaluation across all 18 canonical dimensions (see Appendix B).",
    "Design evaluation covering visual hierarchy, colour contrast, touch targets, CTA prominence, and consistency.",
    "Content evaluation covering financial clarity, jargon appropriateness, disclosure completeness, and micro-copy quality.",
    "Accessibility evaluation covering text size, colour contrast, touch targets, form accessibility, and language simplicity.",
    "Per-persona Markdown audit reports and a master cross-persona aggregated report.",
    "JSON data files and a step-level JSONL log for every persona journey.",
    "Static HTML dashboard for results visualisation — no backend server required.",
    "All supporting modules: Persona Parser, Browser Controller, Decision Engine, Journey Memory, Controller / Orchestrator, Evaluation Suite, and Report Generator.",
    "All Agentic AI documentation supporting this use case, including this PRD, BRD v2.0, and associated SOPs.",
    "All participating teams: Business / CX, Operations, Risk, Compliance, AI Engineering, and Platforms.",
]
for item in in_scope:
    body(doc, item, bullet=True)

h2(doc, "3.2", "Out of Scope")
out_scope = [
    "Actual form submission, loan application execution, or any real financial transaction of any kind.",
    "Real OTP verification — this requires a physical device and is explicitly excluded.",
    "Testing authenticated flows that require real customer account credentials.",
    "Load testing, performance benchmarking, or stress testing of the target website.",
    "Multi-session simulation or returning-user state modelling across sessions.",
    "Automated remediation, code fixes, or UI changes to the target website based on findings.",
    "The mobile/ directory and mobile_main.py file (Appium native app testing — a separate initiative out of BRD scope).",
    "Any backend server component for the dashboard — the dashboard is fully static.",
    "Real-time monitoring or continuous scheduled crawling without human-initiated run commands.",
]
for item in out_scope:
    body(doc, item, bullet=True)

h2(doc, "3.3", "Participating Teams")
teams = [
    ("Process Owner / CX Team",       "Owns the business purpose, persona definitions, and acts on CX recommendations."),
    ("Agentic AI Domain Team",         "Owns agent design, BRD quality, use case approval, and agent boundary definitions."),
    ("AI Engineering Team",            "Implements the Decision Engine, Evaluation Suite, and all agent modules."),
    ("Platform / Browser Team",        "Manages Playwright infrastructure, viewport configuration, and screenshot pipeline."),
    ("AI Security & Compliance Team",  "Reviews data handling, screenshot storage, PII treatment, and audit log completeness."),
    ("PMO Team",                        "Reviews and accepts this PRD against SOP V1.0 criteria."),
    ("Development / Platform Team",    "Verifies technical feasibility, system integrations, and implementation approach."),
]
grid_table(doc, ["Team / Role", "Responsibility"], teams, [2.2, 5.1])

# ══════════════════════════════════════════════════════
# SECTION 4 – CORE PRINCIPLES
# ══════════════════════════════════════════════════════
h1(doc, "4", "Core Principles")
body(doc,
    "The following principles are non-negotiable and guide all design, documentation, and "
    "execution decisions for the CX Audit Agent. They are directly derived from the BFL "
    "Enterprise Agentic Requirement SOP V1.0.")

principles = [
    ("1. Process First, Technology Second",
     "All documentation in this PRD first explains what the business process is, independent "
     "of AI tools or APIs. The technology stack exists to serve the process, not to define it. "
     "The agent's behaviour is derived from process clarity, not from model capabilities."),
    ("2. Human Clarity Before Agent Execution",
     "Every step executed by the CX Audit Agent must be clearly understandable by a new "
     "business user with no AI background. If a step cannot be explained in plain business "
     "language from start to finish, it must not be automated."),
    ("3. Small, Clear Steps Over Broad Descriptions",
     "Every process step has a precise trigger, a single clear purpose, and a defined outcome. "
     "Vague steps such as 'evaluate the page' are broken into atomic actions with specific outputs."),
    ("4. Explicit Is Safer Than Implicit",
     "No assumptions, institutional knowledge, or unwritten rules are permitted in this document. "
     "Every constraint, exception, and escalation path is explicitly stated with specific conditions."),
    ("5. Governance Without Slowing Business",
     "This PRD enables strong governance and audit readiness while allowing the CX team to "
     "move quickly, with confidence that the agent operates within defined and auditable boundaries."),
    ("6. Screen-Only Reasoning",
     "Every action the agent takes must be justified by something physically visible in the "
     "current screenshot. Training-data knowledge about website layouts is explicitly prohibited "
     "as justification for any action."),
    ("7. No Fabricated Navigation",
     "The agent never navigates to a URL it has not reached through visible clicks. Any navigate "
     "action targeting a URL not in the current session history is rejected before execution and "
     "logged as a policy violation."),
    ("8. Persona Integrity Over Exploration",
     "The agent embodies the assigned persona at all times. It does not explore the website "
     "like a QA tester. It behaves exactly as that specific human being would behave, including "
     "giving up when the persona would genuinely give up."),
]
for title, text in principles:
    p = doc.add_paragraph()
    para_spacing(p, before=60, after=40)
    set_indent(p, 360)
    r1 = p.add_run(title + "\n")
    r1.bold = True; r1.font.size = Pt(10.5); r1.font.color.rgb = BFL_BLUE
    r2 = p.add_run(text)
    r2.font.size = Pt(10.5)

# ══════════════════════════════════════════════════════
# SECTION 5 – STANDARD DEFINITIONS
# ══════════════════════════════════════════════════════
h1(doc, "5", "Standard Definitions")
body(doc, "The following standard granularity levels apply to all process documentation in this PRD:")
gran_data = [
    ("Use Case", "Business objective",     "CX Audit of bajajfinserv.in"),
    ("Process",  "Start-to-end journey",   "Simulate a persona journey and produce an audit report"),
    ("Step",     "Meaningful stage",       "Execute Decision Engine at current page state"),
    ("Action",   "Atomic activity",        "Click 'Apply Now' button at coordinate (180, 420)"),
]
grid_table(doc, ["Level", "Meaning", "Example"], gran_data, [1.4, 2.0, 3.9])
body(doc, "Business stakeholders should always think in processes and steps, not in technical identifiers or API calls.")
doc.add_paragraph()

h2(doc, "5.1", "Glossary of Key Terms")
glossary = [
    ("Persona",            "A user-defined profile describing a specific customer type being simulated. May include name, age, occupation, financial literacy, intent, constraints, and test data."),
    ("Journey",            "The complete sequence of agent steps from landing on the start URL to reaching a terminal condition."),
    ("Terminal Condition", "The defined reason a journey ends: Goal Achieved, Login Wall – OTP Requested, Login Wall – Persona Abandoned, Valid Journey Abandoned, Loop Detected, Consecutive Failures, Max Steps Reached, or Timeout."),
    ("Decision Engine",    "The AI core (OpenAI GPT-4.1 with vision) that receives a screenshot and page state and returns a structured JSON action decision with full CX observation fields."),
    ("Step",               "One complete cycle: screenshot captured → Decision Engine called → action returned → action executed → result recorded."),
    ("CX Dimension",       "One of 18 named aspects of customer experience evaluated after the journey by the CX Evaluator."),
    ("Loop Detection",     "The mechanism that identifies when the agent repeats the same step fingerprint within a sliding window and injects a warning or terminates the journey."),
    ("Login Wall",         "A sign-in or OTP popup that blocks the user from proceeding without providing a mobile number or credentials."),
    ("Evaluation Suite",   "The four post-journey modules: CX Evaluator, Design Evaluator, Content Analyzer, Accessibility Auditor."),
    ("Friction Point",     "A specific moment or UI element in the journey that caused confusion, delay, or abandonment for the persona."),
    ("P1 / P2 / P3",       "Recommendation priority levels. P1 = Critical (journey-blocking). P2 = High (significantly degrades experience). P3 = Enhancement (improves without blocking)."),
]
grid_table(doc, ["Term", "Definition"], glossary, [1.8, 5.5])

# ══════════════════════════════════════════════════════
# SECTION 6 – KNOWLEDGE CLEAN-UP
# ══════════════════════════════════════════════════════
h1(doc, "6", "Knowledge Clean-up and Standardization")
body(doc,
    "All process descriptions, agent instructions, and evaluation criteria in this PRD are "
    "written to the following standardisation rules. These rules are mandatory for any "
    "supplementary documentation (SOPs, knowledge bases, process notes) submitted alongside "
    "this PRD or used to update agent behaviour.")

h2(doc, "6.1", "Prohibited Phrases")
body(doc, "The following phrases introduce ambiguity and are not permitted anywhere in this PRD or supporting documentation:")
for phrase in [
    '"As per policy" — must be replaced with the specific policy rule, section reference, or data condition.',
    '"If required" — must be replaced with the specific condition that determines when the step is required.',
    '"Handle appropriately" — must be replaced with the exact handling instruction.',
    '"As applicable" — must be replaced with the specific applicability condition.',
    '"Ensure compliance before proceeding" — must be replaced with the specific compliance check, e.g., "Verify that the page state popup indicator is False before calling the Decision Engine."',
    'Any statement that assumes prior knowledge of the system, the website, or the persona\'s behaviour.',
]:
    body(doc, phrase, bullet=True)

h2(doc, "6.2", "Rewrite Standard")
body(doc, "Every statement in this PRD must:")
for rule in [
    "Refer to a specific system, rule, data element, or visible UI element.",
    "Be interpretable in exactly one way — no two readers should derive different meanings.",
    "Name the actor performing the action: human, agent, or system.",
    "State the input required and the output or outcome produced.",
]:
    body(doc, rule, bullet=True)
info_box(doc,
    "Example of rewrite standard applied to this product:\n"
    "  NON-COMPLIANT: 'The agent should check the page before clicking.'\n"
    "  COMPLIANT: 'Before executing a CLICK action, the Decision Engine must verify that the "
    "target element label is present in the page state interactive elements list extracted by "
    "the Browser Controller from the current DOM. If the label is absent, the Decision Engine "
    "must use click_x / click_y coordinates from the screenshot instead.'")

h2(doc, "6.3", "PII and Data Handling Rules")
body(doc, "The following explicit rules govern all personally identifiable information (PII) in the system:")
for rule in [
    "Test mobile numbers used by personas are synthetic values provided in the persona definition file. They do not belong to real customers and must not be substituted with real numbers.",
    "Screenshots captured during a journey are stored locally in the reports/{persona-slug}/steps/ directory. They are not transmitted to external systems beyond the transient OpenAI Vision API call.",
    "The OpenAI API processes screenshots transiently within the API response lifecycle. No screenshot data persists on OpenAI servers beyond this scope.",
    "Journey logs (JSONL files in logs/) contain step-level observations and may include text extracted from pages. These files are classified as Internal Use and must not be shared externally.",
    "No real customer account credentials are ever used. Authenticated journeys use only configured test accounts.",
]:
    body(doc, rule, bullet=True)

# ══════════════════════════════════════════════════════
# SECTION 7 – USE CASE DEFINITION
# ══════════════════════════════════════════════════════
h1(doc, "7", "Use Case Definition")

h2(doc, "7.1", "Mandatory Use Case Statement")
info_box(doc,
    "This use case enables the Bajaj Finance Digital CX team to audit the customer experience "
    "quality of bajajfinserv.in by deploying AI agents that simulate real Indian customer "
    "personas navigating the live website, within data privacy, audit trail, and no-form-submission "
    "constraints, using OpenAI Vision API (GPT-4.1), Playwright Browser Automation (Chromium), "
    "and a static HTML reporting dashboard.")

h2(doc, "7.2", "Use Case Decomposition")
body(doc, "The top-level use case breaks into the following constituent processes:")
uc_data = [
    ("UC-001", "Persona Journey Simulation",   "Simulate a complete website journey for a defined persona from start URL to terminal condition."),
    ("UC-002", "CX Dimension Evaluation",      "Score the completed journey across 18 CX dimensions with evidence-grounded rationale."),
    ("UC-003", "Design Quality Evaluation",    "Assess visual design quality of pages visited from the perspective of the specific persona."),
    ("UC-004", "Content Quality Evaluation",   "Assess written content quality, financial clarity, jargon appropriateness, and disclosure completeness."),
    ("UC-005", "Accessibility Evaluation",     "Assess inclusivity and accessibility of the experience for the specific persona."),
    ("UC-006", "Audit Report Generation",      "Produce per-persona and master Markdown reports and structured JSON data files."),
    ("UC-007", "Dashboard Visualisation",      "Present all audit results in an interactive static HTML dashboard requiring no backend server."),
]
grid_table(doc, ["Use Case ID", "Name", "Description"], uc_data, [1.1, 2.0, 4.2])

h2(doc, "7.3", "Actor Definitions")
actor_data = [
    ("Decision Engine (Agentic AI)",   "Primary journey actor. Makes all navigation and interaction decisions autonomously within defined constraints."),
    ("Browser Controller (System)",    "Executes actions decided by the Decision Engine against the live Chromium browser."),
    ("CX Evaluator (Agent)",           "Post-journey agent that evaluates the completed journey across 18 CX dimensions."),
    ("Design Evaluator (Agent)",       "Post-journey agent that assesses visual design quality."),
    ("Content Analyzer (Agent)",       "Post-journey agent that assesses content quality and financial clarity."),
    ("Accessibility Auditor (Agent)",  "Post-journey agent that assesses accessibility and inclusive design."),
    ("Report Generator (System)",      "Produces all output files from evaluation results and journey logs."),
    ("Human Operator",                 "Configures and launches the audit, reviews output reports, and acts on recommendations."),
    ("PMO / Audit Reviewer",           "Reviews this PRD, agent boundaries, and audit output for governance and SOP compliance."),
]
grid_table(doc, ["Actor", "Role in System"], actor_data, [2.3, 5.0])

# ══════════════════════════════════════════════════════
# SECTION 8 – PROCESS DECOMPOSITION
# ══════════════════════════════════════════════════════
h1(doc, "8", "Process Decomposition")
body(doc,
    "The complete CX Audit Agent workflow is decomposed below into processes, steps, and actions "
    "following the SOP V1.0 granularity standard. Each step states: what happens, who performs "
    "it, what information is used, what decision is made, and what happens next. No step uses a "
    "prohibited phrase from Section 6.1.")

# ── Process 1
h2(doc, "8.1", "Process: Audit Session Initialisation")
bold_label(doc, "Trigger: ", "Human operator executes the run command with configuration parameters.")
bold_label(doc, "Actor: ", "Human Operator and System")
bold_label(doc, "Output: ", "All personas loaded, Chromium browser running, output directories confirmed, browser at target URL.")
doc.add_paragraph()

p1_steps = [
    ("Step 1.1", "Read and Validate Configuration",
     "System (config.py + .env loader)",
     "Configuration file values: TARGET_URL, OPENAI_API_KEY, OPENAI_MODEL, viewport dimensions, mobile emulation flag, headed/headless flag, maximum steps, wall-clock timeout, output directory path.",
     "Is TARGET_URL a reachable URL? Is OPENAI_API_KEY present and non-empty? Are all required parameters present?",
     "All valid → proceed to Step 1.2. Any value missing or invalid → terminate with configuration error written to console log. Do not proceed."),
    ("Step 1.2", "Parse Persona Definition File",
     "Persona Parser (System)",
     "Persona file at the configured path (default: personas/bajaj_personas.md).",
     "Does each persona block contain sufficient information for a meaningful journey? Are test data fields present for personas that require form filling?",
     "All personas parsed → Persona objects list available; proceed to Step 1.3. Parse error on one persona → log the error for that persona, skip it, continue with remaining personas. Zero valid personas → terminate."),
    ("Step 1.3", "Initialise Output Directories",
     "System (Report Generator)",
     "Output directory path and log directory path from configuration.",
     "Do the required directories (reports/, logs/, dashboard/) exist? If not, can they be created?",
     "All directories confirmed present or created → proceed to Step 1.4. Directory creation failure → terminate with file system error log."),
    ("Step 1.4", "Launch Chromium Browser",
     "Browser Controller (System)",
     "Viewport width, viewport height, mobile emulation flag, headed/headless mode flag, device pixel ratio (if mobile emulation active).",
     "Is the Chromium process launched and the Playwright page object responsive within 10 seconds?",
     "Browser active → proceed to Step 1.5. Launch failure → retry once after 5 seconds. Second failure → terminate with browser launch error log."),
    ("Step 1.5", "Navigate to Target URL",
     "Browser Controller (System)",
     "TARGET_URL from configuration.",
     "Did the page load successfully (HTTP 200 or equivalent success response) within the configured timeout?",
     "Page loaded → record start URL in session history; proceed to Persona Journey Execution (Step 2.1 for first persona). Load failure (timeout, DNS error, HTTP error) → log failure; terminate session."),
]
for sid, sname, actor, info, decision, outcome in p1_steps:
    h3(doc, f"{sid}: {sname}")
    bold_label(doc, "What happens: ", f"The system executes {sname.lower()} as part of audit session initialisation.", indent_twips=360)
    bold_label(doc, "Who performs it: ", actor, indent_twips=360)
    bold_label(doc, "Information used: ", info, indent_twips=360)
    bold_label(doc, "Decision logic: ", decision, indent_twips=360)
    bold_label(doc, "Next outcome: ", outcome, indent_twips=360)
    doc.add_paragraph()

# ── Process 2
h2(doc, "8.2", "Process: Persona Journey Execution")
bold_label(doc, "Trigger: ", "Audit session initialised; browser at target URL; Persona object available.")
bold_label(doc, "Actor: ", "Decision Engine (Agentic AI), Browser Controller (System), Controller / Orchestrator")
bold_label(doc, "Output: ", "Complete step log with screenshots, CX observation fields for every step, and a single terminal condition.")
doc.add_paragraph()

p2_steps = [
    ("Step 2.1", "Capture Screenshot and Extract Page State",
     "Browser Controller (System)",
     "Current browser viewport state.",
     "Was the screenshot captured successfully? Was the DOM parsed successfully to extract page state?",
     "Screenshot captured and compressed to JPEG 72% quality, max 1280px → page state extracted → proceed to Step 2.2. Screenshot capture failure → increment consecutive failure counter; if counter equals CONSECUTIVE_FAIL_LIMIT, terminate with reason 'consecutive_failures'."),
    ("Step 2.2", "Check All Terminal Conditions",
     "Controller / Orchestrator",
     "Current step number vs. configured maximum, wall-clock elapsed time vs. configured timeout, loop fingerprint registry (last N steps), consecutive failure counter.",
     "Has any of the following terminal conditions been met: (a) step count equals MAX_STEPS, (b) elapsed time exceeds WALL_CLOCK_TIMEOUT, (c) loop detected (same fingerprint recurs within sliding window after warning injected), (d) consecutive failure counter equals CONSECUTIVE_FAIL_LIMIT?",
     "Terminal condition met → record terminal reason; end journey; proceed to Post-Journey Evaluation (Section 8.4). No terminal condition → proceed to Step 2.3."),
    ("Step 2.3", "Call Decision Engine",
     "Decision Engine (Agentic AI – OpenAI GPT-4.1 Vision)",
     "Current screenshot (JPEG, upscaled to visual clarity width if viewport < 400px), extracted page state text, full persona description, complete journey history (all previous steps, actions, CX notes, login wall encounters), viewport dimensions in original pixel space, current step number, configured maximum steps, session type (logged_out or logged_in), test data for form filling, loop warning flag (True if loop pattern detected in Step 2.2).",
     "Based on the persona description and what is physically visible in the screenshot right now: what is the single best next action this persona would take? What are the full CX observations for this step?",
     "Structured JSON response received → proceed to Step 2.4. OpenAI API call failure → increment consecutive failure counter; retry once; if second failure → increment counter again; if counter equals CONSECUTIVE_FAIL_LIMIT → terminate with reason 'consecutive_failures'."),
    ("Step 2.4", "Validate Returned Action",
     "Controller / Orchestrator",
     "Action JSON from Decision Engine, browser session URL history (all URLs loaded in current session), final-submit button pattern library (configured list of labels that indicate a final submission action).",
     "Is the action a NAVIGATE to a URL not present in the current session history (fabricated URL)? Does the action target a button label matching the final-submit pattern library? Is the JSON structurally valid?",
     "Fabricated URL detected → reject action; write violation to JSONL issue log; ask Decision Engine for an alternate action (return to Step 2.3, injection of rejection notice). Final submit detected → intercept without executing; record terminal condition 'goal_achieved'; end journey. Structurally invalid JSON → increment failure counter; treat as failure; if limit reached → terminate. Valid action → proceed to Step 2.5."),
    ("Step 2.5", "Execute Action via Browser Controller",
     "Browser Controller (System)",
     "Validated action JSON from Decision Engine: action type, selector, text, url, scroll_direction, click_x, click_y, selector_hints, fallback.",
     "Did the Browser Controller successfully execute the primary action within the configured element timeout? If primary action failed, did the fallback action succeed?",
     "Primary action succeeded → record success in step log; proceed to Step 2.6. Primary action failed → attempt fallback if provided; if fallback also fails → increment consecutive failure counter; record failure details in JSONL issue log; if counter equals CONSECUTIVE_FAIL_LIMIT → terminate."),
    ("Step 2.6", "Record Step in Journey Memory",
     "Journey Memory (System)",
     "Action taken, action result (success/failure), screenshot file path, all CX observation fields from Decision Engine output: observation, reasoning, emotion, cx_note, cognitive_load, trust_signals, unanswered_questions, guiding_factors, visible_content, state_of_mind, login_wall_decision, login_wall_reasoning.",
     "Is the step fingerprint (action type + target element label + current URL) a repeat within the sliding window of the last N steps?",
     "Repeat fingerprint detected → set loop warning flag True for next Decision Engine call (Step 2.3). If loop warning was already True at this step and fingerprint still repeats → record terminal condition 'loop_detected'; end journey. No repeat → clear loop warning flag if previously set; consecutive failure counter reset if this step succeeded; return to Step 2.1 for next step."),
]
for sid, sname, actor, info, decision, outcome in p2_steps:
    h3(doc, f"{sid}: {sname}")
    bold_label(doc, "What happens: ", f"The Controller executes {sname.lower()} in the step loop.", indent_twips=360)
    bold_label(doc, "Who performs it: ", actor, indent_twips=360)
    bold_label(doc, "Information used: ", info, indent_twips=360)
    bold_label(doc, "Decision logic: ", decision, indent_twips=360)
    bold_label(doc, "Next outcome: ", outcome, indent_twips=360)
    doc.add_paragraph()

# ── Process 3
h2(doc, "8.3", "Process: Login Wall Handling")
bold_label(doc, "Trigger: ", "Decision Engine identifies a login wall, OTP popup, or sign-in modal on the current screenshot.")
bold_label(doc, "Actor: ", "Decision Engine (Agentic AI)")
bold_label(doc, "Output: ", "Login wall decision recorded (dismissed / entered_mobile / ignored) with full reasoning; journey continues or terminates.")
doc.add_paragraph()

lw_steps = [
    ("Step 3.1", "Identify Login Wall",
     "Decision Engine (Agentic AI)",
     "Current screenshot; page state popup indicator; login_wall_encounter_log from journey history.",
     "Is a login wall, OTP entry popup, or sign-in modal physically visible in the current screenshot — not merely a header Sign In link?",
     "Login wall visually confirmed → proceed to Step 3.2. No login wall visible → return to standard step cycle (Step 2.3)."),
    ("Step 3.2", "Assess Persona Disposition Toward the Login Wall",
     "Decision Engine (Agentic AI)",
     "Full persona description (financial literacy, risk tolerance, data-sharing comfort, technology confidence), total number of login walls encountered so far in this journey, the persona's current emotional state and motivation level from journey history, the value the persona has received from public-facing content so far.",
     "Given everything this persona has seen and experienced in this journey up to this exact moment: would they (A) dismiss the wall and continue exploring, (B) enter their test mobile number to request an OTP, or (C) abandon the journey entirely?",
     "Decision A (Dismiss) → execute DISMISS POPUP action; record login_wall_decision = 'dismissed' with full reasoning. Decision B (Enter Mobile) → execute TYPE action entering configured test mobile number; then execute CLICK action on the OTP / submit button using its exact visible label; record login_wall_decision = 'entered_mobile'; terminal condition 'login_wall_otp_requested' triggered. Decision C (Abandon) → execute DONE action; record login_wall_decision = 'ignored'; terminal condition 'login_wall_persona_abandoned' triggered."),
    ("Step 3.3", "Record Login Wall Encounter",
     "Journey Memory (System)",
     "Login wall decision, full reasoning, screenshot reference, step number, running count of login wall encounters.",
     "Is this the third login wall encounter where the persona chose to dismiss (Option A) and the journey is still active?",
     "Third dismiss and journey blocked → terminal condition 'login_wall_persona_abandoned'. Fewer than three dismissals or journey still progressing → continue journey normally."),
]
for sid, sname, actor, info, decision, outcome in lw_steps:
    h3(doc, f"{sid}: {sname}")
    bold_label(doc, "Who performs it: ", actor, indent_twips=360)
    bold_label(doc, "Information used: ", info, indent_twips=360)
    bold_label(doc, "Decision logic: ", decision, indent_twips=360)
    bold_label(doc, "Next outcome: ", outcome, indent_twips=360)
    doc.add_paragraph()

# ── Process 4
h2(doc, "8.4", "Process: Post-Journey Evaluation")
bold_label(doc, "Trigger: ", "Journey has reached a terminal condition. Journey log is complete.")
bold_label(doc, "Actor: ", "Four independent Evaluation Agents running post-journey.")
bold_label(doc, "Output: ", "Structured evaluation results with dimension scores, friction points, positive moments, and prioritised recommendations from all four modules.")
doc.add_paragraph()

eval_steps = [
    ("Step 4.1", "CX Evaluator — Score 18 Dimensions",
     "CX Evaluator (Agent – GPT-4.1 Vision)",
     "Complete journey log (all steps, all CX observation fields), sampled screenshots (JPEG 72%, max 1280px), full persona description, the 18 canonical CX dimension names from BRD Section 8.1.",
     "For each of the 18 CX dimensions: what score from 0 to 10 is supported by the specific journey evidence? Which steps provide the primary evidence for each score? What friction points occurred per dimension? What recommendations emerge?",
     "18 dimension scores with evidence-grounded rationale + overall CX score + journey verdict (one paragraph) + TL;DR (2–3 sentences) + key takeaways (3–5 specific insights) + friction points (each with severity, page/location, description, customer impact) + positive moments + recommendations (P1/P2/P3 each with area, action, expected impact) + emotional arc stage-by-stage + persona emotional narrative → proceed to Step 4.2."),
    ("Step 4.2", "Design Evaluator — Assess Visual Design Quality",
     "Design Evaluator (Agent – GPT-4.1 Vision)",
     "Sampled journey screenshots, page state data, persona description.",
     "For each design dimension (visual hierarchy, colour contrast, touch target adequacy, CTA prominence, use of imagery, whitespace/visual density, design consistency across pages, form field design): what specific evidence from the visited pages supports the assessment?",
     "Design dimension scores, specific observations for each page visited, friction points related to design, and design-specific recommendations (P1/P2/P3) → proceed to Step 4.3."),
    ("Step 4.3", "Content Analyzer — Assess Written Content Quality",
     "Content Analyzer (Agent – GPT-4.1 Vision)",
     "Sampled screenshots, extracted page text content, persona description and financial literacy level.",
     "Is financial information (rates, fees, charges) clear and prominently displayed before any login gate? Is the language appropriate for this persona's financial literacy level? Are regulatory disclosures present? Does the pre-gate content give enough information for this persona to make an informed decision?",
     "Content findings and content-specific friction points, specific content gaps, jargon assessment per page, micro-copy quality assessment, and content recommendations (P1/P2/P3) → proceed to Step 4.4."),
    ("Step 4.4", "Accessibility Auditor — Assess Inclusivity",
     "Accessibility Auditor (Agent – GPT-4.1 Vision)",
     "Sampled screenshots, page state text, persona description and any accessibility-relevant attributes.",
     "Are text sizes adequate for this persona? Is colour contrast sufficient across visited pages? Are touch targets sized appropriately for mobile? Is the language simple enough for this persona's literacy level? Are form fields accessible?",
     "Accessibility friction points with severity levels, WCAG-aligned recommendations (P1/P2/P3), and specific findings per page visited → all four evaluation modules complete; proceed to Report Generation (Step 5.1)."),
]
for sid, sname, actor, info, decision, outcome in eval_steps:
    h3(doc, f"{sid}: {sname}")
    bold_label(doc, "Who performs it: ", actor, indent_twips=360)
    bold_label(doc, "Information used: ", info, indent_twips=360)
    bold_label(doc, "Decision logic: ", decision, indent_twips=360)
    bold_label(doc, "Next outcome: ", outcome, indent_twips=360)
    doc.add_paragraph()

# ── Process 5
h2(doc, "8.5", "Process: Report Generation and Dashboard Update")
bold_label(doc, "Trigger: ", "All four evaluation modules have completed for a persona.")
bold_label(doc, "Actor: ", "Report Generator (System)")
bold_label(doc, "Output: ", "Per-persona report.md, journey_log.json, step screenshots, updated master_report.md, updated session_index.json.")
doc.add_paragraph()

rpt_steps = [
    ("Step 5.1", "Generate Per-Persona Markdown Report",
     "Report Generator (System)",
     "CX evaluation results, Design evaluation results, Content analysis results, Accessibility audit results, complete journey log.",
     "Are all 13 required report sections present and populated with evidence-grounded content? (Sections: TL;DR and Key Takeaways, Score Snapshot, Journey Outcome, Critical Issues, All Recommendations, What Worked Well, Full CX Dimension Analysis, Design Audit Findings, Content Analysis Findings, Accessibility Audit Findings, Friction Point Reference List, Emotional Journey Arc, Step-by-Step Journey Log Appendix)",
     "reports/{persona-slug}/report.md written with all 13 sections → proceed to Step 5.2. File write failure → log error; continue to next step."),
    ("Step 5.2", "Write Journey Log JSON",
     "Report Generator (System)",
     "Complete structured journey data including all step fields, all CX observation fields, evaluation results, scores, friction points, and recommendations.",
     "Does the JSON file include all fields required by the dashboard? Are screenshot file paths relative (not absolute) to ensure portability?",
     "reports/{persona-slug}/journey_log.json written with all required fields → proceed to Step 5.3."),
    ("Step 5.3", "Update Master Report",
     "Report Generator (System)",
     "All per-persona reports from the current session.",
     "Do the cross-persona findings include: score comparison table, shared friction points appearing in 2 or more personas, divergent findings by persona type, and site-level recommendations?",
     "reports/master_report.md updated with cross-persona aggregated content → proceed to Step 5.4."),
    ("Step 5.4", "Update Dashboard Session Index",
     "Report Generator (System)",
     "All persona outcomes, overall scores, terminal conditions, and report file paths from the current session.",
     "Does session_index.json accurately reflect all personas run in this session, including their outcome labels and scores?",
     "dashboard/session_index.json updated. Dashboard is ready to open in any browser by opening dashboard/index.html directly. No web server required."),
]
for sid, sname, actor, info, decision, outcome in rpt_steps:
    h3(doc, f"{sid}: {sname}")
    bold_label(doc, "Who performs it: ", actor, indent_twips=360)
    bold_label(doc, "Information used: ", info, indent_twips=360)
    bold_label(doc, "Decision logic: ", decision, indent_twips=360)
    bold_label(doc, "Next outcome: ", outcome, indent_twips=360)
    doc.add_paragraph()

# ══════════════════════════════════════════════════════
# SECTION 9 – SYSTEM ACTIONS
# ══════════════════════════════════════════════════════
h1(doc, "9", "System Actions (Agent Execution Layer)")
body(doc,
    "While business teams describe steps (Section 8), agents perform system actions. System "
    "actions are the atomic browser instructions executed by the Browser Controller on behalf "
    "of the Decision Engine. Each action is described in plain business language with its "
    "trigger condition, how it is executed, and the validation outcome.")

h2(doc, "9.1", "Action Catalogue")
actions = [
    ("CLICK",
     "The Decision Engine identifies a specific visible UI element the persona would interact with.",
     "Browser Controller locates the element using the selector (visible text label, accessible name, CSS selector, or XPath) or pixel coordinate (click_x, click_y) in the original viewport pixel space. Executes a click event on the element.",
     "Element clicked and page response confirmed within timeout. If element not found or click fails → increment consecutive failure counter; record in issue log."),
    ("TYPE",
     "A form field, search bar, range slider, date picker, or select dropdown is present and the persona needs to enter a value.",
     "Browser Controller focuses the input field identified by selector and types the text value specified in the action JSON 'text' field. For sliders and select elements, the value is entered as the desired value string.",
     "Value entered and field contains expected text. If field not found → increment failure counter."),
    ("SCROLL DOWN",
     "The persona's target content is not visible in the current viewport and may be below the current scroll position.",
     "Browser Controller executes a vertical downward scroll by one viewport height on the main page document.",
     "Page scrolled downward. New screenshot captured at next step. If already at page bottom → log 'scroll_boundary_reached'; this prevents infinite downward scrolling loops."),
    ("SCROLL UP",
     "The persona needs to return to content above the current scroll position.",
     "Browser Controller executes a vertical upward scroll by one viewport height.",
     "Page scrolled upward. New screenshot at next step."),
    ("SCROLL RIGHT",
     "A horizontal carousel or product strip is visible and the persona needs to see more items to the right.",
     "Browser Controller identifies the scrollable carousel container element and scrolls it horizontally to the right by the configured carousel scroll distance.",
     "Carousel advanced to show next set of items. New screenshot at next step."),
    ("SCROLL LEFT",
     "A horizontal carousel or product strip is visible and the persona needs to go back to items on the left.",
     "Browser Controller identifies the scrollable carousel container and scrolls it horizontally to the left.",
     "Carousel position moved left. New screenshot at next step."),
    ("NAVIGATE",
     "Recovery only: the persona needs to return to the start URL or a specific page previously loaded in the current browser session.",
     "Browser Controller checks whether the target URL is present in the current session URL history. If present: navigates to the URL. If the URL is NOT in session history: the action is rejected before execution. The rejection is logged as a policy violation in the JSONL issue log and the Decision Engine is asked to choose a different action.",
     "Navigation executed to a confirmed session URL. New screenshot at next step. Fabricated URL → rejected, violation logged, alternate action requested."),
    ("BACK",
     "The persona wants to return to the previous page in browser history.",
     "Browser Controller triggers the browser back action.",
     "Previous page loaded. New screenshot at next step."),
    ("WAIT",
     "A page loading state is visible, or the persona is pausing to read content before proceeding.",
     "Browser Controller pauses execution for the number of milliseconds specified in the action JSON 'text' field.",
     "Wait duration elapsed. Continue to next step without capturing a new screenshot mid-wait."),
    ("PLAY VIDEO",
     "A video element is visible on the current page and is in a paused state; the persona would watch it.",
     "Browser Controller identifies the video HTML element and triggers the play() event on it.",
     "Video element playing confirmed. New screenshot at next step."),
    ("DISMISS POPUP",
     "A genuine obstruction — cookie banner, promotional popup, chat bubble, or modal that contains no navigation links or product links — is blocking the main page content.",
     "Browser Controller identifies the close or dismiss element of the popup (typically an X button or 'Not now' label) and clicks it. CRITICAL: The agent must NOT use this action on a navigation drawer or side menu that contains product links. A navigation panel is not an obstruction — it is the navigation.",
     "Popup closed and main content confirmed visible. If same popup recurs three times at the same URL → terminal condition 'popup_loop' triggered."),
    ("SEARCH",
     "The persona cannot find their target product or information through visible page content after exploring available options, and the site's internal search bar is visible.",
     "Browser Controller locates the site search input element and types 1 to 3 keywords that reflect the persona's intent. Submits the search query.",
     "Search results page loaded. New screenshot at next step."),
    ("DONE",
     "The persona has genuinely reached a natural journey endpoint — either the persona's goal has been fulfilled or the persona has made a real decision to leave the website.",
     "Controller records the terminal condition with the specific reason. No further browser action is executed. Journey Memory is finalised.",
     "Journey ends. Terminal reason recorded. Post-journey Evaluation Suite triggered (Section 8.4)."),
]
for action_name, trigger, execution, validation in actions:
    h3(doc, action_name)
    bold_label(doc, "Trigger: ", trigger, indent_twips=360)
    bold_label(doc, "Execution: ", execution, indent_twips=360)
    bold_label(doc, "Validation: ", validation, indent_twips=360)
    doc.add_paragraph()

h2(doc, "9.2", "Action Owner Matrix")
body(doc, "S = System automated  |  A = Agent decides  |  M = Manual human approval required")
owner_data = [
    ("CLICK",           "A – Decision Engine decides target and coordinates", "S – Browser Controller executes", "No"),
    ("TYPE",            "A – Decision Engine decides field and value",         "S – Browser Controller executes", "No"),
    ("SCROLL DOWN/UP",  "A – Decision Engine decides",                         "S – Browser Controller executes", "No"),
    ("SCROLL LEFT/RIGHT","A – Decision Engine decides",                        "S – Browser Controller executes", "No"),
    ("NAVIGATE",        "A – Decision Engine requests target URL",             "S – Validated against session history before execution", "No — but URL validated; fabricated URLs blocked"),
    ("BACK",            "A – Decision Engine decides",                         "S – Browser Controller executes", "No"),
    ("WAIT",            "A – Decision Engine decides duration",                "S – Browser Controller executes", "No"),
    ("PLAY VIDEO",      "A – Decision Engine decides",                         "S – Browser Controller executes", "No"),
    ("DISMISS POPUP",   "A – Decision Engine decides",                         "S – Browser Controller executes; navigation panels NOT dismissable", "No"),
    ("SEARCH",          "A – Decision Engine decides query keywords",          "S – Browser Controller executes", "No"),
    ("DONE",            "A – Decision Engine decides",                         "S – Controller terminates journey and triggers evaluation", "No"),
]
grid_table(doc, ["Action", "Decision Actor", "Execution Actor", "Human Approval?"],
           owner_data, [1.3, 2.1, 2.3, 1.5])

# ══════════════════════════════════════════════════════
# SECTION 10 – UI CHANGE POLICY
# ══════════════════════════════════════════════════════
h1(doc, "10", "UI Change Policy (Vibe-Coded UI)")
body(doc,
    "The CX Audit Agent includes a static HTML dashboard (dashboard/index.html). The following "
    "policy governs all changes to the dashboard UI and any configuration screens introduced "
    "by this product.")

h2(doc, "10.1", "When UI Documentation Is Required")
body(doc, "UI documentation is mandatory whenever:")
for cond in [
    "A new screen or view is introduced in the dashboard.",
    "An existing dashboard view is modified in a way that changes how data is presented or how users interact with it.",
    "A filter, sort, or export behaviour on any dashboard view changes.",
    "A configuration input (run parameters, persona file path, output directory path) is exposed through any UI element.",
    "A UI element is added to or removed from any of the seven defined dashboard views.",
]:
    body(doc, cond, bullet=True)

h2(doc, "10.2", "How to Describe UI Changes")
body(doc,
    "UI requirements must describe behaviour, guidance text, data binding, validation rules, "
    "and accessibility requirements. Visual design (specific colours, font sizes, exact spacing) "
    "is out of scope for this PRD.")
for rule in [
    "State what data is displayed: name the exact JSON field from session_index.json or journey_log.json that populates each UI element.",
    "State what interactions are available: which filters exist, what they filter on, what the default state is.",
    "State validation: what the UI must display when a JSON file is missing, malformed, or contains zero steps.",
    "State accessibility: all interactive elements must have accessible names; all score values must be readable without colour alone.",
]:
    body(doc, rule, bullet=True)

h2(doc, "10.3", "Dashboard Screen Inventory")
ui_screens = [
    ("Run Audit Tab",
     "Session metadata (date, time, target URL, persona count) and a summary table of all personas with columns: Name, Intent, Outcome Label, Overall Score. Quick filters: by outcome type (dropdown), by score range (slider 0–10).",
     "session_index.json → fields: session_date, target_url, personas[].name, personas[].intent, personas[].terminal_condition, personas[].overall_score",
     "If session_index.json is absent: display 'No audit data found. Run the agent to generate results.' If a persona row has no journey_log.json: display persona row with status field showing 'Data unavailable'."),
    ("Persona Journey View",
     "Persona card (name, intent, outcome, overall score, emotional arc summary) and a step-by-step timeline. Each step shows: screenshot thumbnail (clickable to full size), action type, target element, emotion label, full CX note text, cognitive load level, trust signals observed, unanswered questions, guiding factors, state of mind, and login wall decision with reasoning where applicable. Steps are expandable / collapsible.",
     "journey_log.json → fields: persona.name, persona.intent, terminal_condition, overall_score, steps[].screenshot_path, steps[].action, steps[].emotion, steps[].cx_note, steps[].cognitive_load, steps[].trust_signals, steps[].state_of_mind, steps[].login_wall_decision",
     "If screenshot file path is broken: display placeholder image with label 'Screenshot unavailable'. If CX note is empty string: display 'No CX note recorded for this step.'"),
    ("Dimension Scores View",
     "All 18 CX dimension scores displayed as colour-coded visual score bars (0–10 scale). Colour coding: 0–3 = Red, 4–6 = Amber, 7–10 = Green. Each dimension is expandable to show rationale text and step-level evidence references. Design score, Content score, and Accessibility score displayed alongside CX dimensions.",
     "journey_log.json → cx_evaluation.dimensions[].name, .score, .rationale, .step_evidence[]",
     "Dimension name must display exactly as stored in JSON — no reformatting. If score field is null: display 'Score unavailable'."),
    ("Friction Points and Issues View",
     "All friction points and issues from all four evaluation modules listed in a filterable table. Columns: Severity, Source Module, Page / URL, Description, Customer Impact. Filterable by severity (High/Medium/Low) and by source module (CX/Design/Content/Accessibility).",
     "journey_log.json → friction_points[].severity, .source, .page, .description, .customer_impact",
     "Empty state: 'No friction points recorded for this persona.' Default sort: High severity first."),
    ("Recommendations View",
     "All recommendations from all four evaluation modules listed in a filterable table. Columns: Priority, Source Module, Area Affected, Specific Action, Expected Impact. Filterable by priority (P1/P2/P3) and source module. P1 items highlighted in red. Default sort: P1 first, then P2, then P3.",
     "journey_log.json → recommendations[].priority, .source, .area, .action, .expected_impact",
     "Empty state: 'No recommendations recorded for this persona.'"),
    ("Issue Logs Tab",
     "Per-step JSONL entries in full. Sections: Action failures and fallback attempts, Loop detection events with step numbers that triggered each warning, Terminal condition trigger with full context, Login wall encounter log with all decisions, Token usage per step, Total token usage for journey, Estimated API cost for the run.",
     "logs/{persona-slug}.jsonl — all step entries in raw JSONL format",
     "If JSONL file is absent: display 'No debug log available for this persona. Ensure logs/ directory is writable.'"),
    ("Emotional Journey View",
     "A visual representation of emotion states across all steps (emotion label on Y-axis, step number on X-axis). Stage-by-stage breakdown table: Stage Name, Emotion, Specific Trigger at that stage. Full persona emotional narrative text displayed below the chart.",
     "journey_log.json → emotional_arc.stages[].stage, .emotion, .trigger; emotional_arc.narrative",
     "If emotional arc data is absent: display 'Emotional arc data not available for this persona.'"),
]
for screen_name, description, data_source, validation in ui_screens:
    h3(doc, screen_name)
    bold_label(doc, "What is displayed: ", description, indent_twips=360)
    bold_label(doc, "Data source field(s): ", data_source, indent_twips=360)
    bold_label(doc, "Validation / error state: ", validation, indent_twips=360)
    doc.add_paragraph()

# ══════════════════════════════════════════════════════
# SECTION 11 – END-TO-END PROCESS FLOW
# ══════════════════════════════════════════════════════
h1(doc, "11", "End-to-End Process Flow Description")
body(doc,
    "Per SOP V1.0 Section 11, all critical Agentic AI use cases must include a visual swimlane-style "
    "process flow. The narrative below specifies the complete flow and serves as the specification "
    "for the formal Miro/Visio swimlane diagram to be produced by the Agentic AI Domain Team and "
    "attached as Appendix A prior to final PMO submission.")

h2(doc, "11.1", "Swimlane Owner Definitions")
swimlane_data = [
    ("Lane 1: Human Operator",              "Configures and initiates the audit run. Reviews output reports. Acts on recommendations."),
    ("Lane 2: Orchestrator (System)",        "Manages the step loop. Executes terminal condition checks. Manages loop detection and failure counting."),
    ("Lane 3: Decision Engine (Agent)",      "Receives screenshot + context at each step. Returns structured JSON: action + CX observations."),
    ("Lane 4: Browser Controller (System)", "Executes all browser actions. Captures screenshots. Extracts page state from DOM."),
    ("Lane 5: Evaluation Suite (Agents)",   "Four post-journey agents: CX Evaluator, Design Evaluator, Content Analyzer, Accessibility Auditor."),
    ("Lane 6: Report Generator (System)",   "Produces all output files after all evaluation modules complete."),
]
grid_table(doc, ["Swimlane", "Owner and Responsibility"], swimlane_data, [2.5, 4.8])

h2(doc, "11.2", "Flow Narrative (Step-by-Step)")
flow = [
    ("[Lane 1 – Human Operator]",
     "Operator runs: python main.py with optional flags. Configuration loaded and validated. Persona file parsed. Output directories confirmed. Browser process launched at configured viewport."),
    ("[Lane 1 → Lane 4 – Handoff]",
     "Browser Controller navigates to TARGET_URL. Page loads. First screenshot captured and compressed. Page state extracted. Step counter initialised to 1."),
    ("[Lane 2 – Orchestrator: Loop Start]",
     "FOR EACH STEP in the journey:\n"
     "  1. Orchestrator checks terminal conditions (max steps, timeout, loop detection, failure limit).\n"
     "  2. If terminal condition met → EXIT LOOP → proceed to Evaluation Suite.\n"
     "  3. If no terminal condition → pass screenshot + context to Decision Engine."),
    ("[Lane 3 – Decision Engine]",
     "Receives screenshot (vision input), page state text, persona description, journey history, loop warning flag.\n"
     "Returns structured JSON with: action, selector, click_x, click_y, scroll_direction, text, confidence, "
     "observation, reasoning, emotion, cx_note, cognitive_load, trust_signals, unanswered_questions, "
     "guiding_factors, visible_content, state_of_mind, login_wall_decision, login_wall_reasoning."),
    ("[Lane 2 – Orchestrator: Validate]",
     "Validates action JSON:\n"
     "  - Fabricated URL? → Reject; log violation; re-call Decision Engine with rejection notice.\n"
     "  - Final submit button? → Intercept; record 'goal_achieved'; EXIT LOOP.\n"
     "  - Structurally invalid? → Increment failure counter.\n"
     "  - Valid → pass to Browser Controller."),
    ("[Lane 4 – Browser Controller]",
     "Executes the validated action. If primary action fails and fallback is provided, attempts fallback.\n"
     "Captures new screenshot after action. Extracts updated page state."),
    ("[Lane 2 – Orchestrator: Record]",
     "Journey Memory records: step fingerprint, action result, all CX observation fields.\n"
     "Loop pattern check: if fingerprint repeats → inject loop warning for next step.\n"
     "If loop warning was already active and pattern persists → EXIT LOOP with 'loop_detected'.\n"
     "If action succeeded → reset consecutive failure counter. Return to top of loop."),
    ("[Lane 3 – Login Wall Branch (within loop)]",
     "When Decision Engine identifies a login wall:\n"
     "  OPTION A: DISMISS → continue loop.\n"
     "  OPTION B: ENTER_MOBILE + CLICK OTP → EXIT LOOP with 'login_wall_otp_requested'.\n"
     "  OPTION C: DONE → EXIT LOOP with 'login_wall_persona_abandoned'."),
    ("[Lane 2 → Lane 5 – Post-Journey Handoff]",
     "Journey loop exited. Terminal reason recorded. Journey log finalised.\n"
     "Evaluation Suite triggered: CX Evaluator → Design Evaluator → Content Analyzer → Accessibility Auditor.\n"
     "Each module receives the complete journey log and sampled screenshots."),
    ("[Lane 5 → Lane 6 – Evaluation to Reporting Handoff]",
     "All four evaluation modules complete. Structured results available.\n"
     "Report Generator produces: report.md (13 sections), journey_log.json, updates master_report.md, "
     "updates session_index.json."),
    ("[Lane 1 – Human Operator]",
     "Opens dashboard/index.html in any browser. Reviews findings. Prioritises P1 and P2 recommendations. "
     "Plans remediation with the CX / Digital team."),
]
for lane, desc in flow:
    p = doc.add_paragraph()
    para_spacing(p, before=60, after=40)
    r1 = p.add_run(lane + "\n")
    r1.bold = True; r1.font.size = Pt(10.5); r1.font.color.rgb = BFL_ORANGE
    r2 = p.add_run(desc)
    r2.font.size = Pt(10.5)

info_box(doc,
    "PMO Note: A formal Miro / Visio swimlane diagram illustrating the above flow with explicit "
    "hand-off arrows, decision diamonds at loop check points, login wall branch paths, and all "
    "terminal condition exit paths is being prepared and will be attached as Appendix A before "
    "final submission. The narrative above serves as the specification for that diagram.")

# ══════════════════════════════════════════════════════
# SECTION 12 – AGENT DEFINITIONS
# ══════════════════════════════════════════════════════
h1(doc, "12", "Agent Definitions and Requirements")
body(doc,
    "Per SOP V1.0 Section 12, the following documentation is provided for each agent before any "
    "agent is built. All five agents in the CX Audit Agent system are defined below with their "
    "name, business purpose, primary users, risk level, allowed actions, prohibited actions, "
    "human approval conditions, escalation scenarios, audit requirements, and human override mechanism.")

agents = [
    {
        "name": "Agent 1: Decision Engine Agent",
        "purpose": "At every step of a persona journey, receives the current screenshot, page state, persona context, and full journey history; and returns the single best next action a genuine human of that persona type would take — along with a complete CX observation record for that step.",
        "users": "Internal only — called by the Orchestrator at every journey step. Indirectly serves the CX / Digital team through the output audit reports.",
        "risk": "HIGH — this agent directly controls browser actions on a live BFSI website and produces observations that constitute official CX audit evidence acted upon by business stakeholders.",
        "can_do": [
            "Select any action from the defined catalogue: CLICK, TYPE, SCROLL DOWN, SCROLL UP, SCROLL LEFT, SCROLL RIGHT, NAVIGATE (recovery only), BACK, WAIT, PLAY VIDEO, DISMISS POPUP, SEARCH, DONE.",
            "Provide all 10 CX observation fields: observation, reasoning, emotion, cx_note, cognitive_load, trust_signals, unanswered_questions, guiding_factors, visible_content, state_of_mind.",
            "Make login wall decisions: dismissed, entered_mobile, or ignored — based on real-time persona reasoning.",
            "Request a NAVIGATE action to a URL already present in the current browser session history.",
            "Request a SEARCH action with 1–3 keywords reflecting the persona's intent.",
        ],
        "not_do": [
            "Navigate to any URL not already loaded in the current browser session history.",
            "Click the final Submit Application button or any button matching the final-submit pattern library. This is intercepted by the Controller before execution.",
            "Click Sign In, Login, Register, or My Account navigation links during a logged-out session.",
            "Report its own action failures (element not found, timeout) as website CX problems in the cx_note or observations.",
            "Predict or describe what will happen after the current action in the cx_note. The cx_note covers only what is currently visible before the action executes.",
            "Use training-data knowledge about bajajfinserv.in website structure to justify any action. Every action must be justified by what is physically visible in the current screenshot.",
            "Dismiss a side navigation drawer or menu panel that contains product links, section headings, or navigation items.",
        ],
        "human_approval": ["No Decision Engine action requires direct human approval during a live journey execution. Human oversight is provided through post-run report review."],
        "escalation": [
            "If the consecutive failure counter reaches CONSECUTIVE_FAIL_LIMIT (configured parameter), the Orchestrator terminates the journey without calling the Decision Engine again.",
            "If a loop warning has been active for more than one step and the fingerprint still repeats, the Orchestrator terminates with 'loop_detected' without calling the Decision Engine again.",
            "If the OpenAI API returns an error or timeout, the Orchestrator retries once; if the second call also fails, the failure counter is incremented.",
        ],
        "audit": [
            "Every Decision Engine call is logged in the JSONL step log with: full input summary, full output JSON, step number, timestamp.",
            "Token usage (prompt tokens + completion tokens) is recorded per call.",
            "Fabricated URL attempts are logged as policy violations with the rejected URL.",
            "Final-submit interceptions are logged with the intercepted button label and step number.",
            "Loop warning injections are logged with the repeating fingerprint.",
        ],
        "override": "The Human Operator may terminate any journey at any time by interrupting the Python process (Ctrl+C). The Orchestrator will complete the current step, write the partial journey log, and proceed to the Evaluation Suite with whatever data is available.",
    },
    {
        "name": "Agent 2: CX Evaluator Agent",
        "purpose": "After a persona journey completes, evaluates the full journey against the 18 canonical CX dimensions and produces a scored, evidence-grounded assessment with an overall journey verdict, emotional arc, and prioritised recommendations.",
        "users": "Internal only — called by the Orchestrator post-journey. Output is consumed by the Report Generator and displayed in the dashboard Dimension Scores View.",
        "risk": "MEDIUM — produces CX scores that directly influence business prioritisation and investment decisions. Dimension names must match the 18 canonical names exactly.",
        "can_do": [
            "Score each of the 18 CX dimensions from 0 to 10 with specific journey step evidence.",
            "Produce: journey verdict (one paragraph), TL;DR (2–3 sentences), key takeaways (3–5 insights), friction points (severity/page/description/customer impact), positive moments.",
            "Generate P1, P2, and P3 recommendations (each with area, specific action, expected impact).",
            "Describe the emotional journey arc stage by stage and produce a full persona emotional narrative.",
            "Access a sampled subset of journey screenshots (JPEG 72%, max 1280px).",
            "Reference specific step numbers and CX observation fields as evidence for scores.",
        ],
        "not_do": [
            "Modify the journey log or any step data.",
            "Use dimension names that deviate from the 18 canonical names in BRD v2.0 Section 8.1. No abbreviations, no snake_case, no '&' in place of 'and'.",
            "Generate recommendations that require code changes to the CX Audit Agent itself — scope is limited to findings about bajajfinserv.in.",
            "Access any external website, API, or data source other than the OpenAI Vision API called with the journey screenshots.",
        ],
        "human_approval": ["No human approval is required for the CX Evaluator to produce its output. Scores and recommendations are reviewed by the CX team after report generation."],
        "escalation": [
            "If the CX Evaluator API call fails (timeout, model error, invalid response), the failure is logged in the JSONL issue log and the persona report is generated with a clearly marked 'CX Evaluation unavailable — API error' notice in the relevant sections.",
        ],
        "audit": [
            "All 18 dimension scores with rationale and step evidence are written to journey_log.json under cx_evaluation.",
            "Token usage for evaluation calls is logged in the JSONL file.",
            "Per-dimension step-level evidence references (step numbers) are preserved in the JSON output.",
        ],
        "override": "Human CX reviewers may add manual annotations or comments to the Markdown report. These are reviewer notes only and do not alter the JSON data files or dashboard display.",
    },
    {
        "name": "Agent 3: Design Evaluator Agent",
        "purpose": "After a persona journey completes, assesses the visual design quality of pages visited during the journey, evaluated from the perspective of the specific persona.",
        "users": "Internal only — called post-journey. Output consumed by Report Generator for Section 8 of the persona report.",
        "risk": "LOW — informs design improvement decisions. Does not control any browser actions and does not produce scores used for compliance purposes.",
        "can_do": [
            "Assess: visual hierarchy effectiveness for this persona, colour contrast and text readability, touch target adequacy on mobile, CTA prominence and visual clarity, use of imagery and illustration, whitespace and visual density, design consistency across visited pages, form field design quality.",
            "Access sampled journey screenshots.",
            "Produce dimension scores, specific observations per page visited, friction points, and design recommendations (P1/P2/P3).",
        ],
        "not_do": [
            "Evaluate pages not visited during the journey.",
            "Propose specific colour hex codes, exact font sizes, or pixel-level spacing values — scope is qualitative assessment and directional recommendations only.",
        ],
        "human_approval": ["None required."],
        "escalation": ["API failure → log failure in JSONL issue log; generate report with 'Design Evaluation unavailable — API error' notice."],
        "audit": ["Design evaluation output written to journey_log.json under design_evaluation and included in persona report.md Section 8."],
        "override": "Human reviewers may annotate the Markdown report section.",
    },
    {
        "name": "Agent 4: Content Analyzer Agent",
        "purpose": "After a persona journey completes, assesses the quality, clarity, and appropriateness of written content on pages visited during the journey.",
        "users": "Internal only — called post-journey.",
        "risk": "LOW — informs content improvement decisions. Findings may have regulatory implications if financial disclosures are absent.",
        "can_do": [
            "Assess: financial clarity (are rates, fees, charges stated clearly before any login gate?), jargon appropriateness relative to persona's financial literacy, disclosure completeness, pre-gate content sufficiency, micro-copy quality (error messages, form labels, button labels, placeholder text), tone and voice appropriateness, and missing content gaps.",
            "Access sampled screenshots and extracted page text.",
        ],
        "not_do": [
            "Rewrite or suggest specific replacement copy text — scope is gap identification and directional recommendations only.",
            "Access content on pages not visited during the journey.",
        ],
        "human_approval": ["None required."],
        "escalation": ["API failure → log failure; generate report with 'Content Analysis unavailable — API error' notice."],
        "audit": ["Content analysis output written to journey_log.json under content_analysis and included in persona report.md Section 9."],
        "override": "Human reviewers may annotate the Markdown report section.",
    },
    {
        "name": "Agent 5: Accessibility Auditor Agent",
        "purpose": "After a persona journey completes, assesses the inclusivity and accessibility of the journey experience for the specific persona, using visual observation from screenshots aligned to WCAG principles.",
        "users": "Internal only — called post-journey.",
        "risk": "MEDIUM — accessibility findings may have regulatory implications (WCAG, RBI digital accessibility guidelines). Findings must be clearly scoped as visual assessment, not certified compliance.",
        "can_do": [
            "Assess: text size and readability for this persona, colour contrast across visited pages, touch target sizing adequacy on mobile viewport, language simplicity relative to persona's literacy level, presence of alt-text indicators, form field accessibility (label association, error identification), keyboard navigability signals, screen reader compatibility indicators in page structure.",
            "Access sampled screenshots and page state text.",
            "Produce accessibility friction points with severity levels and WCAG-aligned recommendations (P1/P2/P3).",
        ],
        "not_do": [
            "Run programmatic automated WCAG tests such as axe or Lighthouse — scope is visual assessment from screenshots only.",
            "Certify or claim WCAG compliance — findings are observations and recommendations, not a compliance certificate.",
        ],
        "human_approval": ["None required."],
        "escalation": ["API failure → log failure; generate report with 'Accessibility Audit unavailable — API error' notice."],
        "audit": ["Accessibility audit output written to journey_log.json under accessibility_audit and included in persona report.md Section 10."],
        "override": "Human reviewers may annotate the Markdown report section.",
    },
]

for agent in agents:
    h2(doc, "", agent["name"])
    bold_label(doc, "Business Purpose: ", agent["purpose"])
    bold_label(doc, "Primary Users: ", agent["users"])
    bold_label(doc, "Risk Level: ", agent["risk"])
    doc.add_paragraph()

    bold_label(doc, "What the Agent Can Do (Allowed Actions):", "")
    for item in agent["can_do"]:
        body(doc, item, bullet=True, indent_twips=360)

    bold_label(doc, "What the Agent Must NOT Do:", "")
    for item in agent["not_do"]:
        body(doc, item, bullet=True, indent_twips=360)

    bold_label(doc, "Actions Requiring Human Approval:", "")
    for item in agent["human_approval"]:
        body(doc, item, bullet=True, indent_twips=360)

    bold_label(doc, "Escalation and Termination Conditions:", "")
    for item in agent["escalation"]:
        body(doc, item, bullet=True, indent_twips=360)

    bold_label(doc, "Audit and Logging Requirements:", "")
    for item in agent["audit"]:
        body(doc, item, bullet=True, indent_twips=360)

    bold_label(doc, "Human Override Mechanism: ", agent["override"])
    doc.add_paragraph()

# ══════════════════════════════════════════════════════
# SECTION 13 – BRD ACCEPTANCE CRITERIA
# ══════════════════════════════════════════════════════
h1(doc, "13", "BRD Acceptance Criteria")
body(doc,
    "This PRD will be accepted by the PMO team and the Agentic AI Domain Team only when all of "
    "the following criteria are satisfied. Each criterion maps to a specific section of this PRD.")

criteria = [
    ("Use Case Defined (Section 7.1)",
     "A single-sentence use case statement is present and follows the SOP format: "
     "'This use case enables [actor] to [outcome] within [constraints] using [systems].'"),
    ("End-to-End Process Documented (Section 8)",
     "The complete process from session initialisation (Step 1.1) through journey execution, "
     "login wall handling, post-journey evaluation, and dashboard update (Step 5.4) is documented "
     "with no gaps or undefined branches between terminal conditions and report generation."),
    ("Steps Are Outcome-Driven (Section 8 all steps)",
     "Every step states: what happens, who performs it, what information is used, what decision "
     "is made, and what happens next. No step uses a prohibited phrase from Section 6.1."),
    ("Systems and Data Flows Listed (Sections 8 and 9)",
     "All systems (OpenAI API, Playwright Chromium, local file system, static HTML dashboard) are "
     "named in the relevant steps. All data elements (screenshot, page state, persona description, "
     "journey log, JSONL log) are explicitly identified with their source and destination."),
    ("UI Documented (Section 10.3)",
     "All seven dashboard screens are documented with: what is displayed, the exact JSON data source "
     "field names, and the specific error state display behaviour."),
    ("Agent Definitions Complete (Section 12)",
     "All five agents are documented with name, business purpose, risk level, allowed actions, "
     "prohibited actions, escalation conditions, audit requirements, and human override mechanism."),
    ("Failure and Escalation Paths Defined (Sections 8 and 12)",
     "All failure paths are explicitly documented: consecutive browser failures, loop detection, "
     "fabricated URL rejection, final-submit interception, popup loop (3 recurrences), API call "
     "failure for each of the four evaluation modules, configuration validation failure, and "
     "directory creation failure."),
    ("Terminal Conditions Defined (Referenced from BRD v2.0 Section 7)",
     "All seven terminal conditions are defined with precise trigger conditions: Goal Achieved, "
     "Login Wall – OTP Requested, Login Wall – Persona Abandoned, Valid Journey Abandoned, "
     "Loop Detected, Consecutive Failures, Max Steps Reached, Timeout."),
    ("Compliance and Audit Controls Defined (Sections 6.3 and 12)",
     "PII handling rules, JSONL logging requirements, screenshot storage policy, and the human "
     "override mechanism for each agent are all explicitly documented."),
    ("Approvals and Ownership Defined (Section 14)",
     "All four approver roles are named with their specific approval responsibilities and what "
     "they validate."),
]
for title, desc in criteria:
    bold_label(doc, f"{title}: ", desc, bullet=True)

# ══════════════════════════════════════════════════════
# SECTION 14 – APPROVALS AND OWNERSHIP
# ══════════════════════════════════════════════════════
h1(doc, "14", "Approvals and Ownership")
body(doc,
    "All four approvals below are mandatory before development or deployment of any component "
    "of the CX Audit Agent begins. Approval of a PoC or pilot run does not constitute approval "
    "for production deployment. A fresh approval review is required when scope, agent boundaries, "
    "or system integrations change materially.")

approval_data = [
    ("Process Owner\n(CX / Digital Team)", "Business approval",
     "Confirms the persona definitions, journey simulation logic, and evaluation dimensions accurately reflect the CX team's real objectives. Validates that the 18 CX dimension names and scoring rationale align with BFL's CX strategy. Confirms terminal condition definitions reflect authentic customer behaviour patterns observed in real research. Validates that the no-submission guardrail is operationally appropriate."),
    ("AI Security &\nCompliance Team", "Regulatory & policy approval",
     "Reviews PII handling rules (Section 6.3) for compliance with applicable data protection obligations. Reviews screenshot storage policy and OpenAI API data processing terms. Validates that the no-submission guardrail (final-submit interception) is technically enforced and is not bypassable through any action sequence. Reviews fabricated URL rejection mechanism and its JSONL audit trail. Confirms audit log retention policy is acceptable."),
    ("Agentic AI\nDomain Team", "Use case & BRD approval",
     "Confirms the use case statement (Section 7.1) meets SOP V1.0 format requirements. Reviews all five agent definitions (Section 12) for boundary completeness and clarity. Validates that no agent has undefined or unbounded permissions. Confirms all BRD acceptance criteria (Section 13) are fully met. Signs off on the Go/No-Go checklist (Section 16) before counter-signing."),
    ("Development /\nPlatform Team", "Technical feasibility review",
     "Verifies OpenAI GPT-4.1 Vision API availability, rate limits, and latency acceptable for production. Confirms Playwright Chromium compatibility with bajajfinserv.in without anti-bot blocking. Validates that the screenshot compression pipeline (JPEG 72%, max 1280px) is implementable consistently across all five agent modules. Reviews JSONL logging architecture for performance under maximum step counts. Confirms static dashboard data binding to session_index.json and journey_log.json is technically sound."),
]
app_tbl = doc.add_table(rows=1 + len(approval_data), cols=3)
app_tbl.style = 'Table Grid'
for i, h in enumerate(["Role", "Approval Responsibility", "What They Validate"]):
    cell = app_tbl.rows[0].cells[i]
    set_cell_bg(cell, '003387')
    set_cell_borders(cell, '003387')
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(h); r.bold = True; r.font.size = Pt(10); r.font.color.rgb = WHITE

for ri, (role, resp, validate) in enumerate(approval_data):
    row = app_tbl.rows[ri + 1]
    fill = 'EBF0FA' if ri % 2 == 0 else 'FFFFFF'
    for ci, val in enumerate([role, resp, validate]):
        cell = row.cells[ci]
        set_cell_bg(cell, fill)
        set_cell_borders(cell, 'D9D9D9')
        p = cell.paragraphs[0]
        r = p.add_run(val); r.font.size = Pt(10)
        if ci == 0:
            r.bold = True; r.font.color.rgb = BFL_BLUE

for i, w in enumerate([1.6, 1.6, 4.1]):
    for row in app_tbl.rows:
        row.cells[i].width = Inches(w)
doc.add_paragraph()

info_box(doc,
    "All four approvals are mandatory before development can begin. Approvals must be recorded "
    "with signature and date in the Sign-Off Sheet (Page 2 of this document). Any conditional "
    "approval must specify in writing the condition to be resolved and the timeline for re-review.",
    fill="EBF0FA")

# ══════════════════════════════════════════════════════
# SECTION 15 – GUIDING ENTERPRISE RULE
# ══════════════════════════════════════════════════════
h1(doc, "15", "Guiding Enterprise Rule")
p = doc.add_paragraph()
para_spacing(p, before=120, after=120)
set_para_shading(p, 'FFF3E0')
set_indent(p, 360)
pPr = p._p.get_or_add_pPr()
ind_el = OxmlElement('w:ind')
ind_el.set(qn('w:left'), '360'); ind_el.set(qn('w:right'), '360')
pPr.append(ind_el)
r = p.add_run(
    "If a process cannot be clearly explained from start to end in business language, "
    "it must not be automated using an agent.")
r.bold = True; r.font.size = Pt(12); r.font.color.rgb = AMBER
doc.add_paragraph()

body(doc,
    "Application to this product: Every step the CX Audit Agent takes — from capturing a "
    "screenshot to scoring a CX dimension to dismissing a popup — must be explainable to a "
    "CX manager at Bajaj Finance without reference to AI model internals, API schemas, "
    "probability distributions, or programming constructs.")
body(doc,
    "If any step in this PRD cannot be so explained, that step must be redesigned before it "
    "is built. This rule is the final arbitration standard for all disputes about scope, agent "
    "boundaries, and process step definitions during development, UAT, and post-launch review.")

# ══════════════════════════════════════════════════════
# SECTION 16 – GO / NO-GO CHECKLIST
# ══════════════════════════════════════════════════════
h1(doc, "16", "Go / No-Go Readiness Checklist (Mandatory)")
body(doc,
    "This checklist is a formal readiness gate required by BFL Enterprise Agentic Requirement "
    "SOP V1.0. Any No result is a No-Go and blocks development commencement. All items must be "
    "reviewed and signed off jointly by the four approving roles identified in Section 14.")

checklist_rows = [
    ("Use Case",    "A single-sentence use case statement in SOP format is present in Section 7.1."),
    ("Use Case",    "The use case names the actor, the specific outcome, the constraints, and all named systems."),
    ("Process",     "End-to-end process is fully documented from session initialisation (Step 1.1) to dashboard update (Step 5.4)."),
    ("Process",     "All seven terminal conditions and their precise trigger definitions are documented."),
    ("Process",     "Login wall handling covers all three outcomes: dismiss, enter mobile number, abandon."),
    ("Process",     "All failure paths are explicitly documented with specific conditions and recovery actions."),
    ("Steps",       "Every step in Section 8 states: what happens, who performs it, information used, decision logic, and next outcome."),
    ("Steps",       "No step in Section 8 uses a prohibited phrase from Section 6.1."),
    ("Steps",       "Every step decision has explicit outcome paths with no undefined branches."),
    ("Systems",     "All systems are named: OpenAI Vision API, Playwright Chromium, local file system, static HTML dashboard."),
    ("Systems",     "All data flows are explicit: screenshot → API → action → JSONL log → journey_log.json → dashboard."),
    ("Systems",     "API failure handling is defined for all five agent modules."),
    ("UI",          "All seven dashboard screens are documented with display content, JSON data source fields, and error states."),
    ("UI",          "UI documentation describes behaviour and validation — not visual design specifications."),
    ("Agent",       "All five agents have documented name, purpose, risk level, allowed actions, and prohibited actions."),
    ("Agent",       "All five agents have documented escalation/termination conditions and human override mechanism."),
    ("Agent",       "Decision Engine prohibited actions explicitly include: fabricated URLs, final-submit clicks, Sign In clicks in logged-out mode."),
    ("Agent",       "No agent has undefined or unbounded permissions — all action boundaries are specific."),
    ("Risk",        "PII handling rules are explicitly defined in Section 6.3."),
    ("Risk",        "Final-submit interception guardrail is documented as technically enforced and non-bypassable."),
    ("Risk",        "Loop detection and consecutive failure handling are defined with specific counter thresholds from configuration."),
    ("Risk",        "Fabricated URL rejection is defined and the rejection is logged as a violation in the JSONL issue log."),
    ("Ownership",   "All four approver roles are named in Section 14 with specific approval responsibilities."),
    ("Ownership",   "BRD acceptance criteria (Section 13) are fully met as verified by the Agentic AI Domain Team."),
    ("Ownership",   "Sign-off sheet (Page 2) is completed with signatures and dates from all four approvers."),
]
checklist_table(doc, checklist_rows)

info_box(doc,
    "FINAL DECISION\n\n"
    "Go / No-Go:  _______________     Date of readiness review:  _______________\n\n"
    "Signed by Process Owner: _______________     Date: _______________\n"
    "Signed by AI Security & Compliance: _______________     Date: _______________\n"
    "Signed by Agentic AI Domain Team: _______________     Date: _______________\n"
    "Signed by Development / Platform Team: _______________     Date: _______________\n\n"
    "Any No result requires the relevant section to be revised and the checklist to be "
    "re-reviewed before development can begin. Conditional Go decisions must state the "
    "condition and resolution timeline in writing.",
    fill="FFF3E0")

# ══════════════════════════════════════════════════════
# APPENDICES
# ══════════════════════════════════════════════════════
doc.add_page_break()

h1(doc, "Appendix A", "Output File Structure Reference")
body(doc, "The following directory structure is produced after a complete audit session:")
p = doc.add_paragraph()
para_spacing(p, before=40, after=40)
set_para_shading(p, 'F2F2F2')
set_indent(p, 360)
r = p.add_run(
    "reports/\n"
    "  {persona-slug}/\n"
    "    report.md                  Per-persona Markdown audit report (13 sections)\n"
    "    journey_log.json           Full structured data for dashboard\n"
    "    steps/\n"
    "      step_000.png             Screenshot at Step 0\n"
    "      step_001.png             Screenshot at Step 1\n"
    "      ...\n"
    "  master_report.md             Aggregated cross-persona master report\n\n"
    "logs/\n"
    "  {persona-slug}.jsonl         Step-by-step JSONL debug and audit log\n\n"
    "dashboard/\n"
    "  index.html                   Dashboard entry point (open in any browser)\n"
    "  session_index.json           Session index consumed by dashboard Run Audit Tab"
)
r.font.name = 'Courier New'
r.font.size = Pt(9.5)
doc.add_paragraph()

h1(doc, "Appendix B", "18 CX Dimension Reference")
body(doc,
    "The following dimension names are canonical per BRD v2.0 Section 8.1. They must appear "
    "exactly as written below in all reports, JSON files, and dashboard displays. Never abbreviated, "
    "never in snake_case, never with '&' substituted for 'and'.")
dims = [
    "1.   Discoverability and Information Architecture",
    "2.   Content Quality, Financial Clarity and Disclosure",
    "3.   Trust and Credibility Signals",
    "4.   Conversion and Task Flow Design",
    "5.   Emotional Experience and Persona Fit",
    "6.   Mobile and Touch Experience",
    "7.   Accessibility and Inclusive Design",
    "8.   Error Handling and Recovery",
    "9.   System Feedback and Load Experience",
    "10.  Micro-copy and Language Quality",
    "11.  Form Design and Data Collection UX",
    "12.  Navigation Depth and Efficiency",
    "13.  Personalisation and Context Awareness",
    "14.  Pre-Sales Support and Help Availability",
    "15.  Visual Hierarchy and Scannability",
    "16.  Consistency and Standards",
    "17.  User Control and Flexibility",
    "18.  Data Privacy and Consent UX",
]
for d in dims:
    body(doc, d)

h1(doc, "Appendix C", "Version History")
grid_table(doc,
    ["Version", "Date", "Description"],
    [("1.0", "May 2026", "Initial PRD submission to PMO team, authored in full compliance with BFL Enterprise Agentic Requirement SOP V1.0.")],
    [1.0, 1.5, 4.8])

# ══════════════════════════════════════════════════════
# SAVE
# ══════════════════════════════════════════════════════
output_path = r"C:\Users\Hiya Bhandari\Desktop\Folders\CX Audit Agent_OpenAI_Fix v11\CX_Audit_Agent_PRD_v1.0.docx"
doc.save(output_path)
print(f"PRD successfully saved to:\n{output_path}")
