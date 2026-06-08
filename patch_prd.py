"""
Patch CX_Audit_Agent_PRD_v1.0.docx:
  1. Fix all external BRD references → self-contained PRD references
  2. Add Appendices D–G: detailed evaluation & scoring frameworks
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

DOC_PATH = r"C:\Users\Hiya Bhandari\Desktop\Folders\CX Audit Agent_OpenAI_Fix v11\CX_Audit_Agent_PRD_v1.0.docx"

BFL_BLUE   = RGBColor(0x00, 0x33, 0x87)
BFL_ORANGE = RGBColor(0xFF, 0x66, 0x00)
WHITE      = RGBColor(0xFF, 0xFF, 0xFF)
BLACK      = RGBColor(0x00, 0x00, 0x00)
AMBER      = RGBColor(0xBF, 0x36, 0x00)

# ─────────────────────────────────────────────────────────────
# TEXT-REPLACEMENT HELPERS
# ─────────────────────────────────────────────────────────────

def replace_in_para(para, old, new):
    """Replace all occurrences of `old` with `new` across runs in a paragraph."""
    full = "".join(r.text for r in para.runs)
    if old not in full:
        return False
    updated = full.replace(old, new)
    # Write back: put everything in first run, blank the rest
    if para.runs:
        para.runs[0].text = updated
        for r in para.runs[1:]:
            r.text = ""
    return True


def replace_in_doc(doc, old, new):
    count = 0
    for para in doc.paragraphs:
        if replace_in_para(para, old, new):
            count += 1
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    if replace_in_para(para, old, new):
                        count += 1
    return count


# ─────────────────────────────────────────────────────────────
# XML / FORMATTING HELPERS  (same as generate_prd.py)
# ─────────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color):
    tc = cell._tc; tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'),'clear'); shd.set(qn('w:color'),'auto'); shd.set(qn('w:fill'),hex_color)
    tcPr.append(shd)

def set_cell_borders(cell, color="D9D9D9"):
    tc = cell._tc; tcPr = tc.get_or_add_tcPr()
    tcB = OxmlElement('w:tcBorders')
    for side in ['top','left','bottom','right']:
        b = OxmlElement(f'w:{side}')
        b.set(qn('w:val'),'single'); b.set(qn('w:sz'),'4')
        b.set(qn('w:space'),'0'); b.set(qn('w:color'),color)
        tcB.append(b)
    tcPr.append(tcB)

def para_spacing(p, before=0, after=0):
    pPr = p._p.get_or_add_pPr()
    sp = OxmlElement('w:spacing')
    sp.set(qn('w:before'),str(before)); sp.set(qn('w:after'),str(after))
    pPr.append(sp)

def set_indent(p, left=0):
    pPr = p._p.get_or_add_pPr()
    ind = OxmlElement('w:ind'); ind.set(qn('w:left'),str(left))
    pPr.append(ind)

def set_shading(p, fill):
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'),'clear'); shd.set(qn('w:color'),'auto'); shd.set(qn('w:fill'),fill)
    pPr.append(shd)

def h1(doc, number, title):
    p = doc.add_paragraph(); para_spacing(p,200,80); set_shading(p,'003387'); set_indent(p,120)
    r = p.add_run(f"{number}. {title.upper()}")
    r.bold=True; r.font.size=Pt(13); r.font.color.rgb=WHITE

def h2(doc, title):
    p = doc.add_paragraph(); para_spacing(p,160,60)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    left = OxmlElement('w:left')
    left.set(qn('w:val'),'single'); left.set(qn('w:sz'),'18')
    left.set(qn('w:space'),'4'); left.set(qn('w:color'),'FF6600')
    pBdr.append(left); pPr.append(pBdr); set_indent(p,180)
    r = p.add_run(title); r.bold=True; r.font.size=Pt(11.5); r.font.color.rgb=BFL_BLUE

def h3(doc, title):
    p = doc.add_paragraph(); para_spacing(p,100,40); set_indent(p,360)
    r = p.add_run(title); r.bold=True; r.font.size=Pt(11); r.font.color.rgb=BFL_ORANGE

def body(doc, text, bullet=False, indent=0):
    p = doc.add_paragraph(style='List Bullet') if bullet else doc.add_paragraph()
    para_spacing(p,40,40)
    if indent: set_indent(p,indent)
    r = p.add_run(text); r.font.size=Pt(10.5)

def bold_label(doc, label, text, bullet=False, indent=0):
    p = doc.add_paragraph(style='List Bullet') if bullet else doc.add_paragraph()
    para_spacing(p,40,40)
    if indent: set_indent(p,indent)
    r1 = p.add_run(label); r1.bold=True; r1.font.size=Pt(10.5); r1.font.color.rgb=BFL_BLUE
    r2 = p.add_run(text);  r2.font.size=Pt(10.5)

def info_box(doc, text, fill="EBF0FA"):
    tbl = doc.add_table(rows=1,cols=1); tbl.alignment=WD_TABLE_ALIGNMENT.LEFT
    cell=tbl.cell(0,0); set_cell_bg(cell,fill); set_cell_borders(cell,'003387')
    p=cell.paragraphs[0]; r=p.add_run(text)
    r.italic=True; r.font.size=Pt(10)
    r.font.color.rgb=BFL_BLUE if fill=="EBF0FA" else AMBER
    doc.add_paragraph()

def grid_table(doc, headers, rows_data, col_widths=None):
    tbl = doc.add_table(rows=1+len(rows_data), cols=len(headers))
    tbl.style='Table Grid'; tbl.alignment=WD_TABLE_ALIGNMENT.LEFT
    hr = tbl.rows[0]
    for i,h in enumerate(headers):
        cell=hr.cells[i]; set_cell_bg(cell,'003387'); set_cell_borders(cell,'003387')
        p=cell.paragraphs[0]; p.alignment=WD_ALIGN_PARAGRAPH.CENTER
        r=p.add_run(h); r.bold=True; r.font.size=Pt(10); r.font.color.rgb=WHITE
    for ri,row_data in enumerate(rows_data):
        row=tbl.rows[ri+1]; fill='F2F2F2' if ri%2==0 else 'FFFFFF'
        for ci,val in enumerate(row_data):
            cell=row.cells[ci]; set_cell_bg(cell,fill); set_cell_borders(cell,'D9D9D9')
            p=cell.paragraphs[0]; r=p.add_run(str(val)); r.font.size=Pt(10)
    if col_widths:
        for i,w in enumerate(col_widths):
            for row in tbl.rows: row.cells[i].width=Inches(w)
    doc.add_paragraph()

def scoring_table(doc, rows):
    """Scoring rubric table: Band | Score | Description."""
    headers = ["Score Band", "Score", "What It Means"]
    colors  = ['1a7a3c','2d8a4e','d97706','c0392b','7c0000']
    tbl = doc.add_table(rows=1+len(rows), cols=3)
    tbl.style='Table Grid'; tbl.alignment=WD_TABLE_ALIGNMENT.LEFT
    hr = tbl.rows[0]
    for i,h in enumerate(headers):
        cell=hr.cells[i]; set_cell_bg(cell,'003387'); set_cell_borders(cell,'003387')
        p=cell.paragraphs[0]; p.alignment=WD_ALIGN_PARAGRAPH.CENTER
        r=p.add_run(h); r.bold=True; r.font.size=Pt(10); r.font.color.rgb=WHITE
    band_fills = ['dcfce7','d1fae5','fef9c3','fee2e2','fecaca']
    for ri,(band,score,desc) in enumerate(rows):
        row=tbl.rows[ri+1]
        for ci,val in enumerate([band,score,desc]):
            cell=row.cells[ci]
            set_cell_bg(cell, band_fills[ri] if ci<2 else 'FFFFFF')
            set_cell_borders(cell,'D9D9D9')
            p=cell.paragraphs[0]
            r=p.add_run(val); r.font.size=Pt(10)
            if ci<2: r.bold=True
    for i,w in enumerate([1.4,0.8,5.1]):
        for row in tbl.rows: row.cells[i].width=Inches(w)
    doc.add_paragraph()


# ─────────────────────────────────────────────────────────────
# STEP 1 — LOAD AND FIX BRD REFERENCES
# ─────────────────────────────────────────────────────────────

doc = Document(DOC_PATH)

replacements = [
    # External BRD document references → self-contained references
    ("BRD v2.0 Section 8.1",          "Appendix B of this PRD"),
    ("BRD v2.0 Section 7",            "Section 8 of this PRD"),
    ("BRD v2.0",                       "this PRD"),
    ("out of BRD scope",               "out of scope for this PRD"),
    # Tidy up residuals
    ("this PRD, this PRD, and associated SOPs",
                                       "this PRD and associated SOPs"),
    ("this PRD, this PRD",             "this PRD"),
    # Appendix B description line
    ("canonical per Appendix B of this PRD. They must appear exactly as written below",
     "canonical as defined below. They must appear exactly as written"),
]

print("── Fixing BRD references ──")
for old, new in replacements:
    n = replace_in_doc(doc, old, new)
    if n:
        print(f"  Replaced {n}× : '{old[:60]}...' → '{new[:60]}'")


# ─────────────────────────────────────────────────────────────
# STEP 2 — UPDATE APPENDIX C VERSION ROW (minor wording)
# ─────────────────────────────────────────────────────────────
replace_in_doc(doc,
    "Initial PRD submission to PMO team, authored in full compliance with BFL Enterprise Agentic Requirement SOP V1.0.",
    "Initial PRD submission to PMO team. All evaluation and scoring frameworks included inline. Authored in full compliance with BFL Enterprise Agentic Requirement SOP V1.0.")

# ─────────────────────────────────────────────────────────────
# STEP 3 — ADD PAGE BREAK + NEW APPENDICES
# ─────────────────────────────────────────────────────────────

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# APPENDIX D — CX EVALUATION FRAMEWORK (18 DIMENSIONS)
# ══════════════════════════════════════════════════════════════

h1(doc, "Appendix D", "CX Evaluation Framework — 18 Dimensions")
body(doc,
    "This appendix defines the complete evaluation framework for the CX Evaluator Agent. "
    "For each of the 18 canonical CX dimensions, the following is specified: what the dimension "
    "evaluates, a five-band scoring rubric (0–10), observable criteria grounded in screenshot "
    "evidence, and the most common friction patterns observed in BFSI website journeys.")
info_box(doc,
    "Scoring Scale:  9–10 = Excellent  |  7–8 = Good  |  5–6 = Moderate  |  3–4 = Poor  |  0–2 = Very Poor\n"
    "All scores must be grounded in specific step evidence from the journey log. A score without "
    "an evidence reference is invalid.")

# ── DIMENSION DEFINITIONS ──
CX_DIMENSIONS = [

  ("1", "Discoverability and Information Architecture",
   "How easily the persona can find their target product or information through the site's navigation structure, labeling, and organisational logic.",
   [
    ("Excellent","9–10","Target product immediately discoverable. Navigation labels perfectly match this persona's vocabulary and intent. Multiple clear pathways exist. Logical product grouping. Site search produces accurate results on first query."),
    ("Good","7–8","Target product discoverable through 2–3 clicks. Navigation labels mostly aligned with persona's mental model. Minor labeling ambiguity but resolvable. Site structure logical."),
    ("Moderate","5–6","Product discoverable after 4–5 clicks with some effort. Navigation labels partially misaligned with persona's vocabulary. Some inconsistent categorisation. Search partially useful."),
    ("Poor","3–4","Product requires significant exploration across 6+ clicks. Navigation labels unclear or counterintuitive for this persona. Inconsistent hierarchy. Search results unhelpful."),
    ("Very Poor","0–2","Target product completely undiscoverable through visible navigation. No logical structure. No functional search. Persona cannot make meaningful progress."),
   ],
   ["Navigation label clarity for this persona's vocabulary", "Number of clicks required to reach target content", "Search bar presence and result relevance", "Breadcrumb presence and accuracy", "Product category grouping logic", "Hamburger menu / side drawer clarity on mobile"],
   ["Product buried under non-intuitive category names", "Search returns irrelevant results for financial product queries", "Navigation drawer dismissed as popup due to labeling ambiguity", "Sub-categories not visible until after unnecessary scrolling"]),

  ("2", "Content Quality, Financial Clarity and Disclosure",
   "Whether financial terms, interest rates, fees, charges, eligibility criteria, and regulatory disclosures are communicated clearly, accurately, and completely — and at the right point in the journey.",
   [
    ("Excellent","9–10","All key financial figures (rate, fee, tenure, eligibility) stated clearly and prominently before any login gate. Disclosures fully present. Language pitched perfectly for this persona's literacy level. No hidden charges."),
    ("Good","7–8","Most financial information clear and accessible pre-gate. Regulatory disclosures present. Minor gaps in prominence or terminology for this persona."),
    ("Moderate","5–6","Basic financial information present but not prominent. Key figures require effort to find. Partial disclosures. Some financial jargon uncorrected for this persona."),
    ("Poor","3–4","Key financial information sparse, vague, or buried. Major disclosure gaps. Heavy use of financial jargon inappropriate for this persona's literacy level."),
    ("Very Poor","0–2","Key financial information absent or hidden entirely behind login gate. Regulatory disclosures missing. Persona cannot make an informed decision from public content."),
   ],
   ["Interest rate / APR visibility pre-login", "Fee and charge breakdown availability", "Eligibility criteria clarity", "Regulatory disclosure presence (RBI registration, T&C, KFS)", "Language complexity relative to persona's financial literacy", "Key Financial Statement (KFS) or Fact Sheet availability"],
   ["Interest rates visible only post-login", "Processing fees disclosed only in step 4 of application", "Eligibility buried in FAQ with 8+ accordion clicks", "Disclaimers in 8pt font below fold", "EMI calculator present but formula unexplained"]),

  ("3", "Trust and Credibility Signals",
   "Presence, prominence, placement, and effectiveness of trust-building elements — regulatory badges, security indicators, social proof, brand consistency, and data safety messaging.",
   [
    ("Excellent","9–10","Comprehensive trust architecture. RBI/NBFC registration visible. SSL indicator present. Star ratings and testimonials contextually placed. Trust signals appear precisely at key decision moments (before data entry, before login wall). Brand consistency reinforces credibility."),
    ("Good","7–8","Multiple trust signals well-placed. Regulatory indication clearly visible. Security messaging present. Some social proof. Trust signals not perfectly timed but broadly accessible."),
    ("Moderate","5–6","Some trust signals visible but not prominent. Single regulatory badge present (typically in footer). Basic security indicator. No social proof at key moments."),
    ("Poor","3–4","Minimal trust signals. Regulatory badge buried in footer. No contextual trust messaging. No ratings or testimonials visible during journey. This persona would feel uncertain."),
    ("Very Poor","0–2","No visible trust signals at any point in the journey. No RBI / regulatory indicators. No security messaging. No social proof. For this persona, the site fails the basic credibility test."),
   ],
   ["RBI / NBFC registration badge visibility and placement", "SSL / security certificate indicator", "Star ratings and customer count display", "Testimonials or case study presence", "Partner / award logos", "Privacy policy link accessibility", "Data safety messaging at point of data entry"],
   ["RBI badge only in footer, not visible on product pages", "No security messaging before mobile number entry", "Testimonials present but clearly generic / non-Indian", "Brand elements inconsistent across pages, eroding trust"]),

  ("4", "Conversion and Task Flow Design",
   "How effectively the site guides this persona through their intended task — clarity of the conversion funnel, CTA quality, step visibility, and minimisation of friction on the path to goal completion.",
   [
    ("Excellent","9–10","Primary CTA immediately visible above fold. Task flow clearly signposted at every step. Progress indicator present for multi-step flows. No dead ends. Persona reaches goal completion with zero structural friction."),
    ("Good","7–8","Primary CTA clearly visible. Task flow well-structured. Minor friction in 1–2 steps. Persona reaches goal with manageable effort."),
    ("Moderate","5–6","CTA present but not optimally placed. Task flow functional but with 2–3 friction points. Some steps require guessing. Persona can complete task but with noticeable effort."),
    ("Poor","3–4","CTA difficult to find or poorly labelled. Task flow fragmented with multiple unclear steps. Multiple dead ends. Persona struggles significantly to progress."),
    ("Very Poor","0–2","No clear CTA visible. Task flow non-existent or broken. Persona unable to make meaningful progress toward their goal through any visible path."),
   ],
   ["Primary CTA visibility above fold on mobile", "CTA label specificity and action-orientation", "Step count to reach goal", "Progress indicator presence in multi-step flows", "Back navigation preservation of form data", "Dead-end page count in journey"],
   ["'Apply Now' CTA below fold on 360px mobile", "CTA labelled 'Know More' instead of product-specific action", "5-step flow with no progress indicator", "Form data lost on browser back", "Application funnel exits to homepage on error"]),

  ("5", "Emotional Experience and Persona Fit",
   "Whether the site's tone, pacing, language, and design create an emotional experience aligned with this persona's needs — confidence-building, anxiety-reducing, and appropriately engaging.",
   [
    ("Excellent","9–10","Experience consistently builds confidence throughout journey. Tone perfectly calibrated to this persona's demographic, financial literacy, and emotional state. Site proactively addresses anxieties (data safety, hidden costs, commitment) at the right moments. Persona feels understood."),
    ("Good","7–8","Predominantly positive emotional journey. Tone largely appropriate. Some minor moments of uncertainty but well-recovered. Persona remains engaged and motivated."),
    ("Moderate","5–6","Mixed emotional journey. Positive moments offset by friction-induced anxiety. Tone partially appropriate. Persona experiences doubt at key decision points but continues."),
    ("Poor","3–4","Predominantly frustrating or anxiety-inducing experience. Tone mismatched to persona. Site fails to address key emotional concerns. Persona's confidence erodes significantly during journey."),
    ("Very Poor","0–2","Experience actively creates anxiety, confusion, or mistrust. Tone completely inappropriate for this persona type. Persona's emotional state deteriorates from landing page onward."),
   ],
   ["Tone appropriateness for this persona's age / occupation / literacy", "Anxiety-addressing messaging at data entry points", "Confidence signals at application commitment points", "Reassurance messaging after login wall", "Emotional arc progression across journey steps"],
   ["Aggressive urgency messaging ('Offer expires in 2 hours') creating pressure on risk-averse persona", "No reassurance before requesting Aadhaar / PAN", "Complex financial language creating anxiety in low-literacy persona", "Pop-ups interrupting flow at high-engagement moments"]),

  ("6", "Mobile and Touch Experience",
   "Quality of the mobile-specific UX for the exact viewport used in the audit, including touch target sizing, mobile navigation patterns, viewport optimisation, and gesture support.",
   [
    ("Excellent","9–10","Fully mobile-optimised. All touch targets ≥ 44×44px. Bottom navigation strip intuitive. One-thumb operation supported. No horizontal scroll on main content. Images and text perfectly sized for viewport. Carousel gestures work correctly."),
    ("Good","7–8","Good mobile experience. Most touch targets adequately sized. Mobile navigation functional. Minor layout issues on specific pages. Generally comfortable for one-thumb use."),
    ("Moderate","5–6","Functional mobile experience with notable pain points. Some touch targets too small. Navigation accessible but not optimised for mobile. Minor horizontal scroll or overflow issues."),
    ("Poor","3–4","Multiple mobile usability failures. Several touch targets too small to tap accurately. Navigation difficult on mobile. Some desktop-specific layout elements visible. Pinch-to-zoom frequently required."),
    ("Very Poor","0–2","Site not meaningfully mobile-optimised for this viewport. Desktop layout rendered at mobile resolution. Touch targets systematically too small. Horizontal scroll on main content. Critical content cut off."),
   ],
   ["Touch target size (minimum 44×44px per WCAG 2.5.5)", "Horizontal scroll presence on main content", "Bottom navigation strip presence and usability", "Text size at 360px viewport width", "Image scaling and compression", "Carousel touch gesture support", "Keyboard popup behaviour on form fields"],
   ["Product CTA button 28×20px — untappable on mobile", "Hero banner forces horizontal scroll at 360px", "Desktop navigation rendered as-is on mobile", "Date picker opens numeric keyboard instead of date picker on mobile"]),

  ("7", "Accessibility and Inclusive Design",
   "How inclusive the experience is for this specific persona — covering visual accessibility (contrast, text size), cognitive accessibility (language complexity), and basic WCAG alignment observable from screenshots.",
   [
    ("Excellent","9–10","WCAG AA contrast throughout. Text ≥16px on all body content. Language perfectly calibrated to this persona's literacy level. Form fields fully labelled. Error messages identify specific fields. Visual information never conveyed by colour alone."),
    ("Good","7–8","WCAG AA contrast on primary content. Adequate text sizing. Language mostly appropriate. Most form fields labelled. Minor accessibility gaps on secondary elements."),
    ("Moderate","5–6","WCAG AA contrast on most elements but with notable failures on secondary text or UI. Language partially accessible for this persona. Form fields mostly labelled."),
    ("Poor","3–4","Multiple contrast failures on important content. Text too small in key areas for this persona. Language significantly mismatched to this persona's literacy level. Several unlabelled form fields."),
    ("Very Poor","0–2","Systematic contrast failures. Text size inadequate throughout. Language completely inaccessible for this persona. Form fields unlabelled. Error states not identified accessibly. Critical information conveyed by colour only."),
   ],
   ["Text contrast ratio (WCAG AA = 4.5:1 for normal text)", "Body text size at viewport (minimum 16px)", "Form field label presence and association", "Error message specificity", "Alt text on informational images", "Colour-alone usage for critical information", "Language reading level relative to persona"],
   ["Grey text on white background failing WCAG AA at 2.8:1", "12px disclaimer text rendering below 1em on mobile", "OTP field with no visible label", "Error state shown only in red without text description", "Rate comparison chart readable only by colour coding"]),

  ("8", "Error Handling and Recovery",
   "How the site handles user mistakes, failed actions, and edge cases — quality of error messages, clarity of recovery instructions, and smoothness of retry flows.",
   [
    ("Excellent","9–10","All errors produce specific, field-level messages in plain language appropriate for this persona. Recovery path clearly stated within the error message. Inline validation prevents errors before submission. Form state preserved on error. Retry seamless."),
    ("Good","7–8","Most errors produce clear, helpful messages. Recovery guidance present. Minor gaps in inline validation. Form state mostly preserved. Retry functional."),
    ("Moderate","5–6","Basic error messages present. Some recovery guidance. Inline validation partial. Form data sometimes lost on error. Retry possible but requiring effort."),
    ("Poor","3–4","Generic error messages ('Something went wrong'). No recovery guidance. No inline validation. Form data frequently lost on error. Retry path unclear."),
    ("Very Poor","0–2","No meaningful error handling. Error states either absent or system-level (HTTP errors). No recovery path visible. Form data consistently lost. User left with no guidance."),
   ],
   ["Error message specificity (field-level vs. page-level)", "Plain language in error messages for this persona", "Recovery instruction completeness", "Inline validation presence and timing", "Form state preservation on back navigation or error", "Retry path visibility"],
   ["'Invalid input' with no field identification", "'Please try again later' with no timeline or alternative", "Form resets entirely on OTP timeout", "Inline validation triggers on every keypress causing anxiety for slow typers"]),

  ("9", "System Feedback and Load Experience",
   "Quality and completeness of system status communication — loading states, progress indicators, action confirmations, and communication of wait times.",
   [
    ("Excellent","9–10","All heavy operations show loading indicators. Action confirmations are specific and include next steps. Progress preserved through interruptions. Page load time acceptable on typical Indian mobile network. System status always clear."),
    ("Good","7–8","Most heavy operations show loading states. Confirmations present and reasonably specific. Minor gaps in status communication. Performance acceptable."),
    ("Moderate","5–6","Loading indicators present on some operations. Confirmations vague but present. Some silent delays. Performance variable."),
    ("Poor","3–4","Loading indicators largely absent. Confirmations generic or missing. Multiple silent delays of 3–5 seconds with no feedback. User uncertain if action was received."),
    ("Very Poor","0–2","No loading feedback. No action confirmations. Long silent delays throughout. User repeatedly clicks CTAs thinking action failed. System status never communicated."),
   ],
   ["Spinner or skeleton screen on page load", "Loading indicator on form submission", "Confirmation message specificity after action", "Page transition speed (subjective at mobile network speed)", "Progress preservation on interruption (call received during application)"],
   ["5-second blank white screen between pages with no loading indicator", "OTP sent with no confirmation message — user doesn't know to check phone", "Application form shows spinner for 8 seconds with no ETA", "Back navigation during loading causes silent session reset"]),

  ("10", "Micro-copy and Language Quality",
   "Quality of all small text elements — button labels, placeholder text, helper text, tooltips, form instructions, confirmation messages, error labels — in terms of clarity, action-orientation, and appropriateness for this persona.",
   [
    ("Excellent","9–10","Every micro-copy element is specific, action-oriented, and calibrated to this persona's literacy level. Button labels describe the exact outcome. Placeholder text provides useful examples. Helper text prevents errors proactively. Consistent voice throughout."),
    ("Good","7–8","Most micro-copy clear and helpful. Button labels mostly action-oriented. Some placeholder examples. Minor inconsistencies in tone. Helper text present on most complex fields."),
    ("Moderate","5–6","Adequate micro-copy. Some generic labels ('Submit', 'Continue') without context. Placeholder text sometimes absent. Helper text sparse. Tone mostly appropriate."),
    ("Poor","3–4","Largely generic micro-copy. Most buttons labelled generically. No placeholder examples. No helper text on complex fields. Tone mismatched to persona. Technical language present."),
    ("Very Poor","0–2","Poor or absent micro-copy throughout. Button labels cryptic or system-generated. No placeholders. No helper text. Technical error codes visible. Micro-copy actively misleads or confuses this persona."),
   ],
   ["CTA button label specificity ('Apply for Gold Loan' vs 'Apply Now' vs 'Submit')", "Form field placeholder text usefulness", "Helper text presence on Aadhaar / PAN / income fields", "Tooltip content quality", "Confirmation message specificity", "Technical jargon in visible copy"],
   ["'OK' and 'Cancel' buttons in OTP dialog — ambiguous for low-literacy persona", "Income field placeholder shows 'Enter amount' instead of '₹ e.g. 35000'", "PAN field with no explanation of why it is required", "Financial jargon ('DPD', 'LTV ratio') unremediated in product copy"]),

  ("11", "Form Design and Data Collection UX",
   "Quality of form design including label placement, input type appropriateness, validation timing, field sequencing, and the overall data collection experience for this persona.",
   [
    ("Excellent","9–10","All form labels clearly placed and persistently visible. Correct input types for all fields (numeric keyboard for phone, date picker for DOB). Minimal required fields — nothing collected before it is needed. Inline validation on blur. Logical sequence respects persona's journey context."),
    ("Good","7–8","Labels clear and visible. Most input types appropriate. Reasonable field sequence. Inline validation on most fields. Minor unnecessary fields present."),
    ("Moderate","5–6","Labels adequate but some disappear on focus (placeholder-only labels). Some wrong input types. Validation only on submit. Sequence mostly logical."),
    ("Poor","3–4","Some labels unclear or absent. Multiple wrong input types. No inline validation. Illogical field sequence. Several fields collecting data prematurely in the journey."),
    ("Very Poor","0–2","Labels absent or confusing. Wrong input types throughout. No validation. Completely illogical sequence. Extensive data collected before any product information shown. Form designed for system convenience, not user experience."),
   ],
   ["Label placement (above field, not placeholder-only)", "Keyboard type triggered by input (numeric/text/email/date)", "Number of fields on single screen", "Inline validation timing (on blur vs on submit)", "Field sequence logic relative to persona's decision stage", "Progress indicator in multi-step forms"],
   ["Floating label disappears on focus — user forgets what field they are filling", "Mobile number field triggering alphabetic keyboard", "PAN, Aadhaar, income all required before rate is revealed", "12-field form on single screen with no grouping", "Date field requiring DD/MM/YYYY format with no example"]),

  ("12", "Navigation Depth and Efficiency",
   "Number of interactions required to reach the target content and availability of efficient pathways for this persona's specific intent.",
   [
    ("Excellent","9–10","Target content reachable in 1–2 interactions from landing page. Multiple parallel pathways available. Quick links or shortcuts aligned with this persona's intent visible above fold. No dead ends encountered."),
    ("Good","7–8","Target content reachable in 2–3 interactions. At least two clear pathways available. Minor inefficiencies. No material dead ends."),
    ("Moderate","5–6","Target content reachable in 4–5 interactions. Limited pathway options. Some navigational inefficiency but resolvable. Occasional dead ends."),
    ("Poor","3–4","Target content requires 6–8 interactions. Single pathway only. Frequent dead ends. Significant backtracking required."),
    ("Very Poor","0–2","Target content unreachable within the step budget via visible navigation. All pathways lead to dead ends or login gates before product information. Persona cannot make meaningful progress."),
   ],
   ["Step count to first product information page", "Step count to first rate / fee information", "Number of pathways available to target product", "Dead end page count", "Breadcrumb presence reducing backtracking", "Internal search effectiveness for this persona's query"],
   ["Gold loan reachable only via: Home > Products > Secured Loans > Loans Against Assets > Gold Loan (5 clicks)", "Search for 'gold loan' returns customer service FAQ, not product page", "Back button returns to homepage instead of previous category page"]),

  ("13", "Personalisation and Context Awareness",
   "Degree to which the site adapts its content, recommendations, and UI to the user's demonstrated context, location, device, and journey behaviour.",
   [
    ("Excellent","9–10","Site actively personalises content to this persona's context (detected location, device type, browsing history). Product recommendations align with stated intent. Pre-filled information where available. Contextually relevant offers shown at appropriate moments."),
    ("Good","7–8","Some personalisation present. Location-based content relevant. Basic intent-aware recommendations. Pre-fill where technically possible."),
    ("Moderate","5–6","Minimal personalisation. Location detected but used only for branch locator. Generic product recommendations. No pre-fill."),
    ("Poor","3–4","Near-generic experience. Personalisation limited to single element (e.g. city name in heading). Recommendations unrelated to persona's intent."),
    ("Very Poor","0–2","Completely generic experience identical for all visitors regardless of context, device, or demonstrated intent. Zero personalisation signals."),
   ],
   ["Location-based content relevance", "Device-appropriate layout (mobile vs desktop)", "Intent-aware product recommendations", "Pre-filled fields in forms (for returning users)", "Contextually triggered offers"],
   ["Mumbai user shown Chennai branch contact details as default", "Same hero banner shown regardless of product page entry intent", "No pre-fill of mobile number already verified at login wall"]),

  ("14", "Pre-Sales Support and Help Availability",
   "Availability, accessibility, discoverability, and usefulness of help resources before the user commits — FAQs, chat, callback, comparison tools, EMI calculators, and eligibility checkers.",
   [
    ("Excellent","9–10","Proactive help offered contextually at decision moments. EMI calculator or eligibility checker visible on product page. FAQ relevant to this persona's specific concerns accessible within 1 click. Live chat or callback available. Help resources calibrated to this persona's literacy level."),
    ("Good","7–8","Help resources available and accessible. Calculator present. FAQ findable within 2 clicks. Chat available but not proactively offered. Content mostly appropriate for this persona."),
    ("Moderate","5–6","Basic help resources present. Generic FAQ accessible but not contextual. Calculator present but hard to find. Chat available as persistent widget."),
    ("Poor","3–4","Minimal help resources. FAQ only accessible via footer. No calculator on product page. Chat either absent or unreliable. Help content in complex language."),
    ("Very Poor","0–2","No pre-sales support resources visible during the journey. No FAQ, calculator, chat, or callback option encountered. Persona has no mechanism to resolve doubts before being asked for personal data."),
   ],
   ["EMI / eligibility calculator presence on product page", "FAQ proximity to product information", "Live chat widget visibility and responsiveness", "Callback option availability", "Help content language level appropriateness"],
   ["EMI calculator buried in footer > Tools section", "FAQ opens in new browser tab breaking journey flow", "Chat widget loads but agent never responds", "Eligibility checker requires login before showing result"]),

  ("15", "Visual Hierarchy and Scannability",
   "How effectively the page's visual design directs this persona's attention and enables them to quickly scan and extract the specific information they are looking for.",
   [
    ("Excellent","9–10","Strong visual hierarchy guides this persona's eye directly to relevant information. F-pattern or Z-pattern layout supports scanning behaviour. Key figures (rates, CTAs, trust signals) immediately visible without reading full content. Content density appropriate for this persona."),
    ("Good","7–8","Clear visual hierarchy. Key information prominent. Most important content scannable without reading in full. Minor secondary elements compete with primary content."),
    ("Moderate","5–6","Some hierarchy present. Headings distinct from body. Key information accessible with moderate scanning effort. Some visual competition for attention."),
    ("Poor","3–4","Weak visual hierarchy. Key information not distinguished from secondary content. Dense text blocks requiring full reading. This persona would miss critical information."),
    ("Very Poor","0–2","No meaningful visual hierarchy. All text same visual weight. Key information buried in dense paragraphs. Persona with this literacy level cannot efficiently locate needed information through scanning."),
   ],
   ["Heading / body contrast", "Key figures (rate, CTA, fee) visual prominence", "White space usage", "Content density per screen relative to persona's literacy level", "Banner image vs information balance on mobile"],
   ["Interest rate buried in 4th paragraph of dense text block", "Product hero banner takes 70% of viewport leaving no above-fold information", "Star rating displayed at 10px size below the fold", "Mobile screen showing 3 lines of real content and 70% image"]),

  ("16", "Consistency and Standards",
   "Consistency of UI patterns, interaction models, terminology, and visual design language across all pages visited in the journey.",
   [
    ("Excellent","9–10","Seamless design language across all pages. Every interaction pattern predictable and consistent. Terminology used identically throughout. Navigation behaviour consistent on every page. Standards match established mobile / web conventions."),
    ("Good","7–8","High consistency. Predictable patterns throughout. Minor deviations on 1–2 pages. Terminology mostly consistent. Navigation fundamentally reliable."),
    ("Moderate","5–6","Mostly consistent but with notable exceptions on specific pages. Some terminology variation. Navigation pattern deviates on certain sections."),
    ("Poor","3–4","Inconsistent experience across pages visited. UI patterns differ significantly between sections. Terminology inconsistent (same product called different names in different pages). Navigation behaves unpredictably."),
    ("Very Poor","0–2","Highly inconsistent. Different design systems appear to operate across different pages. Interaction patterns broken or contradictory. Terminology completely inconsistent. Each page feels like a different product."),
   ],
   ["Navigation header consistency across pages", "CTA styling consistency", "Product naming consistency", "Back button behaviour consistency", "Form styling consistency across multi-step flows"],
   ["'Personal Loan' on homepage called 'Consumer Loan' on product page", "Primary CTA is blue on some pages and orange on others", "Back button navigates in-page on some sections, full browser back on others", "Mobile bottom nav disappears on certain product pages"]),

  ("17", "User Control and Flexibility",
   "Degree to which the user can direct their own journey — go back, undo selections, edit previously entered information, save progress, and navigate freely without penalty.",
   [
    ("Excellent","9–10","Full user control throughout. Back navigation works correctly and preserves all form data. Users can edit any previous step in a multi-step flow. Progress auto-saved. Multiple entry and exit points to suit different user needs."),
    ("Good","7–8","Good user control. Back navigation functional. Most form data preserved. Users can generally edit previous selections. Minor friction in edge cases."),
    ("Moderate","5–6","Basic control. Back navigation sometimes resets sections. Form data partially preserved. Editing previous steps possible but clunky."),
    ("Poor","3–4","Limited user control. Back navigation frequently resets form progress. Previous steps cannot be edited. Linear-only flow with no flexibility. Mandatory fields that can't be skipped."),
    ("Very Poor","0–2","User has no meaningful control. Back button destroys all form progress. No ability to edit. Forced linear flow with no flexibility. User locked into a path they cannot exit without losing everything."),
   ],
   ["Back navigation form data preservation", "Edit capability for previously submitted form steps", "Progress save / resume functionality", "Ability to skip optional steps", "Browser back button consistency"],
   ["Complete form reset on browser back from step 3 of 5", "No ability to edit mobile number once submitted and OTP sent", "Application form cannot be saved and resumed", "Mandatory 'nominee' field blocking progression with no 'add later' option"]),

  ("17.1", "Data Privacy and Consent UX",
   "",
   [], [], []),  # placeholder, we'll handle dim 18 specially below
]

# Remove placeholder
CX_DIMENSIONS = [d for d in CX_DIMENSIONS if d[0] != "17.1"]

# Add dim 18 properly
CX_DIMENSIONS.append(
  ("18", "Data Privacy and Consent UX",
   "Transparency of data collection, quality of consent mechanisms, clarity of privacy communication, and ease of understanding what data is collected and why — from the perspective of this persona.",
   [
    ("Excellent","9–10","Data collection explained clearly and specifically at point of collection. Consent is genuinely informed — persona understands what they are agreeing to. Privacy policy accessible and in plain language for this persona. Data minimisation evident. Opt-out easy and visible. Trust built, not assumed."),
    ("Good","7–8","Privacy information accessible. Consent flows reasonable. Data use broadly explained. Minor gaps in contextual explanation. Opt-out available but not prominent."),
    ("Moderate","5–6","Basic privacy information present. Cookie consent present. General privacy policy accessible. Limited explanation of specific data use at collection points."),
    ("Poor","3–4","Privacy information buried or generic. Consent mechanisms nudge toward acceptance. Little to no contextual explanation of why specific data is collected. Opt-out difficult to find."),
    ("Very Poor","0–2","No meaningful privacy communication. Data collection implicit with no explanation. Consent absent or hidden. No opt-out visible. For this persona, the experience creates significant data-sharing anxiety with no resolution."),
   ],
   ["Privacy policy link visibility and placement", "Contextual data explanation at point of collection", "Cookie consent design (opt-in vs opt-out vs implied)", "PAN / Aadhaar / income collection justification copy", "Opt-out / data deletion option visibility"],
   ["'By proceeding you agree to our T&C' in 9pt font below CTA", "Aadhaar requested with no explanation of purpose or storage", "Cookie banner pre-checked with no opt-out option", "Privacy policy accessible only in English — no vernacular option for relevant personas"])
)

# Now draw all 18 dimensions
for num, name, what, rubric, criteria, friction in CX_DIMENSIONS:
    h2(doc, f"D.{num}  {name}")
    if what:
        bold_label(doc, "What this evaluates: ", what)
    doc.add_paragraph()
    if rubric:
        scoring_table(doc, rubric)
    if criteria:
        bold_label(doc, "Observable evidence criteria:", "")
        for c in criteria:
            body(doc, c, bullet=True, indent=360)
    if friction:
        bold_label(doc, "Common BFSI friction patterns:", "")
        for f in friction:
            body(doc, f, bullet=True, indent=360)
    doc.add_paragraph()

# ══════════════════════════════════════════════════════════════
# APPENDIX E — DESIGN EVALUATION FRAMEWORK
# ══════════════════════════════════════════════════════════════

doc.add_page_break()
h1(doc, "Appendix E", "Design Evaluation Framework")
body(doc,
    "This appendix defines the evaluation framework used by the Design Evaluator Agent. "
    "Design evaluation is persona-specific — the same design may score differently for a "
    "52-year-old semi-urban farmer persona versus a 28-year-old urban salaried professional. "
    "All assessments are grounded in screenshots captured during the actual journey.")
info_box(doc,
    "Scoring Scale:  9–10 = Excellent  |  7–8 = Good  |  5–6 = Moderate  |  3–4 = Poor  |  0–2 = Very Poor\n"
    "Design dimensions are evaluated on the pages actually visited during the journey. "
    "Design quality on unvisited pages is out of scope.")

DESIGN_DIMS = [
  ("E.1", "Visual Hierarchy and Layout Effectiveness",
   "How effectively the visual design of each page directs this persona's attention to the information and actions that matter most for their journey.",
   [
    ("Excellent","9–10","Visual hierarchy immediately directs this persona to their priority information (rates, CTA, trust signals) without conscious effort. Clear distinction between primary, secondary, and tertiary content levels. F-pattern or Z-pattern layout naturally guides the eye for this persona type."),
    ("Good","7–8","Strong visual hierarchy with minor competing elements. Priority information prominent. Secondary content clearly subordinate. Mostly supports natural eye movement for this persona."),
    ("Moderate","5–6","Some hierarchy present. Headings and body distinguished. Key information accessible but requiring deliberate scanning. 2–3 competing elements disrupt natural flow."),
    ("Poor","3–4","Weak hierarchy. Key information competes with secondary content. Dense layout. This persona would need to read in full rather than scan."),
    ("Very Poor","0–2","No meaningful visual hierarchy. Everything appears equal weight. Critical information invisible or buried. Layout impedes this persona's ability to extract information efficiently."),
   ]),

  ("E.2", "Colour Contrast and Text Readability",
   "Adequacy of colour contrast between text and background across all pages visited, evaluated against WCAG 2.1 standards and relative to this persona's visual acuity needs.",
   [
    ("Excellent","9–10","All body text meets WCAG AA (4.5:1 minimum contrast). Large text meets WCAG AA (3:1). Primary CTAs meet WCAG AA. No text uses colour as the only differentiator. Font size ≥16px for body content at 360px viewport."),
    ("Good","7–8","Primary content meets WCAG AA. Minor contrast failures only on tertiary elements (disclaimers, placeholder text). Body font size adequate."),
    ("Moderate","5–6","Primary body text adequate. Notable contrast failures on secondary elements. Some small text (12–14px) in important areas. Disclaimer text often below standard."),
    ("Poor","3–4","Multiple contrast failures on important content (product names, rates, CTAs). Body text at or below 14px in key sections. Colour palette choices create systematic readability issues."),
    ("Very Poor","0–2","Systematic contrast failures throughout. Key financial information unreadable for this persona. Text sizes critically small. Colour scheme creates readability barriers throughout the journey."),
   ]),

  ("E.3", "Touch Target Adequacy (Mobile)",
   "Sizing and spacing of interactive elements on the mobile viewport — ensuring all touchable elements are large enough for reliable activation without accidental triggers.",
   [
    ("Excellent","9–10","All interactive elements (buttons, links, form fields, navigation items, carousel controls) are ≥44×44px per WCAG 2.5.5. Adequate spacing between adjacent targets (≥8px) prevents mis-taps. One-thumb operation comfortable for primary actions."),
    ("Good","7–8","Primary interactive elements adequately sized. Secondary elements (inline links, small icons) may be slightly undersized but not critically so. Spacing mostly adequate."),
    ("Moderate","5–6","Primary CTA and navigation adequately sized. Several secondary interactive elements (filter chips, accordion headers, FAQ expand links) below minimum. Occasional mis-taps likely."),
    ("Poor","3–4","Multiple primary interactive elements below minimum size. Navigation items too small for reliable tapping. Carousel controls particularly problematic. Frequent mis-taps likely for this persona."),
    ("Very Poor","0–2","Systematic touch target failures. Primary CTAs, navigation, and form controls all below minimum size. Site is functionally difficult to operate on mobile for this persona. Desktop interaction model applied to mobile."),
   ]),

  ("E.4", "CTA Design and Prominence",
   "Visual clarity, prominence, label specificity, and placement of all call-to-action buttons across pages visited in the journey.",
   [
    ("Excellent","9–10","Primary CTA immediately visible above fold on every page where an action is expected. Button design distinguishes primary from secondary actions. Labels specific and action-oriented ('Apply for Gold Loan', 'Calculate My EMI'). Consistent CTA design language throughout journey."),
    ("Good","7–8","Primary CTAs clearly visible. Distinction between primary and secondary actions present. Labels mostly action-oriented. Minor positioning issues on 1–2 pages."),
    ("Moderate","5–6","CTAs present and findable but requiring some effort. Generic labels ('Apply', 'Proceed'). Primary vs secondary distinction inconsistent. Below-fold placement on some key pages."),
    ("Poor","3–4","CTAs difficult to find on some pages. Generic labels ('Submit', 'OK'). No clear visual distinction between primary and secondary actions. Inconsistent placement."),
    ("Very Poor","0–2","CTAs absent, hidden, or visually indistinguishable from body text. No consistent CTA design language. This persona cannot easily identify where to take action on any page."),
   ]),

  ("E.5", "Use of Imagery and Illustration",
   "Relevance, quality, cultural appropriateness, and informational value of images and illustrations on pages visited in the journey.",
   [
    ("Excellent","9–10","Images directly relevant to product and this persona's demographic context. Imagery features representative Indian customers in relatable scenarios. Images contribute to trust and understanding rather than occupying space. High quality and optimised for mobile loading."),
    ("Good","7–8","Most images relevant and appropriate. Broadly representative. Minor instances of generic stock photography. Images don't actively mislead or distract."),
    ("Moderate","5–6","Some relevant imagery. Mix of relevant and generic stock photos. Images occasionally dominate viewport at the expense of informational content. Representation partially appropriate."),
    ("Poor","3–4","Largely generic stock imagery with no relevance to this persona or Indian financial context. Images consume significant viewport without informational value. Some images slow page significantly."),
    ("Very Poor","0–2","Imagery irrelevant, culturally inappropriate, or actively misleading for this persona. Hero images consume most of mobile viewport. Critical informational content displaced by decorative imagery."),
   ]),

  ("E.6", "Whitespace and Visual Density",
   "Balance of content density and whitespace across pages, assessed relative to this persona's ability to process information at the observed density level.",
   [
    ("Excellent","9–10","Content density perfectly calibrated for this persona's cognitive processing capacity. Adequate whitespace between sections. Information grouped logically. No content claustrophobia on mobile. Each screen presents a manageable amount of information."),
    ("Good","7–8","Good use of whitespace. Content density appropriate for most screens. Minor density spikes on some pages. Generally comfortable reading experience."),
    ("Moderate","5–6","Adequate whitespace on primary pages. Some screens notably dense with information. Minor claustrophobia on mobile. This persona would need to slow down on denser pages."),
    ("Poor","3–4","Consistently high information density. Whitespace sparse throughout. This persona would feel overwhelmed on multiple pages. Wall-of-text presentation in several sections."),
    ("Very Poor","0–2","Extreme information density throughout. No meaningful whitespace. Every page presents an overwhelming amount of content. This persona would abandon rather than process the content volume."),
   ]),

  ("E.7", "Design Consistency Across Pages",
   "Consistency of visual design language, colour system, typography, spacing, and component design across all pages visited in the journey.",
   [
    ("Excellent","9–10","Seamless visual design language across all pages. Identical colour palette, type scale, spacing system, and component library throughout. Every page feels part of the same product. No jarring transitions between sections."),
    ("Good","7–8","High design consistency. Minor variations on 1–2 pages. Core visual language consistent. No significant jarring transitions."),
    ("Moderate","5–6","Mostly consistent. Notable design variations on specific sections (e.g. application funnel looks different from marketing pages). Core branding consistent but detail-level deviations."),
    ("Poor","3–4","Significant design inconsistencies. Multiple distinct visual styles operating across the journey. Colour, type, and component inconsistencies create a fragmented experience."),
    ("Very Poor","0–2","Highly inconsistent design. Different sections appear to belong to different products or companies. Brand colours, typography, and component styles conflict across pages. This persona would question if they are on the same site."),
   ]),

  ("E.8", "Form Field Visual Design",
   "Visual design quality of all form elements encountered in the journey — label placement, input field sizing, focus state design, error state design, and overall form visual coherence.",
   [
    ("Excellent","9–10","All form labels persistently visible above or beside input fields (not placeholder-only). Input fields sized for comfortable touch entry. Clear focus states indicating active field. Distinct error states (colour + icon + text). Helper text visually distinct from labels. Logical visual grouping of related fields."),
    ("Good","7–8","Labels persistent and clear. Most fields adequately sized. Focus states present. Error states present with most required elements. Minor visual inconsistencies."),
    ("Moderate","5–6","Labels adequate but some use placeholder-only approach. Fields mostly usable. Focus states basic. Error states rely primarily on colour. Some visual inconsistency."),
    ("Poor","3–4","Multiple placeholder-only labels (disappear on focus). Some fields too small for comfortable mobile entry. Focus states absent or barely visible. Error states colour-only. Poor visual grouping."),
    ("Very Poor","0–2","Systematic form design failures. Placeholder-only labels throughout. No focus states. Error states absent or indistinguishable. Form fields too small for mobile. This persona would struggle with basic form interaction throughout the application flow."),
   ]),
]

for num, name, what, rubric in DESIGN_DIMS:
    h2(doc, f"{num}  {name}")
    bold_label(doc, "What this evaluates: ", what)
    doc.add_paragraph()
    scoring_table(doc, rubric)
    doc.add_paragraph()

# ══════════════════════════════════════════════════════════════
# APPENDIX F — CONTENT EVALUATION FRAMEWORK
# ══════════════════════════════════════════════════════════════

doc.add_page_break()
h1(doc, "Appendix F", "Content Evaluation Framework")
body(doc,
    "This appendix defines the evaluation framework used by the Content Analyzer Agent. "
    "Content quality is always assessed relative to this specific persona — the same content "
    "may be appropriately clear for a financially literate urban professional but "
    "completely inaccessible for a first-time rural borrower with limited English literacy.")
info_box(doc,
    "Scoring Scale:  9–10 = Excellent  |  7–8 = Good  |  5–6 = Moderate  |  3–4 = Poor  |  0–2 = Very Poor\n"
    "Content on pages not visited during the journey is out of scope. "
    "All assessments are grounded in page text extracted from the journey log.")

CONTENT_DIMS = [
  ("F.1", "Financial Clarity and Rate Transparency",
   "Whether key financial information — interest rates, fees, EMI structure, total cost of credit, prepayment charges — is communicated clearly, accurately, and prominently before the user is asked to share personal data.",
   [
    ("Excellent","9–10","All key financial figures prominently displayed on product page before any login gate. Rates shown as actual APR or monthly rate alongside 'starting from' figures. All fees itemised with amounts or ranges. EMI calculation clearly linked to tenure/amount. No hidden charges introduced post-login. Language fully accessible for this persona's literacy level."),
    ("Good","7–8","Most financial information visible pre-gate. Rates and main fees present. Minor gaps (e.g. processing fee range not specified). Language mostly appropriate."),
    ("Moderate","5–6","Basic financial information present but not complete. 'Interest rate from X%' without clarity on what determines the actual rate. Fees mentioned but not detailed. EMI calculator present but formula opaque."),
    ("Poor","3–4","Financial information sparse, vague, or requires significant exploration to find. Rate ranges very wide without context. Fees unspecified. Key information available only post-login."),
    ("Very Poor","0–2","Financial information absent from public pages. Rate only shown as 'Competitive rates — apply to know more'. All financial details gated behind login. Persona has no basis for pre-decision."),
   ]),

  ("F.2", "Pre-Gate Content Sufficiency",
   "Whether the content available before any login wall or data-collection gate provides this persona with enough information to make a meaningful, informed decision about whether to proceed.",
   [
    ("Excellent","9–10","Persona can answer all of these questions from public content alone: What is the product? What are the key terms (rate, fee, tenure)? Am I likely to be eligible? What happens if I apply? What happens to my data? This persona does not need to log in to understand whether the product is right for them."),
    ("Good","7–8","Most key questions answerable from public content. Minor gaps (e.g. eligibility criteria partially described). Persona has sufficient context to make a reasonable pre-decision."),
    ("Moderate","5–6","Basic product description present. Rate range available. Eligibility criteria vague or partially visible. Persona can form a rough understanding but with notable uncertainty."),
    ("Poor","3–4","Minimal pre-gate information. Product described in marketing language without substantive details. Persona is asked to 'apply to know more' on most key questions."),
    ("Very Poor","0–2","Essentially no pre-gate content beyond product name and marketing headline. All substantive information hidden behind login. Persona is asked to commit personal data before understanding the product."),
   ]),

  ("F.3", "Jargon Assessment and Language Appropriateness",
   "Whether the language used in all content on visited pages is accessible and appropriate for this persona's financial literacy level, education background, and language preference.",
   [
    ("Excellent","9–10","Zero unexplained financial jargon for this persona's literacy level. Technical terms are either avoided or immediately explained in plain language with examples. Reading level calibrated to this persona. Vernacular language available if relevant to this persona's context."),
    ("Good","7–8","Minimal unexplained jargon. Most technical terms explained or contextualised. Language broadly accessible. 1–2 jargon instances present but not critical to core journey."),
    ("Moderate","5–6","Some jargon present without explanation. Important terms (e.g. 'DPD', 'LTV', 'CIBIL' for low-literacy persona) used as-is. Language partially accessible but would create confusion at 2–3 points."),
    ("Poor","3–4","Significant jargon throughout. Multiple unexplained financial, regulatory, or legal terms. Language clearly written for an audience more financially sophisticated than this persona. Persona would be confused on most pages."),
    ("Very Poor","0–2","Pervasive jargon. Content appears written entirely for financially expert audience. This persona cannot extract meaningful information from most pages encountered. Core product benefits unclear due to language complexity."),
   ]),

  ("F.4", "Regulatory and Legal Disclosure Completeness",
   "Whether all required regulatory and legal disclosures are present and accessible on pages visited, including RBI-mandated disclosures, interest rate disclosure norms, KFS requirements, and T&C accessibility.",
   [
    ("Excellent","9–10","RBI registration number clearly visible. Interest rate disclosure compliant with RBI Fair Practices Code. Key Fact Statement (KFS) or equivalent accessible from product page. T&C accessible in plain language. Processing fee and prepayment charge disclosures present. MITC (Most Important Terms and Conditions) accessible."),
    ("Good","7–8","Most required disclosures present. RBI registration visible. Rate disclosure adequate. T&C accessible. Minor gaps in non-critical disclosures."),
    ("Moderate","5–6","Basic regulatory disclosures present. RBI registration in footer. T&C linked from application form. KFS absent but key terms visible. Interest rate methodology partially explained."),
    ("Poor","3–4","Several required disclosures absent or inaccessible. RBI registration not visible from main pages. T&C behind separate link only at application. No rate methodology disclosure."),
    ("Very Poor","0–2","Critical regulatory disclosures absent from all pages visited. No RBI registration visible. No T&C accessible pre-application. No rate disclosure methodology. Site would fail basic RBI Fair Practices Code audit on pages visited."),
   ]),

  ("F.5", "Micro-copy Quality and Effectiveness",
   "Quality of all small text elements encountered in the journey — form labels, button text, error messages, placeholder text, helper text, tooltips, confirmation messages, loading messages — assessed for clarity, specificity, and appropriateness for this persona.",
   [
    ("Excellent","9–10","Every micro-copy element is specific, action-oriented, and calibrated for this persona. Form field labels persistently visible and descriptive. Placeholder text provides contextual examples relevant to Indian users (e.g. '₹ e.g. 35,000'). Error messages identify the specific issue and the specific remedy. Confirmation messages describe what just happened and what happens next."),
    ("Good","7–8","Most micro-copy clear and helpful. Button labels mostly specific. Good placeholder examples on most fields. Error messages mostly actionable. Minor generic labels remaining."),
    ("Moderate","5–6","Adequate micro-copy. Some generic elements ('Enter value', 'Submit', 'Error'). Placeholder text sometimes absent. Errors sometimes generic. Core navigation labelled adequately."),
    ("Poor","3–4","Predominantly generic micro-copy. Most buttons labelled generically. No contextual placeholder examples. Error messages generic ('Invalid input', 'Try again'). Technical language in some elements."),
    ("Very Poor","0–2","Micro-copy absent, generic, or misleading throughout. Button labels provide no information about outcome. Error codes visible. Technical system messages presented to user. This persona cannot reliably interpret the interface."),
   ]),

  ("F.6", "Tone, Voice, and Brand Appropriateness",
   "Whether the content tone and brand voice across all visited pages is appropriately calibrated for this persona's demographic, emotional state during the journey, and the product category.",
   [
    ("Excellent","9–10","Tone consistently appropriate for this persona's demographic, literacy level, and emotional state. Financial product language appropriately serious without being intimidating. Warmth and trust calibrated to the persona. Brand voice consistent across all pages. Urgency messaging absent or appropriate."),
    ("Good","7–8","Tone largely appropriate. Minor mismatches on 1–2 pages. Brand voice broadly consistent. No actively inappropriate tone."),
    ("Moderate","5–6","Tone mostly appropriate with notable exceptions. Some sections feel too formal or too casual for this persona. Brand voice inconsistent between marketing and product pages."),
    ("Poor","3–4","Tone frequently mismatched to this persona. Marketing sections over-enthusiastic. Legal/product sections overly formal. Urgency tactics that create pressure rather than confidence for this persona."),
    ("Very Poor","0–2","Tone actively counterproductive for this persona throughout. Aggressive sales language creating anxiety in risk-averse persona. Overly simplified language feeling patronising for sophisticated persona. Tone inconsistent across every section of the journey."),
   ]),

  ("F.7", "Missing Content Assessment",
   "Identification of specific informational gaps — content this persona needed at specific journey steps that was absent from all pages visited.",
   [
    ("Excellent","9–10","No critical content gaps identified for this persona. All information needed to complete the journey goal is present and accessible on pages visited. Proactive content addresses questions this persona would have before they need to ask."),
    ("Good","7–8","Minor content gaps only. Most information needed is present. 1–2 secondary information needs unaddressed (e.g. loan insurance optionality not explained)."),
    ("Moderate","5–6","Several content gaps affecting journey quality. Some information needed by this persona at decision points is absent from pages visited, leading to uncertainty. Persona makes decisions with incomplete information."),
    ("Poor","3–4","Significant content gaps. Multiple decision points where this persona lacks critical information. Several questions that would prevent a real customer from proceeding are unaddressed."),
    ("Very Poor","0–2","Major content gaps throughout. This persona cannot answer the most basic questions about product suitability, eligibility, costs, or process from any page visited. Site fails its fundamental informational purpose for this persona."),
   ]),
]

for num, name, what, rubric in CONTENT_DIMS:
    h2(doc, f"{num}  {name}")
    bold_label(doc, "What this evaluates: ", what)
    doc.add_paragraph()
    scoring_table(doc, rubric)
    doc.add_paragraph()

# ══════════════════════════════════════════════════════════════
# APPENDIX G — ACCESSIBILITY EVALUATION FRAMEWORK
# ══════════════════════════════════════════════════════════════

doc.add_page_break()
h1(doc, "Appendix G", "Accessibility Evaluation Framework")
body(doc,
    "This appendix defines the evaluation framework used by the Accessibility Auditor Agent. "
    "Accessibility is evaluated through visual observation of screenshots and page state text. "
    "This is a qualitative assessment — not a programmatic WCAG compliance audit. "
    "Findings are indicative and should be verified with formal tooling (axe, Lighthouse) "
    "before remediation.")
info_box(doc,
    "Scoring Scale:  9–10 = Excellent  |  7–8 = Good  |  5–6 = Moderate  |  3–4 = Poor  |  0–2 = Very Poor\n"
    "WCAG Reference: Web Content Accessibility Guidelines 2.1 (Level AA as minimum target).\n"
    "All observations are persona-contextualised — assessed relative to this persona's "
    "specific accessibility needs and literacy level.")

ACC_DIMS = [
  ("G.1", "Text Size and Readability",
   "Whether text sizes across all visited pages are adequate for this persona to comfortably read all informational content without assistive tools.",
   [
    ("Excellent","9–10","All body text ≥16px at viewport width. Heading hierarchy clear and appropriately sized. Line height ≥1.5 for body content. Text scales appropriately with browser zoom. No text rendered as non-scalable image. Maximum line length ≤80 characters for comfortable reading."),
    ("Good","7–8","Most body text ≥16px. Headings clearly sized. Minor instances of 14px text in secondary areas. Generally comfortable for this persona."),
    ("Moderate","5–6","Body text 14–16px with some 12–13px in secondary areas. Comfortable for most content but some fine print in important areas. This persona may need to lean in."),
    ("Poor","3–4","Multiple areas of important text at 12–14px. Disclaimer text and fee information at 10–12px. This persona would strain to read key financial information."),
    ("Very Poor","0–2","Systematic text size inadequacy. Key information (rates, fees, T&C) rendered at 8–10px. Text in images not scalable. This persona cannot comfortably read critical content without assistive magnification."),
   ]),

  ("G.2", "Colour Contrast (WCAG Standards)",
   "Adequacy of colour contrast between all foreground text and background colours across pages visited, evaluated against WCAG 2.1 AA standards.",
   [
    ("Excellent","9–10","All normal text achieves minimum 4.5:1 contrast ratio (WCAG AA). All large text (≥18pt or ≥14pt bold) achieves minimum 3:1. All UI components and graphical objects achieve 3:1. No information conveyed by colour alone. Contrast maintained in both light and any dark-mode rendering."),
    ("Good","7–8","Primary body text meets WCAG AA. Large text meets WCAG AA. Minor failures on tertiary elements (footer text, placeholder text). No failures on critical financial information."),
    ("Moderate","5–6","Primary content contrast adequate. Notable failures on secondary elements (grey-on-white subheadings, light placeholder text). Disclaimer text commonly failing. No failures on primary CTAs."),
    ("Poor","3–4","Multiple contrast failures on secondary content. Some failures on important content (product descriptions, fee information). CTA contrast potentially marginal. This persona would notice readability difficulty."),
    ("Very Poor","0–2","Systematic contrast failures. Key financial information (rates, fees, eligibility) fails WCAG AA. CTA buttons fail contrast requirements. Colour scheme creates fundamental readability barrier. This persona cannot reliably read significant portions of page content."),
   ]),

  ("G.3", "Touch Target Sizing (WCAG 2.5.5)",
   "Whether all interactive elements on the mobile viewport meet minimum touch target size requirements to ensure reliable activation for users with typical manual dexterity.",
   [
    ("Excellent","9–10","All interactive elements ≥44×44 CSS pixels (WCAG 2.5.5 AAA). Minimum 8px spacing between adjacent targets. Primary navigation items, CTAs, form fields, and carousel controls all meet or exceed standard. Comfortable for one-thumb operation."),
    ("Good","7–8","Primary interactive elements meet ≥44×44px. Secondary elements (inline text links, icon buttons) mostly meet standard with minor exceptions. Spacing between targets generally adequate."),
    ("Moderate","5–6","CTAs and navigation meet standard. Multiple secondary interactive elements (accordion headers, filter chips, FAQ expanders) at 32–40px. Occasional mis-taps likely."),
    ("Poor","3–4","Multiple primary and secondary interactive elements below minimum. Navigation items at 28–36px. Close buttons on modals at 20–28px. Frequent mis-taps likely for this persona."),
    ("Very Poor","0–2","Systematic touch target failures. Primary CTAs below minimum size. Navigation items too small for reliable mobile operation. Close buttons on overlays practically untappable. This persona would frequently mis-tap, closing important content or triggering unintended actions."),
   ]),

  ("G.4", "Language Simplicity and Cognitive Accessibility",
   "Whether the reading level, sentence complexity, and use of plain language across all visited pages matches this persona's cognitive and language accessibility needs.",
   [
    ("Excellent","9–10","Content consistently at appropriate reading level for this persona. Short sentences and simple vocabulary throughout. Complex concepts broken into steps. Indian English used appropriately. Vernacular elements where culturally relevant. No acronyms or technical terms without explanation."),
    ("Good","7–8","Content largely accessible. Minor complexity spikes in legal/regulatory sections. Most concepts well-explained. 1–2 unexplained acronyms."),
    ("Moderate","5–6","Content partially accessible. Financial/regulatory sections notably more complex than product sections. Several unexplained terms. This persona would struggle with approximately 25–30% of content."),
    ("Poor","3–4","Content frequently inaccessible for this persona. Heavy use of financial and legal terminology. Long, complex sentences. This persona would struggle with 50%+ of content encountered."),
    ("Very Poor","0–2","Content systematically inaccessible for this persona's literacy level. Written for expert audience. This persona cannot extract meaningful information from the majority of content on pages visited."),
   ]),

  ("G.5", "Form Field Accessibility",
   "Accessibility of all form interactions in the journey — label association, error identification, required field marking, focus management, and keyboard interaction support.",
   [
    ("Excellent","9–10","All form fields have persistent, programmatically associated labels (not placeholder-only). Required fields clearly marked with standard indicator. Error messages identify the specific field and describe the specific remedy in plain text (not just colour change). Focus management logical after error. Keyboard tab order correct."),
    ("Good","7–8","Most fields have clear persistent labels. Required field marking consistent. Error messages mostly field-specific. Minor focus management issues. Keyboard navigation broadly functional."),
    ("Moderate","5–6","Labels present but some placeholder-only. Required field marking present but inconsistent. Error states include text but sometimes generic. Focus sometimes lost after form submission."),
    ("Poor","3–4","Several fields with placeholder-only labels. Required fields inconsistently marked. Error states primarily colour-only. Focus management unreliable. Keyboard navigation problematic."),
    ("Very Poor","0–2","Systematic form accessibility failures. Placeholder-only labels throughout. No required field marking. Error states colour-only with no text description. Focus lost after errors. Keyboard users face significant barriers throughout the application flow."),
   ]),

  ("G.6", "Alternative Text for Informational Images",
   "Whether informational images, product diagrams, rate tables rendered as images, and trust badges have meaningful alternative text for screen reader users and situations where images fail to load.",
   [
    ("Excellent","9–10","All informational images have descriptive alt text that conveys the same information as the image. Rate comparison tables rendered as HTML tables, not images. Decorative images have empty alt attributes. Trust badges have alt text conveying their purpose. All product diagrams described adequately in alt text."),
    ("Good","7–8","Most informational images have meaningful alt text. Minor gaps in secondary imagery. Rate tables rendered as HTML in most cases. Trust badges alt text present."),
    ("Moderate","5–6","Primary product images have basic alt text. Some informational images lack alt text. Rate information occasionally in image format without alt text. Mixed approach to decorative image alt treatment."),
    ("Poor","3–4","Alt text largely absent or file-name based ('image001.jpg'). Important financial information in images without alt text. Trust badges without alt text. Persona relying on images for information would face barriers."),
    ("Very Poor","0–2","No meaningful alt text on any image. Key financial information (rate tables, eligibility infographics, process diagrams) rendered as images with no text alternative. Complete barrier for screen reader users or slow connections."),
   ]),

  ("G.7", "Navigation and Structural Accessibility",
   "Observable accessibility of the navigation structure — skip links, heading hierarchy, landmark structure, and keyboard navigation support — as evidenced from page state and screenshot analysis.",
   [
    ("Excellent","9–10","Logical heading hierarchy (H1 → H2 → H3) observable from page structure. Navigation landmarks structurally consistent. Focus visible on interactive elements. Tab order follows visual reading order. Mobile navigation keyboard-accessible equivalent. No keyboard traps observable."),
    ("Good","7–8","Heading hierarchy mostly logical. Navigation landmarks present. Focus visible on primary elements. Minor structural deviations. Keyboard navigation broadly functional on primary paths."),
    ("Moderate","5–6","Heading hierarchy partially consistent. Some structural landmarks present. Focus visibility variable. Tab order mostly correct. Mobile navigation accessible with effort."),
    ("Poor","3–4","Heading hierarchy inconsistent. Limited structural landmarks. Focus visibility poor. Tab order deviates from visual order. Mobile navigation may be keyboard-inaccessible."),
    ("Very Poor","0–2","No observable heading hierarchy. No structural landmarks. Focus invisible throughout. Tab order completely disconnected from visual layout. Navigation functionally inaccessible for keyboard-only or assistive-technology users."),
   ]),
]

for num, name, what, rubric in ACC_DIMS:
    h2(doc, f"{num}  {name}")
    bold_label(doc, "What this evaluates: ", what)
    doc.add_paragraph()
    scoring_table(doc, rubric)
    doc.add_paragraph()

# Add closing note to Appendix G
info_box(doc,
    "Important Scope Note: All accessibility findings from this automated agent review are based on "
    "visual observation and page state text analysis. They are indicative findings, not a certified "
    "WCAG compliance audit. Before reporting accessibility compliance or non-compliance to regulatory "
    "bodies, findings must be verified using formal automated tooling (axe-core, WAVE, Lighthouse) "
    "and manual screen reader testing (NVDA/TalkBack on actual devices).",
    fill="FFF3E0")

# ── Update Appendix C version table ──
replace_in_doc(doc,
    "Initial PRD submission to PMO team. All evaluation and scoring frameworks included inline. Authored in full compliance with BFL Enterprise Agentic Requirement SOP V1.0.",
    "v1.0 — May 2026 — Initial PRD. Appendices D–G added: full CX, Design, Content and Accessibility evaluation and scoring frameworks. All external BRD cross-references resolved. Authored in compliance with BFL Enterprise Agentic Requirement SOP V1.0.")

# ─────────────────────────────────────────────────────────────
# SAVE
# ─────────────────────────────────────────────────────────────
doc.save(DOC_PATH)
print(f"\nSaved: {DOC_PATH}")
