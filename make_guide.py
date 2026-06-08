"""Generate CX Audit Agent â€” Full Guide.docx"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# â”€â”€ Page margins â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€
for section in doc.sections:
    section.top_margin    = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin   = Inches(1.2)
    section.right_margin  = Inches(1.2)

# â”€â”€ Style helpers â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€
def h1(text):
    p = doc.add_heading(text, level=1)
    p.runs[0].font.color.rgb = RGBColor(0x1A, 0x56, 0xDB)   # blue
    doc.add_paragraph()

def h2(text):
    p = doc.add_heading(text, level=2)
    p.runs[0].font.color.rgb = RGBColor(0x1E, 0x40, 0xAF)

def h3(text):
    doc.add_heading(text, level=3)

def body(text):
    doc.add_paragraph(text)

def bullet(text, level=0):
    p = doc.add_paragraph(text, style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.3 * (level + 1))

def numbered(text, level=0):
    p = doc.add_paragraph(text, style="List Number")
    p.paragraph_format.left_indent = Inches(0.3 * (level + 1))

def code(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Courier New"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x1F, 0x29, 0x37)
    p.paragraph_format.left_indent   = Inches(0.4)
    p.paragraph_format.space_before  = Pt(2)
    p.paragraph_format.space_after   = Pt(2)
    # Light grey shading
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  "F3F4F6")
    p._p.get_or_add_pPr().append(shd)

def divider():
    doc.add_paragraph("â”€" * 80)

def spacer():
    doc.add_paragraph()

# â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•
# COVER
# â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run("CX Audit Agent")
run.bold = True
run.font.size = Pt(28)
run.font.color.rgb = RGBColor(0x1A, 0x56, 0xDB)

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
run2 = sub.add_run("Complete Product Guide â€” What It Is, What It Does, How to Run It")
run2.font.size = Pt(13)
run2.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)

spacer()
divider()
spacer()

# â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•
# SECTION 1 â€” WHAT HAS BEEN BUILT
# â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•
h1("1.  What Has Been Built")

body(
    "The CX Audit Agent is an AI-powered customer experience testing system. "
    "It acts like a real user visiting a website or mobile app, navigates through it, "
    "records everything that works well or causes friction, and produces detailed audit reports "
    "with scores, findings, and recommendations â€” all without any manual testing."
)
spacer()
body(
    "It was built specifically to audit the Bajaj Finserv digital properties but the architecture "
    "is general enough to point at any website or Android app."
)
spacer()

h2("Core Components Built")

h3("1.1  Web Audit Agent  (main.py)")
body(
    "Uses a Playwright-controlled browser (Chromium) and OpenAI to visit a target website. "
    "OpenAI sees the page content, decides what to click or type next, "
    "executes that action through the browser, and records the experience at every step. "
    "The agent runs autonomously until the persona's goal is achieved or the step limit is reached."
)

h3("1.2  Mobile / App Audit Agent  (mobile_main.py)")
body(
    "Uses Appium + Android emulator and OpenAI to test the mobile experience. "
    "OpenAI receives a screenshot of the device screen, decides the next action (tap, swipe, type, back), "
    "and executes it on the real emulator. This covers both the Bajaj Finserv mobile website running in "
    "Chrome on Android and can be extended to a native Android app."
)

h3("1.3  Persona System  (personas/bajaj_personas.md)")
body(
    "The agent does not test as a generic user. It tests as specific, realistic personas â€” "
    "each with a name, age, occupation, financial literacy level, goal, device type, and behavioural constraints. "
    "Multiple personas run in sequence so you get results from multiple user perspectives in one run."
)
body("Personas currently defined for Bajaj Finserv:")
bullet("Arjun Mehra â€” 32, salaried software engineer, researching gold loan rates")
bullet("Priya Nair â€” 28, freelance designer, checking personal loan eligibility")
bullet("Ramesh Patel â€” 55, small business owner, exploring business loan options")
bullet("Sunita Sharma â€” 45, homemaker, looking for fixed deposit information")
bullet("Vikram Singh â€” 38, government employee, comparing insurance products")

h3("1.4  CX Evaluator  (evaluation/cx_evaluator.py)")
body(
    "After each journey completes, OpenAI acts as a CX expert and scores the experience across "
    "seven dimensions on a scale of 0 to 10:"
)
bullet("Navigation & Findability")
bullet("Content Clarity")
bullet("Trust & Credibility")
bullet("Load & Performance")
bullet("Mobile Responsiveness")
bullet("Conversion Ease")
bullet("Emotional Tone")
body("It also extracts friction points (high / medium / low severity), positive moments, and prioritised recommendations (P1 / P2 / P3).")

h3("1.5  Report Generator  (reporting/report_generator.py)")
body("Produces three output formats after every run:")
bullet("Per-persona Markdown report  â€”  reports/{persona-slug}/report.md")
bullet("Per-persona JSON log  â€”  reports/{persona-slug}/journey_log.json")
bullet("Master summary report  â€”  reports/master_report.md  (aggregates all personas)")
body("Mobile audit reports go to reports/mobile/ so they never overwrite web reports.")

h3("1.6  Dashboard  (dashboard/index.html)")
body(
    "A fully interactive web dashboard that displays all audit results visually. "
    "It reads the JSON logs and renders:"
)
bullet("Run history list with filter by audit type (web / app)")
bullet("CX scores per dimension with colour-coded bars")
bullet("Persona details tab â€” who was tested and their profile")
bullet("Friction points sorted by severity")
bullet("Recommendations sorted by priority (P1 first)")
bullet("Step-by-step execution log with success/fail status per step")
bullet("Screenshots tab showing captured screenshots per step")

h3("1.7  Web Server  (server.py)")
body(
    "A Flask API server that sits between the dashboard and the audit engine. "
    "It exposes REST endpoints so the dashboard can:"
)
bullet("List all past runs")
bullet("Fetch report data for any run")
bullet("Trigger a new web or app audit run from the browser")
bullet("Stream real-time status back to the UI while a run is in progress")

spacer()
divider()
spacer()

# â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•
# SECTION 2 â€” WHAT IT CAN DO
# â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•
h1("2.  What It Can Do")

h2("2.1  Audit Modes")

h3("Web Audit â€” Logged Out")
body(
    "Visits the target website as an unauthenticated user. Tests the public-facing experience: "
    "navigation, product discovery, information clarity, and lead generation flows."
)

h3("Web Audit â€” Logged In")
body(
    "Visits the website with a test account already authenticated. Tests the post-login experience: "
    "dashboard usability, account management, and transactional flows."
)

h3("Mobile / App Audit â€” Chrome on Android")
body(
    "Opens the mobile website in Chrome on an Android emulator. Tests how the responsive web "
    "experience feels on a real Android device, including tap accuracy, scroll behaviour, and "
    "mobile-specific navigation patterns."
)

h3("Mobile / App Audit â€” Native Android App  (extendable)")
body(
    "By setting the app package name in config/app_config.yaml, the same agent can audit "
    "a native Android app installed on the emulator, using the exact same screenshot-based "
    "OpenAI decision loop."
)

h2("2.2  Multi-Persona Testing")
body(
    "One command runs the audit for all defined personas sequentially. Each persona explores "
    "the site differently based on their goal and profile. The master report shows side-by-side "
    "comparisons across all personas."
)

h2("2.3  Debug Mode")
body(
    "Run with --debug flag to test only the first persona with a maximum of 10 steps. "
    "Useful for verifying configuration before a full run."
)

h2("2.4  Autonomous Navigation")
body(
    "The agent decides every action itself â€” what to click, what to type, when to scroll, "
    "when to declare success or failure. No scripts or selectors need to be maintained."
)

h2("2.5  Observation Recording")
body(
    "As it navigates, the agent records friction observations (slow load, confusing label, "
    "broken link, hidden CTA) and delight observations (clear pricing, fast response, "
    "reassuring copy) in real time."
)

h2("2.6  Screenshot Capture")
body(
    "Screenshots are saved at configurable points during the journey â€” on every step (verbose mode) "
    "or only on failure. Mobile audit always saves a screenshot per step. All screenshots are "
    "accessible from the dashboard Screenshots tab."
)

h2("2.7  Retry and Resilience")
body(
    "If an action fails (element not found, page load timeout), the agent retries automatically "
    "up to the configured retry count before moving on. Failures are logged and included in the report."
)

h2("2.8  Dashboard Filtering and History")
body(
    "The dashboard keeps a history of every run. You can filter by audit type (Web vs App), "
    "view individual persona reports, compare scores across runs, and drill down into any step."
)

spacer()
divider()
spacer()

# â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•
# SECTION 3 â€” PREREQUISITES
# â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•
h1("3.  Prerequisites â€” Install These Once")

body("Before running anything, ensure the following are installed on your machine.")

h2("3.1  Python and Packages")

numbered("Install Python 3.11 or higher from python.org")
numbered("Open a terminal in the CX Audit Agent folder and install all Python packages:")
code("pip install -r requirements.txt")
numbered("Install Playwright browsers:")
code("playwright install chromium")
numbered("Install python-docx if not present (only needed to regenerate this document):")
code("pip install python-docx")

h2("3.2  Android Emulator  (for mobile audit only)")

numbered("Install Android Studio from developer.android.com/studio")
numbered("Open Android Studio, go to Device Manager, create a virtual device:")
bullet("Select Pixel 6 or Pixel 6a")
bullet("Select API Level 33 or higher")
bullet("Finish and download the system image when prompted")
numbered("Note the Android SDK path â€” typically:")
code("C:\\Users\\YourName\\AppData\\Local\\Android\\Sdk")

h2("3.3  Appium  (for mobile audit only)")
numbered("Install Node.js from nodejs.org")
numbered("Install Appium globally:")
code("npm install -g appium")
numbered("Install the UiAutomator2 driver:")
code("appium driver install uiautomator2")

h2("3.4  API Key")
numbered("Create a file called  .env  in the CX Audit Agent folder")
numbered("Add your OpenAI API key:")
code("OPENAI_API_KEY=sk-your-openai-key-here")
body("You can get an API key from platform.openai.com/api-keys")

spacer()
divider()
spacer()

# â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•
# SECTION 4 â€” HOW TO RUN: WEB AUDIT
# â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•
h1("4.  How to Run â€” Web Audit")

body("The web audit requires only Python, Playwright, and an OpenAI API key. No emulator needed.")

h2("Step 1 â€” Open a terminal in the project folder")
body("Open Command Prompt or PowerShell, then navigate to the folder:")
code("cd \"C:\\Users\\Hiya Bhandari\\Desktop\\Folders\\CX Audit Agent\"")

h2("Step 2 â€” Run the web audit")

h3("Option A â€” All personas, full run")
code("python -X utf8 main.py")

h3("Option B â€” Debug mode  (1 persona, 10 steps, fastest)")
code("python -X utf8 main.py --debug")

h3("Option C â€” Headless mode  (no visible browser window)")
code("python -X utf8 main.py --no-headed")

h3("Option D â€” Logged-in audit")
code("python -X utf8 main.py --auth-mode logged_in")

h3("Option E â€” Single persona by name")
code("python -X utf8 main.py --persona \"Arjun Mehra\"")

h2("Step 3 â€” Watch the terminal")
body(
    "The terminal shows live progress: which persona is running, each step the agent takes, "
    "actions executed, and any failures. A summary table appears at the end with CX scores."
)

h2("Step 4 â€” Find your reports")
body("After the run, reports are written to:")
bullet("reports/{persona-slug}/report.md  â€”  full Markdown report per persona")
bullet("reports/{persona-slug}/journey_log.json  â€”  raw JSON data used by the dashboard")
bullet("reports/master_report.md  â€”  combined summary of all personas")

spacer()
divider()
spacer()

# â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•
# SECTION 5 â€” HOW TO RUN: MOBILE AUDIT
# â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•
h1("5.  How to Run â€” Mobile / App Audit")

body("The mobile audit requires the Android emulator and Appium in addition to Python.")

h2("Step 1 â€” Start the Android emulator")
body("Open Android Studio. In Device Manager, click the Play button next to your Pixel 6 device.")
body("Or start it from terminal using the emulator command:")
code("emulator -avd Pixel_6 -no-snapshot-load")
body("Wait until the Android home screen is fully loaded before proceeding.")

h2("Step 2 â€” Confirm the emulator is detected")
body("Run this command to verify Android can see the device:")
code("C:\\Users\\Hiya Bhandari\\AppData\\Local\\Android\\Sdk\\platform-tools\\adb.exe devices")
body("You should see:  emulator-5554   device")

h2("Step 3 â€” Start Appium with ANDROID_HOME set")
body(
    "This is the most important step. Appium needs to know where the Android SDK is. "
    "Open a new terminal window and run exactly this command:"
)
code(
    "set ANDROID_HOME=C:\\Users\\Hiya Bhandari\\AppData\\Local\\Android\\Sdk\n"
    "set ANDROID_SDK_ROOT=C:\\Users\\Hiya Bhandari\\AppData\\Local\\Android\\Sdk\n"
    "appium --port 4723 --allow-insecure uiautomator2:chromedriver_autodownload"
)
body(
    "Leave this terminal open. Appium must stay running for the entire duration of the mobile audit. "
    "You should see:  Appium REST http interface listener started on http://0.0.0.0:4723"
)

h2("Step 4 â€” Run the mobile audit  (in a second terminal)")
body("Open a second terminal window in the project folder and run:")

h3("Option A â€” Debug mode  (1 persona, 10 steps)")
code("python -X utf8 mobile_main.py --target-url https://www.bajajfinserv.in --auth-mode logged_out --debug")

h3("Option B â€” Full run, all personas")
code("python -X utf8 mobile_main.py --target-url https://www.bajajfinserv.in --auth-mode logged_out")

h3("Option C â€” Single persona")
code("python -X utf8 mobile_main.py --target-url https://www.bajajfinserv.in --persona \"Arjun Mehra\" --debug")

h2("Step 5 â€” Watch the emulator and terminal")
body(
    "The emulator screen will show Chrome opening Bajaj Finserv. "
    "The terminal shows each step OpenAI decides to take â€” taps, swipes, observations, and the final score."
)

h2("Step 6 â€” Find your reports")
body("Mobile reports are written to a separate folder to avoid overwriting web reports:")
bullet("reports/mobile/{persona-slug}/report.md")
bullet("reports/mobile/{persona-slug}/journey_log.json")
bullet("reports/mobile/master_report.md")
bullet("reports/mobile/manifest.json")
bullet("screenshots/mobile/{persona-slug}/step_001.png  (one screenshot per step)")

spacer()
divider()
spacer()

# â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•
# SECTION 6 â€” HOW TO VIEW RESULTS ON THE DASHBOARD
# â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•
h1("6.  How to View Results on the Dashboard")

body(
    "The dashboard is a web interface that displays all audit results visually. "
    "You can use it two ways: via the Flask server (recommended) or by opening the HTML file directly."
)

h2("Method A â€” Via the Flask Server  (recommended)")

h3("Step 1 â€” Start the server")
body("Open a terminal in the project folder and run:")
code("python -X utf8 server.py")
body("You should see:  Running on http://127.0.0.1:5000")

h3("Step 2 â€” Open the dashboard")
body("Open a browser and go to:")
code("http://127.0.0.1:5000")
body(
    "The dashboard loads automatically. It reads all report data from the reports/ folder "
    "and displays every run in the Run History panel on the left."
)

h3("Step 3 â€” Trigger a new audit from the dashboard")
body(
    "Click the New Audit button in the top right of the dashboard. "
    "A dialog appears where you can:"
)
bullet("Choose audit type: Web or App")
bullet("Enter a target URL")
bullet("Select logged_out or logged_in")
bullet("Click Run")
body(
    "The server starts the audit in a background subprocess. "
    "The dashboard shows a running indicator. When complete, the new run appears in the history list."
)

h2("Method B â€” Open the HTML file directly  (offline, no server needed)")

body("Navigate to the dashboard folder in File Explorer:")
code("C:\\Users\\Hiya Bhandari\\Desktop\\Folders\\CX Audit Agent\\dashboard\\index.html")
body(
    "Double-click index.html to open it in your default browser. "
    "In this mode, the dashboard can only read reports that are already in the dashboard/reports/ folder. "
    "To copy the latest reports there, run the following after an audit:"
)
code(
    "xcopy /E /I /Y \"reports\\arjun-mehra\" \"dashboard\\reports\\arjun-mehra\"\n"
    "xcopy /E /I /Y \"reports\\mobile\\arjun-mehra\" \"dashboard\\reports\\arjun-mehra-mobile\""
)

spacer()
divider()
spacer()

# â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•
# SECTION 7 â€” NAVIGATING THE DASHBOARD
# â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•
h1("7.  Navigating the Dashboard")

h2("Run History Panel  (left side)")
body(
    "Lists every completed audit run. Each entry shows the persona name, audit type (Web / App), "
    "auth mode, and overall CX score. Click any run to load its full report on the right."
)
body("Use the filter buttons at the top to show only Web runs or only App runs.")

h2("Report Panel  (right side â€” five tabs)")

h3("Overview Tab")
body(
    "Shows the overall CX score, journey outcome (goal achieved / blocked / max steps reached), "
    "number of steps taken, and pages visited."
)

h3("Persona Tab")
body(
    "Shows who was tested: name, age, occupation, location, device, financial literacy level, "
    "intent, constraints, and success criteria."
)

h3("Scores Tab")
body(
    "Shows the seven CX dimension scores as bars. Each dimension has a numeric score out of 10 "
    "and the reasoning OpenAI used to arrive at that score."
)

h3("Friction & Recommendations Tab")
body(
    "Lists all friction points sorted by severity (red = high, yellow = medium, green = low). "
    "Below that, all recommendations sorted by priority (P1 = immediate action, P2 = important, P3 = nice to have)."
)

h3("Steps Tab")
body(
    "Shows every action the agent took â€” the action type, the target element or URL, "
    "whether it succeeded or failed, and the CX note OpenAI recorded about that step."
)

spacer()
divider()
spacer()

# â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•
# SECTION 8 â€” CONFIGURATION
# â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•
h1("8.  Configuration Files")

h2("config.py  â€”  Main Settings")
body("Controls the web audit behaviour:")
bullet("TARGET_URL  â€”  the website to audit  (default: https://www.bajajfinserv.in)")
bullet("OpenAI_MODEL  â€”  which OpenAI model to use")
bullet("MAX_STEPS  â€”  maximum steps before the agent stops  (default: 25)")
bullet("HEADLESS  â€”  whether to show the browser window or run silently")

h2("config/app_config.yaml  â€”  Mobile Settings")
body("Controls the mobile audit:")
bullet("deviceName  â€”  the emulator to connect to  (emulator-5554)")
bullet("browserName  â€”  set to Chrome for mobile web, or blank for native app")
bullet("appPackage / appActivity  â€”  for native app mode, the Android package and launch activity")
bullet("max_steps, action_delay_ms, element_timeout_secs  â€”  journey runtime settings")

h2("personas/bajaj_personas.md  â€”  Persona Definitions")
body(
    "Each persona is defined in a Markdown block with fields for name, age, gender, occupation, "
    "location, device, financial_literacy, intent, constraints, behaviour, and success_criteria. "
    "Add or edit personas here to change who the agent tests as."
)

h2(".env  â€”  Secrets")
body("Stores the API key. Never commit this file to version control.")
code("OPENAI_API_KEY=sk-your-openai-key-here")

spacer()
divider()
spacer()

# â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•
# SECTION 9 â€” QUICK REFERENCE COMMAND CHEAT SHEET
# â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•
h1("9.  Quick Reference â€” All Commands at a Glance")

h2("Web Audit Commands")
code("python -X utf8 main.py                              # all personas")
code("python -X utf8 main.py --debug                     # 1 persona, 10 steps")
code("python -X utf8 main.py --no-headed                 # no browser window")
code("python -X utf8 main.py --auth-mode logged_in       # logged-in audit")
code("python -X utf8 main.py --persona \"Arjun Mehra\"     # single persona")

h2("Mobile Audit Commands")
code("# Step 1: Start emulator (Android Studio or terminal)")
code("emulator -avd Pixel_6 -no-snapshot-load")
spacer()
code("# Step 2: Start Appium (run in its own terminal window)")
code(
    "set ANDROID_HOME=C:\\Users\\Hiya Bhandari\\AppData\\Local\\Android\\Sdk\n"
    "appium --port 4723 --allow-insecure uiautomator2:chromedriver_autodownload"
)
spacer()
code("# Step 3: Run mobile audit")
code("python -X utf8 mobile_main.py --target-url https://www.bajajfinserv.in --auth-mode logged_out --debug")
code("python -X utf8 mobile_main.py --target-url https://www.bajajfinserv.in --auth-mode logged_out")

h2("Dashboard / Server Commands")
code("python -X utf8 server.py                           # start dashboard server")
body("Then open:  http://127.0.0.1:5000  in a browser")

h2("Verify Emulator Connection")
code("C:\\Users\\Hiya Bhandari\\AppData\\Local\\Android\\Sdk\\platform-tools\\adb.exe devices")

spacer()
divider()
spacer()

# â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•
# SECTION 10 â€” FOLDER STRUCTURE
# â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•
h1("10.  Folder Structure")

code(
    "CX Audit Agent/\n"
    "  main.py                  Web audit entry point\n"
    "  mobile_main.py           Mobile / app audit entry point\n"
    "  server.py                Flask API server for the dashboard\n"
    "  config.py                Global settings and constants\n"
    "  .env                     API key  (create this yourself)\n"
    "  requirements.txt         Python dependencies\n"
    "\n"
    "  agents/                  AI agent logic\n"
    "    controller.py          Main web journey loop\n"
    "    decision_engine.py     OpenAI prompt and action parsing\n"
    "    memory.py              Step-by-step state tracking\n"
    "\n"
    "  browser/                 Playwright browser helpers\n"
    "    controller.py          Click, type, scroll, navigate\n"
    "    state_extractor.py     Extract page text and structure\n"
    "\n"
    "  mobile/\n"
    "    drivers/appium_driver.py   Appium session factory\n"
    "\n"
    "  evaluation/\n"
    "    cx_evaluator.py        Scores journey against CX dimensions\n"
    "\n"
    "  reporting/\n"
    "    report_generator.py    Writes Markdown + JSON reports\n"
    "\n"
    "  parsers/\n"
    "    persona_parser.py      Reads persona definitions from Markdown\n"
    "\n"
    "  personas/\n"
    "    bajaj_personas.md      Persona definitions\n"
    "\n"
    "  config/\n"
    "    app_config.yaml        Mobile / Appium configuration\n"
    "\n"
    "  dashboard/               Frontend web dashboard\n"
    "    index.html\n"
    "    script.js\n"
    "    styles.css\n"
    "    reports/               Report data read by dashboard\n"
    "\n"
    "  reports/                 Web audit output\n"
    "  reports/mobile/          Mobile audit output\n"
    "  screenshots/             Screenshots taken during audits\n"
    "  logs/                    Raw step logs"
)

spacer()
divider()
spacer()

# â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•
# SECTION 11 â€” TROUBLESHOOTING
# â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•
h1("11.  Troubleshooting")

h2("OPENAI_API_KEY not found")
body("Make sure the .env file exists in the project root folder and contains your key on the first line:")
code("OPENAI_API_KEY=sk-your-openai-key-here")

h2("Appium cannot find Android SDK")
body("You must set ANDROID_HOME before starting Appium. Use the set commands shown in Section 5 Step 3. Do not skip this step.")

h2("emulator-5554 not found / adb devices shows nothing")
body("The Android emulator is not running. Start it from Android Studio Device Manager before running the mobile audit.")

h2("Playwright browser not found")
body("Run this once:")
code("playwright install chromium")

h2("UnicodeEncodeError / emoji crashes on Windows")
body("Always run Python with the -X utf8 flag as shown in all commands above. This is already handled â€” do not remove that flag.")

h2("Port 4723 already in use")
body("A previous Appium process is still running. Kill it:")
code("taskkill /F /IM node.exe")
body("Then restart Appium using the command in Section 5 Step 3.")

h2("Dashboard shows no runs")
body(
    "If using the server (Method A), make sure the server is running and you are on http://127.0.0.1:5000. "
    "If using the HTML file directly (Method B), copy the report folders into dashboard/reports/ first."
)

spacer()
divider()
spacer()

# â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•
# FOOTER
# â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•
foot = doc.add_paragraph()
foot.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = foot.add_run("CX Audit Agent  |  Built with OpenAI + Playwright + Appium")
r.font.size = Pt(9)
r.font.color.rgb = RGBColor(0x9C, 0xA3, 0xAF)

# â”€â”€ Save â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€
out = r"C:\Users\Hiya Bhandari\Desktop\Folders\CX Audit Agent\CX Audit Agent â€” Full Guide.docx"
doc.save(out)
print(f"Saved: {out}")


