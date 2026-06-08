"""
Word Report Generator
---------------------
Generates a single combined .docx CX audit report per run.

Structure:
  MAIN BODY
  • Title page
  • Table of Contents (static, pre-filled)
  • Section 1 : Executive Summary
  • Section 2–N: Per-Persona Briefing (one compact page each)
      – Smart persona profile
      – Outcome + score strip
      – TL;DR
      – Compact CX dimension table (all 17 dims, crisp)
      – Brief emotional arc (2–3 lines)
      – Key Takeaways
      – Friction points as crisp bullets (HIGH → MED → LOW)
      – Recommendations table

  ANNEXURES  (detailed reference material)
  • Annexure A: Full Friction & Delight Analysis
  • Annexure B: CX Score Dimensions (full rationale)
  • Annexure C: Emotional Journey Details
  • Annexure D: Complete Journey Steps with Screenshots (2-col layout)

Requires: python-docx, Pillow
"""
from __future__ import annotations

import io
import os
from datetime import datetime, timezone
from pathlib import Path
from typing import Optional
from urllib.parse import urlparse

try:
    from PIL import Image as _PILImage
    _PIL_AVAILABLE = True
except ImportError:
    _PIL_AVAILABLE = False

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor, Emu

import config
from agents.controller import JourneyResult
from evaluation.cx_evaluator import CXAuditResult
from reporting.report_generator import _outcome_label

# ── Colour palette ────────────────────────────────────────────────────────────
_CLR_BLUE        = RGBColor(0x1A, 0x56, 0xDB)
_CLR_DARK_BLUE   = RGBColor(0x1E, 0x40, 0xAF)
_CLR_GREEN       = RGBColor(0x06, 0x5F, 0x46)
_CLR_GREEN_BG    = "E6F9F1"
_CLR_RED         = RGBColor(0x9B, 0x1C, 0x1C)
_CLR_RED_BG      = "FEE2E2"
_CLR_ORANGE_BG   = "FEF3C7"
_CLR_YELLOW_BG   = "FFFBEB"
_CLR_GREY_BG     = "F3F4F6"
_CLR_GREY_TEXT   = RGBColor(0x6B, 0x72, 0x80)
_CLR_WHITE       = RGBColor(0xFF, 0xFF, 0xFF)

_EMOTION_EMOJI = {
    "confident": "😊", "curious": "🤔", "confused": "😕",
    "frustrated": "😤", "anxious": "😰", "overwhelmed": "😵",
    "relieved": "😌", "disappointed": "😞", "suspicious": "🤨",
    "hopeful": "🙂", "reassured": "✅",
}

# Screenshot widths
_SCREENSHOT_FULL_WIDTH   = Inches(5.8)   # annexure full-page
_SCREENSHOT_INLINE_WIDTH = Inches(2.2)   # 2-col journey steps (~35% of page)

# Dimension short descriptions (shown as fine print under dimension name)
_DIM_HINTS: dict[str, str] = {
    "discoverability":        "Can the product be found in 2–3 taps?",
    "content quality":        "Are rates, fees & eligibility shown upfront?",
    "trust":                  "Are trust signals at key anxiety moments?",
    "conversion":             "How smooth is the path to the primary CTA?",
    "emotional experience":   "Does the tone match this persona's mindset?",
    "mobile":                 "Are touch targets & forms mobile-ready?",
    "accessibility":          "Is the UX inclusive for low-literacy users?",
    "error handling":         "Are errors plain-language and recoverable?",
    "system feedback":        "Are loading & confirmation states visible?",
    "micro-copy":             "Are labels, CTAs and copy clear & motivating?",
    "form design":            "Is data collection lean, sequenced, guided?",
    "navigation depth":       "How many taps to goal? Any dead-ends?",
    "personalisation":        "Does the site adapt to this persona's history?",
    "pre-sales":              "Is help & EMI calculator accessible pre-apply?",
    "visual hierarchy":       "Is key info prominent? Is noise minimised?",
    "consistency":            "Are patterns & labels consistent site-wide?",
    "user control":           "Can users undo, pause or exit without loss?",
}


# ── Main entry point ──────────────────────────────────────────────────────────

def generate_word_report(
    results: list[tuple[JourneyResult, CXAuditResult]],
    out_dir: Optional[Path] = None,
) -> Path:
    base     = out_dir or config.REPORTS_DIR
    now_str  = datetime.now(timezone.utc).strftime("%d %b %Y %H%M UTC")
    out_path = base / f"CX_Audit_Report_{now_str.replace(' ', '_').replace(':', '')}.docx"

    doc = Document()
    _configure_document(doc)

    _add_title_page(doc, results, now_str)
    _add_toc(doc, results)
    _add_executive_summary(doc, results)

    for idx, (journey, audit) in enumerate(results, 1):
        doc.add_page_break()
        _add_persona_brief(doc, journey, audit, idx)

    _add_annexures(doc, results)

    doc.save(str(out_path))
    return out_path


# ── Document configuration ────────────────────────────────────────────────────

def _configure_document(doc: Document) -> None:
    for section in doc.sections:
        section.top_margin    = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin   = Inches(1.1)
        section.right_margin  = Inches(1.1)
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10.5)


# ── Title page ────────────────────────────────────────────────────────────────

def _add_title_page(
    doc: Document,
    results: list[tuple[JourneyResult, CXAuditResult]],
    now_str: str,
) -> None:
    doc.add_paragraph()
    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("CX AUDIT REPORT")
    run.bold = True; run.font.size = Pt(28); run.font.color.rgb = _CLR_BLUE

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    site_name = urlparse(config.TARGET_URL).netloc.replace("www.", "") or config.TARGET_URL
    r2 = p2.add_run(f"{site_name}  |  Customer Experience Deep-Dive")
    r2.font.size = Pt(14); r2.font.color.rgb = _CLR_GREY_TEXT

    doc.add_paragraph()

    scores    = [a.overall_score for _, a in results]
    avg_score = sum(scores) / len(scores) if scores else 0

    tbl = doc.add_table(rows=4, cols=2)
    tbl.style = "Table Grid"
    _set_table_width(tbl, Inches(4.5))
    for row_idx, (label, value) in enumerate([
        ("Target URL",       config.TARGET_URL),
        ("Audit Date",       now_str),
        ("Personas Audited", str(len(results))),
        ("Avg CX Score",     f"{avg_score:.1f} / 10  —  {_score_label(avg_score)}"),
    ]):
        row = tbl.rows[row_idx]
        _cell_text(row.cells[0], label, bold=True, bg=_CLR_GREY_BG)
        _cell_text(row.cells[1], value)
    tbl.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()
    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r3 = p3.add_run("Prepared by: CX Audit Agent  (OpenAI + Playwright)")
    r3.font.size = Pt(9); r3.font.color.rgb = _CLR_GREY_TEXT; r3.italic = True

    doc.add_page_break()


# ── Table of Contents (static, pre-filled) ───────────────────────────────────

def _add_toc(doc: Document, results: list[tuple[JourneyResult, CXAuditResult]]) -> None:
    h = doc.add_heading("Table of Contents", level=1)
    h.runs[0].font.color.rgb = _CLR_BLUE

    entries: list[tuple[str, str, bool]] = [
        ("1", "Executive Summary", False),
    ]
    for idx, (journey, _) in enumerate(results, 1):
        entries.append((str(idx + 1), f"Persona: {journey.persona.name}", False))
    entries.extend([
        ("A", "Annexure A — Full Friction & Delight Analysis", True),
        ("B", "Annexure B — CX Score Dimensions",             True),
        ("C", "Annexure C — Emotional Journey Details",       True),
        ("D", "Annexure D — Complete Journey Log with Screenshots", True),
    ])

    tbl = doc.add_table(rows=len(entries), cols=2)
    tbl.style = "Table Grid"
    _set_table_width(tbl, Inches(5.5))

    for i, (num, title, is_annex) in enumerate(entries):
        row = tbl.rows[i]
        bg = _CLR_GREY_BG if (i % 2 == 0) else "FFFFFF"
        _cell_text(row.cells[0], num,   bold=True,     bg=bg, font_size=10)
        _cell_text(row.cells[1], title, bold=is_annex, bg=bg, font_size=10)
        _set_cell_width_inches(row.cells[0], 0.6)
        _set_cell_width_inches(row.cells[1], 4.9)

    doc.add_page_break()


# ── Executive Summary ─────────────────────────────────────────────────────────

def _add_executive_summary(
    doc: Document,
    results: list[tuple[JourneyResult, CXAuditResult]],
) -> None:
    scores      = [a.overall_score for _, a in results]
    avg_score   = sum(scores) / len(scores) if scores else 0
    best_score  = max(scores) if scores else 0
    worst_score = min(scores) if scores else 0
    success_count = sum(1 for j, _ in results if j.success)

    _heading1(doc, "1. Executive Summary")

    stats_tbl = doc.add_table(rows=1, cols=4)
    stats_tbl.style = "Table Grid"
    _set_table_width(stats_tbl, Inches(6.3))
    for cell, label, value in zip(
        stats_tbl.rows[0].cells,
        ["Overall Avg Score", "Best Experience", "Worst Experience", "Goals Achieved"],
        [f"{avg_score:.1f}/10", f"{best_score:.1f}/10", f"{worst_score:.1f}/10",
         f"{success_count}/{len(results)}"],
    ):
        _cell_stat(cell, label, value)
    doc.add_paragraph()

    _heading2(doc, "Persona Verdicts at a Glance")
    vtbl = doc.add_table(rows=1 + len(results), cols=4)
    vtbl.style = "Table Grid"
    _set_table_width(vtbl, Inches(6.3))
    for cell, hdr in zip(vtbl.rows[0].cells, ["Persona", "Score", "Outcome", "One-line Verdict"]):
        _cell_text(cell, hdr, bold=True, bg=_CLR_GREY_BG)
    for i, (journey, audit) in enumerate(results, 1):
        row = vtbl.rows[i]
        outcome = "✅ Achieved" if journey.success else f"🚫 {journey.terminal_reason.replace('_',' ').title()}"
        _cell_text(row.cells[0], journey.persona.name, bold=True)
        _cell_score(row.cells[1], audit.overall_score)
        _cell_text(row.cells[2], outcome)
        _cell_text(row.cells[3], audit.tldr or audit.journey_verdict)
    doc.add_paragraph()

    _heading2(doc, "Critical Issues — High Severity")
    high_issues = [
        (j.persona.name, fp)
        for j, a in results
        for fp in a.friction_points
        if fp.severity.lower() == "high"
    ]
    if high_issues:
        itbl = doc.add_table(rows=1 + len(high_issues), cols=3)
        itbl.style = "Table Grid"
        _set_table_width(itbl, Inches(6.3))
        for cell, hdr in zip(itbl.rows[0].cells, ["Persona", "Location", "Issue"]):
            _cell_text(cell, hdr, bold=True, bg=_CLR_RED_BG)
        for i, (pname, fp) in enumerate(high_issues, 1):
            row = itbl.rows[i]
            _cell_text(row.cells[0], pname, bold=True)
            _cell_text(row.cells[1], fp.location)
            _cell_text(row.cells[2], fp.description)
    else:
        doc.add_paragraph("No critical (high-severity) issues identified.")
    doc.add_paragraph()

    _heading2(doc, "P1 Recommendations — Fix Immediately")
    p1_recs = [
        (j.persona.name, r)
        for j, a in results
        for r in a.recommendations
        if r.priority.upper() == "P1"
    ]
    if p1_recs:
        rtbl = doc.add_table(rows=1 + len(p1_recs), cols=3)
        rtbl.style = "Table Grid"
        _set_table_width(rtbl, Inches(6.3))
        for cell, hdr in zip(rtbl.rows[0].cells, ["Persona", "Area", "Action"]):
            _cell_text(cell, hdr, bold=True, bg=_CLR_ORANGE_BG)
        for i, (pname, rec) in enumerate(p1_recs, 1):
            row = rtbl.rows[i]
            _cell_text(row.cells[0], pname, bold=True)
            _cell_text(row.cells[1], rec.area)
            _cell_text(row.cells[2], rec.action)
    else:
        doc.add_paragraph("No P1 recommendations generated.")
    # No page break here — the loop adds one before each persona section


# ── Per-persona briefing ──────────────────────────────────────────────────────

def _add_persona_brief(
    doc: Document,
    journey: JourneyResult,
    audit: CXAuditResult,
    section_num: int,
) -> None:
    persona = journey.persona
    _heading1(doc, f"{section_num + 1}. Persona: {persona.name}")

    # 1. Smart persona profile (no rigid table)
    _add_persona_profile_smart(doc, persona)

    # 2. Outcome + CX score strip
    outcome_label = f"✅ {_outcome_label(journey)}" if journey.success else f"⏹ {_outcome_label(journey)}"
    otbl = doc.add_table(rows=1, cols=3)
    otbl.style = "Table Grid"
    _set_table_width(otbl, Inches(6.3))
    _cell_stat(otbl.rows[0].cells[0], "Outcome",  outcome_label)
    _cell_stat(otbl.rows[0].cells[1], "CX Score", f"{audit.overall_score:.1f} / 10  {_score_label(audit.overall_score)}")
    _cell_stat(otbl.rows[0].cells[2], "Steps",    f"{journey.memory.step_count}  ({journey.memory.failure_count} failures)")
    doc.add_paragraph()

    # 3. TL;DR
    if audit.tldr:
        _callout_box(doc, "TL;DR", audit.tldr, bg_hex=_CLR_GREY_BG, bold_body=True)

    # 4. Compact CX dimension table
    _add_compact_dimension_table(doc, audit)

    # 5. Brief emotional arc (2–3 lines max)
    _add_brief_emotion_journey(doc, audit)

    # 6. Key Takeaways
    if audit.key_takeaways:
        _heading2(doc, "Key Takeaways")
        for kt in audit.key_takeaways[:5]:
            p = doc.add_paragraph(style="List Bullet")
            p.add_run(kt)
        doc.add_paragraph()

    # 7. Friction points as crisp bullets (HIGH → MEDIUM → LOW)
    _add_friction_bullets_section(doc, audit)

    # 8. Recommendations table
    _heading2(doc, "Recommendations")
    if audit.recommendations:
        recs_sorted = sorted(audit.recommendations, key=lambda r: r.priority.upper())
        rtbl = doc.add_table(rows=1 + len(recs_sorted), cols=4)
        rtbl.style = "Table Grid"
        _set_table_width(rtbl, Inches(6.3))
        for cell, hdr in zip(rtbl.rows[0].cells, ["Priority", "Area", "Action", "Expected Impact"]):
            _cell_text(cell, hdr, bold=True, bg=_CLR_GREY_BG)
        _priority_bg = {"P1": _CLR_RED_BG, "P2": _CLR_ORANGE_BG, "P3": _CLR_YELLOW_BG}
        for i, rec in enumerate(recs_sorted, 1):
            row  = rtbl.rows[i]
            bg   = _priority_bg.get(rec.priority.upper(), _CLR_GREY_BG)
            _cell_text(row.cells[0], rec.priority.upper(), bold=True, bg=bg)
            _cell_text(row.cells[1], rec.area)
            _cell_text(row.cells[2], rec.action)
            _cell_text(row.cells[3], rec.expected_impact or "")
    else:
        doc.add_paragraph("No recommendations generated.")


# ── Persona profile — smart inline format ────────────────────────────────────

def _add_persona_profile_smart(doc: Document, persona) -> None:
    # Line 1: demographic + occupation + location + device + literacy
    chips: list[str] = []
    demo_parts = []
    if persona.age:
        demo_parts.append(str(persona.age))
    if persona.gender:
        demo_parts.append(persona.gender)
    if demo_parts:
        chips.append(", ".join(demo_parts))
    if persona.occupation:
        chips.append(persona.occupation)
    if persona.location:
        chips.append(persona.location)
    if persona.device:
        chips.append(f"Device: {persona.device}")
    if persona.financial_literacy:
        chips.append(f"Fin. Literacy: {persona.financial_literacy}")

    if chips:
        p = doc.add_paragraph()
        r = p.add_run("  ·  ".join(chips))
        r.font.size = Pt(10)
        r.font.color.rgb = _CLR_GREY_TEXT

    # Goal
    if persona.intent:
        p2 = doc.add_paragraph()
        rl = p2.add_run("Goal: "); rl.bold = True; rl.font.size = Pt(10)
        rv = p2.add_run(persona.intent); rv.font.size = Pt(10)

    # Constraints
    if persona.constraints:
        p3 = doc.add_paragraph()
        rl = p3.add_run("Constraints: "); rl.bold = True; rl.font.size = Pt(10)
        rv = p3.add_run(persona.constraints); rv.font.size = Pt(10)

    # Behaviour / success criteria if present
    if getattr(persona, "behaviour", None):
        p4 = doc.add_paragraph()
        rl = p4.add_run("Behaviour: "); rl.bold = True; rl.font.size = Pt(10)
        rv = p4.add_run(persona.behaviour); rv.font.size = Pt(10)

    # Any extra raw attributes not captured above
    known_keys = {
        "age", "gender", "occupation", "location", "device",
        "financial literacy", "intent", "constraints", "behaviour",
        "behavior", "success criteria", "background",
    }
    extras = [
        (k.title(), v) for k, v in persona.raw_attributes.items()
        if k.lower() not in known_keys and v
    ]
    if extras:
        p5 = doc.add_paragraph()
        p5.add_run("  ·  ".join(f"{k}: {v}" for k, v in extras[:4])).font.size = Pt(9.5)

    doc.add_paragraph()


# ── Compact CX dimension table ────────────────────────────────────────────────

def _add_compact_dimension_table(doc: Document, audit: CXAuditResult) -> None:
    if not audit.dimensions:
        return
    _heading2(doc, "CX Performance — Dimension Scores")

    dtbl = doc.add_table(rows=1 + len(audit.dimensions), cols=3)
    dtbl.style = "Table Grid"
    _set_table_width(dtbl, Inches(6.3))

    for cell, hdr in zip(dtbl.rows[0].cells, ["CX Parameter", "Score", "Key Finding"]):
        _cell_text(cell, hdr, bold=True, bg=_CLR_GREY_BG)
    _set_cell_width_inches(dtbl.rows[0].cells[0], 2.1)
    _set_cell_width_inches(dtbl.rows[0].cells[1], 0.9)
    _set_cell_width_inches(dtbl.rows[0].cells[2], 3.3)

    for i, dim in enumerate(audit.dimensions, 1):
        row = dtbl.rows[i]
        bg  = "FFFFFF" if i % 2 == 0 else "F9FAFB"

        # Column 1: dimension name + hint as fine print
        c0 = row.cells[0]
        _shade_cell(c0, bg)
        p = c0.paragraphs[0]; p.clear()
        r_name = p.add_run(dim.name)
        r_name.bold = True; r_name.font.size = Pt(9.5)
        hint = _get_dim_hint(dim.name)
        if hint:
            r_hint = p.add_run(f"\n{hint}")
            r_hint.font.size = Pt(7.5)
            r_hint.font.color.rgb = _CLR_GREY_TEXT
            r_hint.italic = True

        # Column 2: score with colour
        _cell_score(row.cells[1], dim.score)
        _shade_cell(row.cells[1], bg)

        _cell_text(row.cells[2], dim.rationale, bg=bg, font_size=9.5)

        _set_cell_width_inches(row.cells[0], 2.1)
        _set_cell_width_inches(row.cells[1], 0.9)
        _set_cell_width_inches(row.cells[2], 3.3)

    doc.add_paragraph()


# ── Brief emotional arc ───────────────────────────────────────────────────────

def _add_brief_emotion_journey(doc: Document, audit: CXAuditResult) -> None:
    _heading2(doc, "Emotional Arc")

    narrative = (audit.persona_emotional_narrative or "").strip()
    if narrative:
        p = doc.add_paragraph()
        r = p.add_run(narrative)
        r.italic = True; r.font.size = Pt(10)
        doc.add_paragraph()
        return

    # Fallback: top 4 emotional moments as inline arc
    if audit.emotional_journey:
        arc_parts = []
        for em in audit.emotional_journey[:4]:
            emoji = _EMOTION_EMOJI.get(em.emotion.lower(), "")
            arc_parts.append(f"{em.stage}: {emoji} {em.emotion.title()}")
        p = doc.add_paragraph()
        p.add_run("  →  ".join(arc_parts)).font.size = Pt(10)
        doc.add_paragraph()


# ── Friction points as crisp bullets ─────────────────────────────────────────

def _add_friction_bullets_section(doc: Document, audit: CXAuditResult) -> None:
    all_fp = audit.friction_points
    if not all_fp:
        return

    high_fp   = [fp for fp in all_fp if fp.severity.lower() == "high"]
    medium_fp = [fp for fp in all_fp if fp.severity.lower() == "medium"]
    low_fp    = [fp for fp in all_fp if fp.severity.lower() == "low"]

    _heading2(doc, "Friction Points")

    for label, icon, items in [
        ("High Severity",   "🔴", high_fp),
        ("Medium Severity", "🟡", medium_fp),
        ("Low Severity",    "🟢", low_fp),
    ]:
        if not items:
            continue
        p_hdr = doc.add_paragraph()
        r_hdr = p_hdr.add_run(f"{icon}  {label}")
        r_hdr.bold = True; r_hdr.font.size = Pt(10)
        for fp in items:
            p = doc.add_paragraph(style="List Bullet")
            p.add_run(f"{fp.location}: ").bold = True
            p.add_run(fp.description)
            p.paragraph_format.left_indent = Inches(0.2)

    doc.add_paragraph()


# ── Annexures ─────────────────────────────────────────────────────────────────

def _add_annexures(
    doc: Document,
    results: list[tuple[JourneyResult, CXAuditResult]],
) -> None:
    # ── A: Full Friction & Delight ─────────────────────────────────────────
    doc.add_page_break()
    _heading1(doc, "Annexure A — Full Friction & Delight Analysis")

    for journey, audit in results:
        _heading2(doc, journey.persona.name)

        if audit.delight_factors:
            p = doc.add_paragraph()
            p.add_run("Delight Factors").bold = True
            for df in audit.delight_factors:
                _callout_box(doc, "✨", df, bg_hex=_CLR_GREEN_BG)

        high_fp   = [fp for fp in audit.friction_points if fp.severity.lower() == "high"]
        medium_fp = [fp for fp in audit.friction_points if fp.severity.lower() == "medium"]
        low_fp    = [fp for fp in audit.friction_points if fp.severity.lower() == "low"]

        if not audit.friction_points:
            doc.add_paragraph("No friction points identified.")
        else:
            for label_text, color, items in [
                ("HIGH SEVERITY",   _CLR_RED,                      high_fp),
                ("MEDIUM SEVERITY", RGBColor(0x92, 0x40, 0x0E),   medium_fp),
                ("LOW SEVERITY",    _CLR_GREY_TEXT,                 low_fp),
            ]:
                if not items:
                    continue
                p = doc.add_paragraph()
                r = p.add_run(label_text); r.bold = True; r.font.color.rgb = color
                bg_map   = {"HIGH SEVERITY": _CLR_RED_BG, "MEDIUM SEVERITY": _CLR_ORANGE_BG, "LOW SEVERITY": _CLR_YELLOW_BG}
                icon_map = {"HIGH SEVERITY": "🔴", "MEDIUM SEVERITY": "🟡", "LOW SEVERITY": "🟢"}
                bg   = bg_map.get(label_text, _CLR_GREY_BG)
                icon = icon_map.get(label_text, "•")
                for fp in items:
                    _callout_box(
                        doc,
                        f"{icon}  {fp.location}",
                        f"{fp.description}\nUser Impact: {fp.impact}",
                        bg_hex=bg,
                    )
        doc.add_paragraph()

    # ── B: CX Score Dimensions (full rationale) ────────────────────────────
    doc.add_page_break()
    _heading1(doc, "Annexure B — CX Score Dimensions (Full Detail)")

    for journey, audit in results:
        _heading2(doc, journey.persona.name)
        if audit.dimensions:
            dtbl = doc.add_table(rows=1 + len(audit.dimensions), cols=3)
            dtbl.style = "Table Grid"
            _set_table_width(dtbl, Inches(6.3))
            for cell, hdr in zip(dtbl.rows[0].cells, ["Dimension", "Score", "Rationale"]):
                _cell_text(cell, hdr, bold=True, bg=_CLR_GREY_BG)
            for i, dim in enumerate(audit.dimensions, 1):
                row = dtbl.rows[i]
                _cell_text(row.cells[0], dim.name, bold=True)
                _cell_score(row.cells[1], dim.score)
                _cell_text(row.cells[2], dim.rationale)
        doc.add_paragraph()

    # ── C: Emotional Journey ───────────────────────────────────────────────
    doc.add_page_break()
    _heading1(doc, "Annexure C — Emotional Journey Details")

    for journey, audit in results:
        _heading2(doc, journey.persona.name)
        if audit.persona_emotional_narrative:
            np = doc.add_paragraph()
            nr = np.add_run(audit.persona_emotional_narrative)
            nr.italic = True; nr.font.size = Pt(10.5)
            doc.add_paragraph()
        if audit.emotional_journey:
            etbl = doc.add_table(rows=1 + len(audit.emotional_journey), cols=3)
            etbl.style = "Table Grid"
            _set_table_width(etbl, Inches(6.3))
            for cell, hdr in zip(etbl.rows[0].cells, ["Stage", "Emotion", "What Triggered It"]):
                _cell_text(cell, hdr, bold=True, bg=_CLR_GREY_BG)
            for i, em in enumerate(audit.emotional_journey, 1):
                row = etbl.rows[i]
                _cell_text(row.cells[0], em.stage)
                emoji = _EMOTION_EMOJI.get(em.emotion.lower(), "")
                _cell_text(row.cells[1], f"{emoji} {em.emotion.title()}")
                _cell_text(row.cells[2], em.trigger)
        doc.add_paragraph()

    # ── D: Complete Journey Steps with Screenshots (2-col layout) ─────────
    doc.add_page_break()
    _heading1(doc, "Annexure D — Complete Journey Log with Screenshots")

    last_persona_idx = len(results) - 1
    for p_idx, (journey, audit) in enumerate(results):
        doc.add_page_break()
        _heading2(doc, f"{journey.persona.name}  —  Step-by-Step Journey")
        _add_journey_steps(doc, journey, is_last=(p_idx == last_persona_idx))


# ── User journey steps (2-column: text | screenshot) ─────────────────────────

def _add_journey_steps(doc: Document, journey: JourneyResult, is_last: bool = False) -> None:
    steps = journey.memory.steps
    for s_idx, step in enumerate(steps):
        status       = "✅" if step.success else "❌"
        action_upper = step.action.upper()
        target_short = (step.target or step.url or "—")[:80]

        h = doc.add_heading(level=3)
        h.clear()
        run = h.add_run(
            f"Step {step.step_number + 1}  {status}  [{action_upper}]  {target_short}"
        )
        run.font.size = Pt(10.5); run.font.color.rgb = _CLR_DARK_BLUE

        # Collect metadata rows
        meta_rows: list[tuple[str, str]] = []
        if step.url:
            meta_rows.append(("URL", step.url[:100]))
        if step.value:
            meta_rows.append(("Typed", step.value))
        if step.emotion:
            emoji = _EMOTION_EMOJI.get(step.emotion.lower(), "")
            meta_rows.append(("Emotion", f"{emoji} {step.emotion.title()}"))
        if step.error:
            meta_rows.append(("Error", step.error))
        if step.cx_note:
            meta_rows.append(("CX Note", step.cx_note))

        has_screenshot = bool(step.screenshot and Path(step.screenshot).exists()
                              and Path(step.screenshot).stat().st_size > 0)

        if has_screenshot:
            # 2-column table: metadata left, screenshot right
            layout_tbl = doc.add_table(rows=1, cols=2)
            layout_tbl.style = "Table Grid"
            _set_table_width(layout_tbl, Inches(6.3))

            left_cell  = layout_tbl.rows[0].cells[0]
            right_cell = layout_tbl.rows[0].cells[1]
            _set_cell_width_inches(left_cell,  3.9)
            _set_cell_width_inches(right_cell, 2.4)

            # Fill left cell with metadata
            if meta_rows:
                lp = left_cell.paragraphs[0]
                lp.clear()
                for label, value in meta_rows:
                    r_lbl = lp.add_run(f"{label}: ")
                    r_lbl.bold = True; r_lbl.font.size = Pt(9)
                    r_val = lp.add_run(value + "\n")
                    r_val.font.size = Pt(9)
            else:
                _cell_text(left_cell, "—", font_size=9)

            # Fill right cell with screenshot
            right_para = right_cell.paragraphs[0]
            right_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            buf = _compress_screenshot_for_doc(step.screenshot, max_width=800)
            try:
                run_pic = right_para.add_run()
                if buf:
                    run_pic.add_picture(buf, width=_SCREENSHOT_INLINE_WIDTH)
                else:
                    run_pic.add_picture(str(Path(step.screenshot)), width=_SCREENSHOT_INLINE_WIDTH)
            except Exception:
                _cell_text(right_cell, f"[Screenshot: {Path(step.screenshot).name}]", font_size=8)
        else:
            # No screenshot — plain metadata table
            if meta_rows:
                mtbl = doc.add_table(rows=len(meta_rows), cols=2)
                mtbl.style = "Table Grid"
                _set_table_width(mtbl, Inches(5.5))
                for row, (label, value) in zip(mtbl.rows, meta_rows):
                    _cell_text(row.cells[0], label, bold=True, bg=_CLR_GREY_BG, font_size=9)
                    _cell_text(row.cells[1], value, font_size=9)

        # Small gap between steps (skip after the very last step of the last persona)
        if not (is_last and s_idx == len(steps) - 1):
            sp = doc.add_paragraph()
            sp.paragraph_format.space_before = Pt(4)
            sp.paragraph_format.space_after  = Pt(4)


# ── Helpers ───────────────────────────────────────────────────────────────────

def _smart_truncate(text: str, max_words: int = 20) -> str:
    """Truncate at word boundary, never mid-sentence. Appends ellipsis only when cut."""
    if not text:
        return ""
    words = text.split()
    if len(words) <= max_words:
        return text
    return " ".join(words[:max_words]).rstrip(".,;:") + "…"


def _get_dim_hint(name: str) -> str:
    name_lower = name.lower()
    for key, hint in _DIM_HINTS.items():
        if key in name_lower:
            return hint
    return ""


def _heading1(doc: Document, text: str) -> None:
    h = doc.add_heading(text, level=1)
    for run in h.runs:
        run.font.color.rgb = _CLR_BLUE
        run.font.size = Pt(16)


def _heading2(doc: Document, text: str) -> None:
    h = doc.add_heading(text, level=2)
    for run in h.runs:
        run.font.color.rgb = _CLR_DARK_BLUE
        run.font.size = Pt(12)


def _score_label(score: float) -> str:
    if score >= 8.0: return "🟢 Strong"
    if score >= 6.0: return "🟡 Moderate"
    if score >= 4.0: return "🟠 Weak"
    return "🔴 Critical"


def _score_bar(score: float, width: int = 12) -> str:
    filled = round((score / 10) * width)
    return "█" * filled + "░" * (width - filled)


def _cell_text(
    cell,
    text: str,
    bold: bool = False,
    bg: Optional[str] = None,
    font_size: float = 10,
) -> None:
    p = cell.paragraphs[0]; p.clear()
    run = p.add_run(text)
    run.bold = bold; run.font.size = Pt(font_size)
    if bg:
        _shade_cell(cell, bg)


def _cell_score(cell, score: float) -> None:
    p = cell.paragraphs[0]; p.clear()
    bar = _score_bar(score, 10)
    run = p.add_run(f"{score:.1f}/10\n{bar}")
    run.font.size = Pt(9)
    if score >= 7:
        run.font.color.rgb = RGBColor(0x06, 0x5F, 0x46)
    elif score >= 5:
        run.font.color.rgb = RGBColor(0x92, 0x40, 0x0E)
    else:
        run.font.color.rgb = RGBColor(0x99, 0x1B, 0x1B)


def _cell_stat(cell, label: str, value: str) -> None:
    p = cell.paragraphs[0]; p.clear()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_label = p.add_run(label + "\n"); r_label.font.size = Pt(8.5); r_label.font.color.rgb = _CLR_GREY_TEXT
    r_val = p.add_run(value); r_val.bold = True; r_val.font.size = Pt(11)
    _shade_cell(cell, _CLR_GREY_BG)


def _callout_box(
    doc: Document,
    title: str,
    body: str,
    bg_hex: str = _CLR_GREY_BG,
    bold_body: bool = False,
    font_size: float = 10,
) -> None:
    tbl = doc.add_table(rows=1, cols=1)
    tbl.style = "Table Grid"
    _set_table_width(tbl, Inches(6.3))
    cell = tbl.rows[0].cells[0]
    _shade_cell(cell, bg_hex)
    p = cell.paragraphs[0]
    if title:
        r_title = p.add_run(title + "\n"); r_title.bold = True; r_title.font.size = Pt(font_size)
    if body:
        r_body = p.add_run(body); r_body.bold = bold_body; r_body.font.size = Pt(font_size)
    doc.add_paragraph()


def _compress_screenshot_for_doc(path_str: str, max_width: int = 800) -> Optional[io.BytesIO]:
    if not _PIL_AVAILABLE:
        return None
    try:
        p = Path(path_str)
        if not p.exists() or p.stat().st_size == 0:
            return None
        with _PILImage.open(p) as img:
            img = img.convert("RGB")
            if img.width > max_width:
                ratio = max_width / img.width
                img   = img.resize((max_width, int(img.height * ratio)), _PILImage.LANCZOS)
        buf = io.BytesIO()
        img.save(buf, format="JPEG", quality=70, optimize=True)
        buf.seek(0)
        return buf
    except Exception:
        return None


def _embed_screenshot(doc: Document, path: str, label: str = "",
                      width: Inches = _SCREENSHOT_FULL_WIDTH) -> None:
    try:
        p = Path(path)
        if not p.exists() or p.stat().st_size == 0:
            return
        cap = doc.add_paragraph()
        cr = cap.add_run(f"📸  {label}  —  {p.name}")
        cr.font.size = Pt(8.5); cr.font.color.rgb = _CLR_GREY_TEXT; cr.italic = True
        buf = _compress_screenshot_for_doc(path)
        if buf:
            doc.add_picture(buf, width=width)
        else:
            doc.add_picture(str(p), width=width)
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    except Exception:
        pass


def _shade_cell(cell, hex_color: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tc_pr.append(shd)


def _set_table_width(tbl, width: Emu) -> None:
    tbl_pr = tbl._tbl.tblPr
    tbl_w  = OxmlElement("w:tblW")
    tbl_w.set(qn("w:w"),    str(int(width.inches * 1440)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_pr.append(tbl_w)


def _set_cell_width_inches(cell, width_inches: float) -> None:
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcW  = OxmlElement("w:tcW")
    tcW.set(qn("w:w"),    str(int(width_inches * 1440)))
    tcW.set(qn("w:type"), "dxa")
    tcPr.append(tcW)
